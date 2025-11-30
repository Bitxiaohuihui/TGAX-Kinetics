# -*- coding: utf-8 -*-
"""
TGAX Kinetics – Comprehensive Kinetic Analysis & Lifetime Prediction
Version 1.0
"""

# ───────── Crash Handler ─────────
import sys, traceback, datetime, pathlib, io, json, tkinter as _tk, tkinter.messagebox as _mb
def _fatal(et, ev, tb):
    log = pathlib.Path(__file__).with_suffix(".error.log")
    with log.open("a", encoding="utf-8") as f:
        f.write(f"\n=== {datetime.datetime.now()} ===\n")
        traceback.print_exception(et, ev, tb, file=f)
    _tk.Tk().withdraw()
    _mb.showerror("Fatal Error", f"{ev}\n\nDetails have been saved to {log.name}")
    sys.exit(1)
sys.excepthook = _fatal

# ───────── Standard Libraries ─────────
import math, re, os, textwrap, ctypes
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, simpledialog
import warnings

# --- START: MODIFICATION (Threading) ---
# Import threading and queue for background processing
import threading
from queue import Queue, Empty
# --- END: MODIFICATION (Threading) ---

# ───────── High-DPI Awareness for Windows ─────────
try:
    if sys.platform == "win32":
        ctypes.windll.shcore.SetProcessDpiAwareness(2) # Enable Per-Monitor v2 DPI Awareness
except (AttributeError, TypeError):
    pass

# ───────── Third-Party Libraries ─────────
import numpy as np, pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from scipy.optimize import least_squares, minimize_scalar
from scipy import integrate
from scipy.integrate import quad, solve_ivp
from scipy.signal import savgol_filter
from scipy.interpolate import PchipInterpolator
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ───────── Matplotlib Style (High-Contrast Journal Style) ─────────
plt.rcParams.update({
    "legend.labelspacing": 0.1,
    "legend.handlelength": 0.8,
    "legend.handletextpad": 0.2,
    "legend.columnspacing": 0.5,
    "legend.borderpad": 0,
    "legend.borderaxespad": 0.3,
    "font.family": "Arial", "font.size": 10, "axes.titlesize": 12, "axes.labelsize": 10,
    "axes.linewidth": 1.0, "lines.linewidth": 1.5, "lines.markersize": 5, "legend.fontsize": 7,
    "legend.frameon": False, "xtick.major.size": 4, "ytick.major.size": 4, "xtick.major.width": 1.0,
    "ytick.major.width": 1.0, "xtick.direction": "in", 
})
# REVISION: High-contrast color palette for visual distinction
NATURE_COLORS = [
    '#EE7733', '#0077BB', '#33BBEE', '#EE3377', '#CC3311', '#009988', '#BBBBBB',
    '#77AADD', '#99DDFF', '#44BB99', '#DDDD77', '#FFAABB', '#EEDD88', '#AAAA00'
]
plt.rcParams['axes.prop_cycle'] = plt.cycler(color=NATURE_COLORS)

# ───────── Constants ─────────
R, EPS = 8.314462618, 1e-6 # Ideal gas constant (J·mol⁻¹·K⁻¹), Epsilon for float comparisons

# --- START: Mechanism Plot Integration ---
# Step 1: Create the Ideal Mechanism Model Library (EXPANDED TO >60 MODELS)
# This library has been expanded to include over 60 unique and variant models from literature tables.
IDEAL_MODELS = {
    # --- Reaction Order Models (Fn) ---
    'F0': ('Zero-Order', lambda a: np.ones_like(a)),
    'F0.5': ('Reaction Order (n=0.5)', lambda a: (1 - a)**0.5),
    'F1': ('First-Order / Mampel', lambda a: 1 - a),
    'F1.5': ('Reaction Order (n=1.5)', lambda a: (1 - a)**1.5),
    'F2': ('Second-Order', lambda a: (1 - a)**2),
    'F2.5': ('Reaction Order (n=2.5)', lambda a: (1 - a)**2.5),
    'F3': ('Third-Order', lambda a: (1 - a)**3),
    'F4': ('Fourth-Order', lambda a: (1 - a)**4),

    # --- Nucleation and Growth Models (Avrami-Erofeev, An) ---
    'A1': ('Avrami-Erofeev (n=1, equiv. F1)', lambda a: 1 - a), # A1 is mathematically equivalent to F1
    'A1.5': ('Avrami-Erofeev (n=1.5)', lambda a: 1.5 * (1 - a) * (-np.log(1 - a))**(1/3)),
    'A2': ('Avrami-Erofeev (n=2)', lambda a: 2 * (1 - a) * (-np.log(1 - a))**0.5),
    'A2.5': ('Avrami-Erofeev (n=2.5)', lambda a: 2.5 * (1 - a) * (-np.log(1 - a))**(3/5)),
    'A3': ('Avrami-Erofeev (n=3)', lambda a: 3 * (1 - a) * (-np.log(1 - a))**(2/3)),
    'A3.5': ('Avrami-Erofeev (n=3.5)', lambda a: 3.5 * (1 - a) * (-np.log(1 - a))**(2.5/3.5)),
    'A4': ('Avrami-Erofeev (n=4)', lambda a: 4 * (1 - a) * (-np.log(1 - a))**(3/4)),
    'A5': ('Avrami-Erofeev (n=5)', lambda a: 5 * (1 - a) * (-np.log(1 - a))**(4/5)),

    # --- Geometrical Contraction Models (Rn) & Shrinking Core (SC) ---
    'R1': ('Contracting Line / Zero-Order', lambda a: np.ones_like(a)),
    'R2': ('Contracting Area / Cylinder', lambda a: 2 * (1 - a)**0.5),
    'R3': ('Contracting Volume / Sphere', lambda a: 3 * (1 - a)**(2/3)),
    'SC-RC': ('Shrinking Core (Reaction)', lambda a: 3 * (1 - a)**(2/3)), # Same as R3

    # --- Power Law Models (Pn) ---
    'P1/5': ('Power Law (n=1/5)', lambda a: 5 * a**(4/5)),
    'P1/4': ('Power Law (n=1/4)', lambda a: 4 * a**(3/4)),
    'P1/3': ('Power Law (n=1/3)', lambda a: 3 * a**(2/3)),
    'P3/5': ('Power Law (n=3/5)', lambda a: (5/3) * a**(2/5)),
    'P1/2': ('Power Law (n=1/2)', lambda a: 2 * a**0.5),
    'P2/3': ('Power Law (n=2/3)', lambda a: (3/2) * a**(1/3)),
    'P3/4': ('Power Law (n=3/4)', lambda a: (4/3) * a**(1/4)),
    'P4/5': ('Power Law (n=4/5)', lambda a: (5/4) * a**(1/5)),
    'P3/2': ('Power Law (n=3/2)', lambda a: (2/3) * a**(-0.5)),
    'P2': ('Power Law (n=2)', lambda a: 0.5 * a**(-1)),

    # --- Diffusion Models (Dn) & Shrinking Core (SC) ---
    'D1': ('1D Diffusion', lambda a: 0.5 / a),
    'D2': ('2D Diffusion (Valensi)', lambda a: 1 / (-np.log(1 - a))),
    'D3-J': ('3D Diffusion (Jander)', lambda a: (3/2) * (1 - a)**(2/3) / (1 - (1 - a)**(1/3))),
    'D4-GB': ('3D Diffusion (Ginstling-Brounshtein)', lambda a: (3/2) / ((1 - a)**(-1/3) - 1)),
    'D5-ZL': ('3D Diffusion (Zhuravlev-Lesokine)', lambda a: (3/2) / (np.power(1 + a, 1/3) - 1)),
    'D6': ('3D Diffusion (Variant D3d)', lambda a: (3/2) * (1-a)**(1/3) / (1 - (1-a)**(1/3))),
    'D7': ('3D Diffusion (Variant D3g)', lambda a: 6*(1-a)*(-np.log(1-a))**0.5 / (1-(1-a)**(1/3))),
    'D-AJ': ('Anti-Jander', lambda a: (3/2) * (1 - (1-a)**(1/3))**(-1)),
    'D-AGB': ('Anti-Ginstling-Brounshtein', lambda a: (3/2) * (1-a)**(-1/3)),
    'D10-Ash': ('Shrinking Core (Ash Diffusion)', lambda a: (3/2) / ((1 - a)**(-1/3) - 1)), # Same shape as D4-GB
    'SC-DA': ('Shrinking Core (Ash Diffusion)', lambda a: (3/2) / ((1 - a)**(-1/3) - 1)), # Alias for D10-Ash
    'D11-Film': ('Shrinking Core (Film Diffusion)', lambda a: np.ones_like(a)), # Same as F0
    'SC-FD': ('Shrinking Core (Film Diffusion)', lambda a: np.ones_like(a)), # Alias for D11-Film

    # --- Chemical Reaction Models (Cn) ---
    'C1': ('Chemical Reaction (C1)', lambda a: (3/2) * (1-a)**(1/3)),
    'C2': ('Chemical Reaction (C2)', lambda a: 4 * (1-a)**(3/4)),
    'C4': ('Chemical Reaction (C4)', lambda a: (1/2) * (1-a)**(-0.5)),
    'C5': ('Chemical Reaction (C5)', lambda a: (1/3) * (1-a)**(-2/3)),
    'C6': ('Chemical Reaction (C6)', lambda a: (1/4) * (1-a)**(-3/4)),
    'C7': ('Chemical Reaction (C7)', lambda a: (1/2) * (1-a)**(-1)),
    'C8': ('Chemical Reaction (C8)', lambda a: (1/3) * (1-a)**(-1)),
    'C9': ('Chemical Reaction (C9)', lambda a: (1/5) * (1-a)**(-4/5)),
    'C10': ('Chemical Reaction (C10)', lambda a: (1/6) * (1-a)**(-5/6)),
    
    # --- Autocatalytic Models (For Shape Comparison) ---
    'PT(1,1)': ('Prout-Tompkins (m=1, n=1)', lambda a: a * (1-a)),
    'AC(0.5,1)': ('Autocatalytic (m=0.5, n=1)', lambda a: (a**0.5) * (1-a)),
    'AC(1.5,1)': ('Autocatalytic (m=1.5, n=1)', lambda a: (a**1.5) * (1-a)),
    'AC(2,1)': ('Autocatalytic (m=2, n=1)', lambda a: (a**2) * (1-a)),
    'AC(1,0.5)': ('Autocatalytic (m=1, n=0.5)', lambda a: a * (1-a)**0.5),
    'AC(1,1.5)': ('Autocatalytic (m=1, n=1.5)', lambda a: a * (1-a)**1.5),
    'AC(1,2)': ('Autocatalytic (m=1, n=2)', lambda a: a * (1-a)**2),
    'AC(0.5,2)': ('Autocatalytic (m=0.5, n=2)', lambda a: (a**0.5) * (1-a)**2),
    'AC(2,0.5)': ('Autocatalytic (m=2, n=0.5)', lambda a: (a**2) * (1-a)**0.5),
    'AC(2,2)': ('Autocatalytic (m=2, n=2)', lambda a: (a**2) * (1-a)**2),
    'SB(0.5,0.5,1)': ('Sestak-Berggren (0.5,0.5,1)', lambda a: (a**0.5) * (1-a)**0.5 * (-np.log(1-a))),
    'SB(0.5,1,1)': ('Sestak-Berggren (0.5,1,1)', lambda a: (a**0.5) * (1-a) * (-np.log(1-a))),
    'SB(1,0.5,1)': ('Sestak-Berggren (1,0.5,1)', lambda a: a * (1-a)**0.5 * (-np.log(1-a))),
    'SB(1,1,0.5)': ('Sestak-Berggren (1,1,0.5)', lambda a: a * (1-a) * (-np.log(1-a))**0.5),
    'SB(1,1,1)': ('Sestak-Berggren (1,1,1)', lambda a: a * (1-a) * (-np.log(1-a))),
    'SB(1,1,2)': ('Sestak-Berggren (1,1,2)', lambda a: a * (1-a) * (-np.log(1-a))**2),
}
# --- END: Mechanism Plot Integration ---


# ───────── MODERNIZED DIALOG WINDOWS ─────────
# --- START: Mechanism Plot Integration ---
# Step 2: Design the "Model Selection" Dialog
class ModelSelectionDialog(tk.Toplevel):
    def __init__(self, parent, model_library):
        super().__init__(parent)
        self.title("Select Ideal Models for Comparison")
        self.result = None

        try:
            # Assuming parent has a 'resource_path' method
            icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico")
            self.iconbitmap(icon_path)
        except Exception:
            pass 

        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        self.configure(bg="#f8f9fa")

        main_frame = ttk.Frame(self, padding="20")
        main_frame.pack(fill="both", expand=True)
        
        ttk.Label(main_frame, text="Select models to compare (use Ctrl or Shift to select multiple):", 
                  font=("Segoe UI", 10)).pack(pady=(0, 10), anchor="w")

        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill="both", expand=True)

        self.listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED, height=15, 
                                 bg="#ffffff", fg="#212529", font=("Segoe UI", 10),
                                 selectbackground="#6366f1", selectforeground="white")
        self.listbox.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox.config(yscrollcommand=scrollbar.set)

        self.model_keys = list(model_library.keys())
        for key in self.model_keys:
            full_name = model_library[key][0]
            self.listbox.insert(tk.END, f"{key} ({full_name})")

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill="x", pady=(15, 0))
        
        ok_btn = ttk.Button(btn_frame, text="OK", command=self.on_ok, style="Primary.TButton")
        ok_btn.pack(side="right")
        cancel_btn = ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton")
        cancel_btn.pack(side="right", padx=(0, 10))

    def on_ok(self):
        selected_indices = self.listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("No Selection", "Please select at least one model.", parent=self)
            return
        self.result = [self.model_keys[i] for i in selected_indices]
        self.destroy()
# --- END: Mechanism Plot Integration ---


class SaveOptionsDialog(tk.Toplevel):
    def __init__(self, parent, fig, is_thermo_plot=False):
        super().__init__(parent)
        self.title("Save Plot Options")
        
        try:
            icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico")
            self.iconbitmap(icon_path)
        except Exception:
            pass 

        self.fig = fig
        self.is_thermo_plot = is_thermo_plot
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        self.configure(bg="#f8f9fa")

        self.presets = {
            "Angew. Chem. (Single Col)": (8.5, 4/3),
            "Angew. Chem. (Double Col)": (17.5, 16/9),
            "ACS Journals (Single Col)": (8.46, 4/3),
            "ACS Journals (Double Col)": (17.78, 16/9),
            "Custom": (None, None)
        }

        main_frame = ttk.Frame(self, padding="25")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        title_label = ttk.Label(main_frame, text="Export Plot Settings",
                               font=("Segoe UI", 14, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20), sticky="w")

        preset_frame = ttk.LabelFrame(main_frame, text="Size Settings", padding="15", style="Modern.TLabelframe")
        preset_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 15))

        ttk.Label(preset_frame, text="Size Preset:", font=("Segoe UI", 10)).grid(row=0, column=0, sticky="w", pady=5)
        self.preset_var = tk.StringVar(value=list(self.presets.keys())[0])
        self.preset_menu = ttk.Combobox(preset_frame, textvariable=self.preset_var,
                                       values=list(self.presets.keys()), state="readonly",
                                       font=("Segoe UI", 10), width=25)
        self.preset_menu.grid(row=0, column=1, columnspan=2, sticky="ew", pady=5)
        self.preset_menu.bind("<<ComboboxSelected>>", self.update_size_entries)

        ttk.Label(preset_frame, text="Width (cm):", font=("Segoe UI", 10)).grid(row=1, column=0, sticky="w", pady=5)
        self.width_var = tk.StringVar()
        self.width_entry = ttk.Entry(preset_frame, textvariable=self.width_var,
                                   state="disabled", font=("Segoe UI", 10), style="Modern.TEntry")
        self.width_entry.grid(row=1, column=1, sticky="ew", pady=5, padx=(10, 0))

        ttk.Label(preset_frame, text="Height (cm):", font=("Segoe UI", 10)).grid(row=2, column=0, sticky="w", pady=5)
        self.height_var = tk.StringVar()
        self.height_entry = ttk.Entry(preset_frame, textvariable=self.height_var,
                                    state="disabled", font=("Segoe UI", 10), style="Modern.TEntry")
        self.height_entry.grid(row=2, column=1, sticky="ew", pady=5, padx=(10, 0))

        quality_frame = ttk.LabelFrame(main_frame, text="Quality & Format", padding="15", style="Modern.TLabelframe")
        quality_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(0, 15))

        ttk.Label(quality_frame, text="DPI:", font=("Segoe UI", 10)).grid(row=0, column=0, sticky="w", pady=5)
        self.dpi_var = tk.StringVar(value="600")
        self.dpi_entry = ttk.Entry(quality_frame, textvariable=self.dpi_var,
                                 font=("Segoe UI", 10), style="Modern.TEntry")
        self.dpi_entry.grid(row=0, column=1, sticky="ew", pady=5, padx=(10, 0))

        self.formats = [("TIFF", "*.tiff"), ("PNG", "*.png"), ("PDF", "*.pdf"), ("SVG", "*.svg")]
        ttk.Label(quality_frame, text="Format:", font=("Segoe UI", 10)).grid(row=1, column=0, sticky="w", pady=5)
        self.format_var = tk.StringVar(value=self.formats[0][0])
        self.format_menu = ttk.Combobox(quality_frame, textvariable=self.format_var,
                                      values=[f[0] for f in self.formats], state="readonly",
                                      font=("Segoe UI", 10))
        self.format_menu.grid(row=1, column=1, sticky="ew", pady=5, padx=(10, 0))
        self.format_menu.bind("<<ComboboxSelected>>", self.update_transparency_option)

        options_frame = ttk.LabelFrame(main_frame, text="Advanced Options", padding="15", style="Modern.TLabelframe")
        options_frame.grid(row=3, column=0, columnspan=2, sticky="ew", pady=(0, 20))

        self.transparent_var = tk.BooleanVar(value=False)
        self.transparent_check = ttk.Checkbutton(options_frame, text="Transparent Background",
                                                variable=self.transparent_var, style="Modern.TCheckbutton")
        self.transparent_check.grid(row=0, column=0, sticky="w", pady=5)

        self.fill_var = tk.BooleanVar(value=False)
        self.fill_check = ttk.Checkbutton(options_frame, text="Fill Curve Area",
                                        variable=self.fill_var, style="Modern.TCheckbutton")
        if self.is_thermo_plot:
            self.fill_check.grid(row=1, column=0, sticky="w", pady=5)

        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=(10, 0), sticky="e")

        save_btn = ttk.Button(btn_frame, text="Save Plot", command=self.save, style="Primary.TButton")
        save_btn.pack(side="right")
        cancel_btn = ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton")
        cancel_btn.pack(side="right", padx=(0, 10))

        preset_frame.grid_columnconfigure(1, weight=1)
        quality_frame.grid_columnconfigure(1, weight=1)

        self.update_size_entries()
        self.update_transparency_option()

        # --- Automatic Sizing Logic ---
        self.update_idletasks()
        w = self.winfo_reqwidth()
        h = self.winfo_reqheight()
        self.geometry(f"{w+20}x{h+20}")
        self.resizable(False, False)

    def update_size_entries(self, event=None):
        preset = self.preset_var.get()
        if preset == "Custom":
            self.width_entry.config(state="normal")
            self.height_entry.config(state="normal")
            self.width_var.set("")
            self.height_var.set("")
        else:
            w_cm, aspect = self.presets[preset]
            h_cm = w_cm / aspect
            self.width_var.set(f"{w_cm:.2f}")
            self.height_var.set(f"{h_cm:.2f}")
            self.width_entry.config(state="disabled")
            self.height_entry.config(state="disabled")

    def update_transparency_option(self, event=None):
        fmt = self.format_var.get()
        if fmt in ["TIFF", "PNG",]:
            self.transparent_check.config(state="normal")
        else:
            self.transparent_check.config(state="disabled")
            self.transparent_var.set(False)

    def save(self):
        try:
            w_cm = float(self.width_var.get())
            h_cm = float(self.height_var.get())
            dpi = int(self.dpi_var.get())
            fmt = self.format_var.get()
            transparent = self.transparent_var.get()
            fill = self.fill_var.get()
        except ValueError:
            return messagebox.showerror("Invalid Input",
                                      "Please enter valid numbers for size and DPI.",
                                      parent=self)

        file_info = next(f for f in self.formats if f[0] == fmt)
        p = filedialog.asksaveasfilename(
            parent=self,
            defaultextension=file_info[1],
            initialfile=f"figure.{file_info[1].split('.')[-1]}",
            filetypes=self.formats
        )
        if not p: return

        # --- MODIFICATION START: Legend Size Coordination ---
        original_legend_fontsizes = {}
        try:
            # Temporarily set a smaller font size for export that the user prefers
            for i, ax in enumerate(self.fig.axes):
                if ax.get_legend():
                    original_legend_fontsizes[i] = ax.get_legend().get_texts()[0].get_fontsize()
                    for text in ax.get_legend().get_texts():
                        text.set_fontsize(7) # Set to a smaller size for the exported file
            # --- MODIFICATION END: Legend Size Coordination ---

            orig_size = self.fig.get_size_inches()
            self.fig.set_size_inches(w_cm/2.54, h_cm/2.54)
            fills = []
            if self.is_thermo_plot and fill:
                fills = [ax.fill_between(l.get_xdata(), l.get_ydata(), 0,
                                       color=l.get_color(), alpha=0.1)
                        for ax in self.fig.axes for l in ax.get_lines()]

            self.fig.savefig(p, dpi=dpi,
                            transparent=transparent if fmt in ["TIFF", "PNG"] else False,
                            bbox_inches='tight', pad_inches=0.05)

            [f.remove() for f in fills]
            self.fig.set_size_inches(orig_size)
        
        finally:
            # --- MODIFICATION START: Legend Size Coordination ---
            # Restore the original, larger font size for the display window
            for i, ax in enumerate(self.fig.axes):
                if ax.get_legend() and i in original_legend_fontsizes:
                    for text in ax.get_legend().get_texts():
                        text.set_fontsize(original_legend_fontsizes[i])
            # --- MODIFICATION END: Legend Size Coordination ---
        
        self.destroy()

class CalculationSettingsDialog(tk.Toplevel):
    def __init__(self, parent, temp_range_suggestion=(35.0, 800.0)):
        super().__init__(parent)
        self.title("Kinetic Analysis Settings")

        try:
            icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico")
            self.iconbitmap(icon_path)
        except Exception:
            pass
            
        self.result = None
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        self.configure(bg="#f8f9fa")
        self.temp_range_suggestion = temp_range_suggestion

        main_frame = ttk.Frame(self, padding="30")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        title_label = ttk.Label(main_frame, text="Kinetic Analysis Configuration",
                               font=("Segoe UI", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 25), sticky="w")
        
        self.preprocess_enabled_var = tk.BooleanVar(value=False)
        preprocess_check = ttk.Checkbutton(main_frame, text="Enable Advanced Data Preprocessing & Truncation",
                                           variable=self.preprocess_enabled_var, style="Modern.TCheckbutton",
                                           command=self.toggle_preprocessing_options)
        preprocess_check.grid(row=1, column=0, columnspan=2, sticky='w', pady=(0, 5))

        temp_range_frame = ttk.LabelFrame(main_frame, text="Temperature Preprocessing Range", padding="20", style="Modern.TLabelframe")
        temp_range_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(0, 20))
        temp_range_frame.grid_columnconfigure(1, weight=1)

        ttk.Label(temp_range_frame, text="Start Temp (°C):",
                 font=("Segoe UI", 11)).grid(row=0, column=0, sticky="w", pady=8)
        self.temp_start_var = tk.StringVar()
        self.temp_start_entry = ttk.Entry(temp_range_frame, textvariable=self.temp_start_var,
                                          font=("Segoe UI", 11), style="Modern.TEntry")
        self.temp_start_entry.grid(row=0, column=1, sticky="ew", padx=(15, 0), pady=8)

        ttk.Label(temp_range_frame, text="End Temp (°C):",
                 font=("Segoe UI", 11)).grid(row=1, column=0, sticky="w", pady=8)
        self.temp_end_var = tk.StringVar()
        self.temp_end_entry = ttk.Entry(temp_range_frame, textvariable=self.temp_end_var,
                                        font=("Segoe UI", 11), style="Modern.TEntry")
        self.temp_end_entry.grid(row=1, column=1, sticky="ew", padx=(15, 0), pady=8)

        method_frame = ttk.LabelFrame(main_frame, text="Analysis Method", padding="20", style="Modern.TLabelframe")
        method_frame.grid(row=3, column=0, columnspan=2, sticky="ew", pady=(0, 20))

        ttk.Label(method_frame, text="Kinetic Method:",
                 font=("Segoe UI", 11, "bold")).grid(row=0, column=0, sticky="w", pady=5)
        self.method_var = tk.StringVar(value="Vyazovkin")
        methods = ["Vyazovkin", "Friedman", "KAS", "OFW"]
        method_combo = ttk.Combobox(method_frame, textvariable=self.method_var,
                                   values=methods, state="readonly",
                                   font=("Segoe UI", 11), width=20)
        method_combo.grid(row=0, column=1, sticky="ew", padx=(15, 0), pady=5)

        descriptions = {
            "Vyazovkin": "Advanced non-linear isoconversional method (Recommended)",
            "Friedman": "Direct differential method with good accuracy",
            "KAS": "Kissinger-Akahira-Sunose integral method",
            "OFW": "Ozawa-Flynn-Wall integral method"
        }
        self.desc_var = tk.StringVar(value=descriptions["Vyazovkin"])
        desc_label = ttk.Label(method_frame, textvariable=self.desc_var,
                              font=("Segoe UI", 9), foreground="#666666", wraplength=350)
        desc_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(5, 0))

        def update_description(*args):
            self.desc_var.set(descriptions.get(self.method_var.get(), ""))
        self.method_var.trace('w', update_description)

        range_frame = ttk.LabelFrame(main_frame, text="Conversion Range Settings", padding="20", style="Modern.TLabelframe")
        range_frame.grid(row=4, column=0, columnspan=2, sticky="ew", pady=(0, 25))
        range_frame.grid_columnconfigure(1, weight=1)

        ttk.Label(range_frame, text="α Range Start:",
                 font=("Segoe UI", 11)).grid(row=0, column=0, sticky="w", pady=8)
        self.alpha_min_var = tk.StringVar(value="0.01")
        alpha_min_entry = ttk.Entry(range_frame, textvariable=self.alpha_min_var,
                                   font=("Segoe UI", 11), style="Modern.TEntry")
        alpha_min_entry.grid(row=0, column=1, sticky="ew", padx=(15, 0), pady=8)

        ttk.Label(range_frame, text="α Range End:",
                 font=("Segoe UI", 11)).grid(row=1, column=0, sticky="w", pady=8)
        self.alpha_max_var = tk.StringVar(value="0.99")
        alpha_max_entry = ttk.Entry(range_frame, textvariable=self.alpha_max_var,
                                   font=("Segoe UI", 11), style="Modern.TEntry")
        alpha_max_entry.grid(row=1, column=1, sticky="ew", padx=(15, 0), pady=8)

        ttk.Label(range_frame, text="α Step Size:",
                 font=("Segoe UI", 11)).grid(row=2, column=0, sticky="w", pady=8)
        self.alpha_step_var = tk.StringVar(value="0.01")
        alpha_step_entry = ttk.Entry(range_frame, textvariable=self.alpha_step_var,
                                    font=("Segoe UI", 11), style="Modern.TEntry")
        alpha_step_entry.grid(row=2, column=1, sticky="ew", padx=(15, 0), pady=8)

        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=5, column=0, columnspan=2, sticky="e", pady=(20,0))
        calc_btn = ttk.Button(btn_frame, text="Start Calculation", command=self.on_ok, style="Primary.TButton")
        calc_btn.pack(side="right")
        cancel_btn = ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton")
        cancel_btn.pack(side="right", padx=(0, 15))

        self.toggle_preprocessing_options() # Set initial state
        self.update_idletasks()
        w = self.winfo_reqwidth()
        h = self.winfo_reqheight()
        self.geometry(f"{w+20}x{h+20}")
        self.resizable(False, False)

    def toggle_preprocessing_options(self):
        state = "normal" if self.preprocess_enabled_var.get() else "disabled"
        self.temp_start_entry.config(state=state)
        self.temp_end_entry.config(state=state)
        if state == "disabled":
            self.temp_start_var.set("")
            self.temp_end_var.set("")
        else:
            self.temp_start_var.set("25.0")
            self.temp_end_var.set("350.0")

    def on_ok(self):
        try:
            a_min = float(self.alpha_min_var.get())
            a_max = float(self.alpha_max_var.get())
            a_step = float(self.alpha_step_var.get())
            if not (0 <= a_min < a_max <= 1 and a_step > 0):
                raise ValueError("Invalid alpha range.")
            
            self.result = {
                "method": self.method_var.get(),
                "alpha_range": (a_min, a_max, a_step),
                "preprocess_enabled": self.preprocess_enabled_var.get()
            }
            
            if self.result["preprocess_enabled"]:
                t_start = float(self.temp_start_var.get())
                t_end = float(self.temp_end_var.get())
                if t_start >= t_end:
                    raise ValueError("Start temperature must be less than end temperature.")
                self.result["temp_range_c"] = (t_start, t_end)

            self.destroy()
        except ValueError as e:
            messagebox.showerror("Invalid Input",
                               str(e) or "Please enter valid numeric values for ranges.",
                               parent=self)

# --- MODIFICATION START: NEW DIALOG FOR CUMULATIVE PREDICTION ---
class ConversionTimeSettingsDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Conversion Time Prediction")
        
        try:
            icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico")
            self.iconbitmap(icon_path)
        except Exception:
            pass

        self.result = None
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        self.configure(bg="#f8f9fa")

        main_frame = ttk.Frame(self, padding="25")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        title_label = ttk.Label(main_frame, text="Prediction Settings",
                               font=("Segoe UI", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20), sticky="w")

        temp_frame = ttk.LabelFrame(main_frame, text="Isothermal Condition", padding="20", style="Modern.TLabelframe")
        temp_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 20))
        temp_frame.grid_columnconfigure(1, weight=1)

        ttk.Label(temp_frame, text="Temperature (K):",
                 font=("Segoe UI", 11, "bold")).grid(row=0, column=0, sticky="w", pady=10)
        self.t0 = tk.StringVar(value="423")
        t0_entry = ttk.Entry(temp_frame, textvariable=self.t0, font=("Segoe UI", 12), style="Modern.TEntry")
        t0_entry.grid(row=0, column=1, sticky="ew", padx=(15, 0), pady=10)
        
        info_label = ttk.Label(temp_frame, 
                               text="This will predict the cumulative time (t-α curve)\nbased on the currently fitted kinetic model.",
                               font=("Segoe UI", 9), foreground="#666666")
        info_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(10, 0))

        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=2, column=0, columnspan=2, sticky="e", pady=(10,0))

        predict_btn = ttk.Button(btn_frame, text="Start Prediction", command=self._ok, style="Primary.TButton")
        predict_btn.pack(side="right")
        cancel_btn = ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton")
        cancel_btn.pack(side="right", padx=(0, 15))

        self.update_idletasks()
        w = self.winfo_reqwidth()
        h = self.winfo_reqheight()
        self.geometry(f"{w+20}x{h+20}")
        self.resizable(False, False)

    def _ok(self):
        try:
            T0 = float(self.t0.get())
            if T0 <= 0: raise ValueError("Temperature (K) must be a positive number.")
            self.result = {"T0": T0}
            self.destroy()
        except ValueError as e:
            messagebox.showerror("Invalid Input", str(e) or "Please check your input values.", parent=self)
# --- MODIFICATION END ---


# #################################################################################
# #################### START: MODIFIED DATA PREPROCESSING SECTION #################
# #################################################################################

# --- 1. NEW: Stricter Column Aliases ---
# Removed single-character aliases ('s', 'm') and added more specific ones.
T_ALIASES = ['temp', 'temperature', '°c', 'degc', 't-probe']
TIME_ALIASES = ['time', '(min)', 'min', '(s)', 'sec', 'seconds']
TG_ALIASES = ['mass', 'weight', 'wt', 'rem wt', 'tg', '(mg)', 'mg', '(%)']
DSC_ALIASES = ['dsc', 'heat flow', 'hf', 'dta', '(mw)', 'mw', '(uw)', 'uw']

# --- 2. NEW: Advanced Column Identification Function ---
def _find_best_col_idx(headers, aliases, priority_aliases=None):
    """
    Finds the best column index for a set of aliases using a scoring system.
    - Higher scores for priority aliases and whole-word matches.
    """
    if priority_aliases is None:
        priority_aliases = []
        
    best_score = -1
    best_idx = None

    for i, h in enumerate(headers):
        if not isinstance(h, str) or not h:
            continue
        h_lower = h.lower()
        
        for alias in aliases:
            score = 0
            # Use regex for whole-word matching to avoid partial matches like 's' in 'dsc'
            
            if re.search(r'\b' + re.escape(alias) + r'\b', h_lower):
                score = 2  # Base score for a whole-word match
                if alias in priority_aliases:
                    score = 3  # Highest score for a priority whole-word match
            elif alias in h_lower:
                score = 1  # Lower score for a substring match (fallback)

            if score > best_score:
                best_score = score
                best_idx = i
                
    return best_idx

def preprocess_file(path, temp_range_c):
    """
    MODIFIED: This function now uses a robust, prioritized column identification
    strategy, warns for ambiguities, ensures time is zeroed correctly, and validates
    the calculated heating rate.
    """
    # --- File Reading Logic (Unchanged) ---
    file_path_str = str(path).lower()
    if file_path_str.endswith((".xlsx", ".xls")):
        df_raw = pd.read_excel(path, header=None)
    else:
        df_raw = pd.read_csv(path, sep=None, engine='python', header=None, on_bad_lines='skip', encoding_errors='ignore')

    # --- Header and Data Location (Unchanged) ---
    header_row_index = -1
    for i, row in df_raw.iterrows():
        row_str = ' '.join(row.dropna().astype(str)).lower()
        # A simple check for 'temp' is sufficient to identify potential header rows
        if any(alias in row_str for alias in T_ALIASES):
            header_row_index = i
            break
            
    if header_row_index == -1: raise ValueError(f"Could not locate a valid header row in {Path(path).name}.")

    header_line = df_raw.iloc[header_row_index]
    possible_headers = header_line.fillna('').astype(str)
    
    data_start_row = -1
    for i in range(header_row_index + 1, len(df_raw)):
        row = df_raw.iloc[i]
        non_empty_cells = row.dropna()
        if non_empty_cells.empty: continue
        numeric_cells = pd.to_numeric(non_empty_cells, errors='coerce')
        if numeric_cells.notna().sum() >= 2:
            data_start_row = i
            break
    if data_start_row == -1: raise ValueError(f"Could not locate valid data in {Path(path).name}.")

    # --- 3. NEW: Robust Column Identification & Ambiguity Check ---
    temp_idx = _find_best_col_idx(possible_headers, T_ALIASES, ['temp', 'temperature'])
    if temp_idx is None: raise ValueError("Essential 'Temperature' column not found.")
    
    time_idx = _find_best_col_idx(possible_headers, TIME_ALIASES, ['time'])
    tg_idx = _find_best_col_idx(possible_headers, TG_ALIASES, ['mass', 'weight', 'tg'])
    dsc_idx = _find_best_col_idx(possible_headers, DSC_ALIASES, ['dsc', 'heat flow'])

    # NEW: Check for conflicts where multiple roles are assigned to the same column
    found_indices = {'Time': time_idx, 'Mass': tg_idx, 'DSC': dsc_idx}
    non_none_indices = {k: v for k, v in found_indices.items() if v is not None}
    
    from collections import Counter
    idx_counts = Counter(non_none_indices.values())
    
    for idx, count in idx_counts.items():
        if count > 1:
            conflicting_roles = [role for role, i in non_none_indices.items() if i == idx]
            header_name = possible_headers.iloc[idx]
            warnings.warn(
                f"\n[Warning] Ambiguous Column in '{Path(path).name}':\n"
                f"  > Header '{header_name}' was matched for multiple roles: {', '.join(conflicting_roles)}.\n"
                f"  > This may cause incorrect calculations. Please rename columns for clarity (e.g., 'Time (min)', 'Mass (mg)').\n"
            )

    # --- 4. Extract and Truncate Raw Data ---
    data_df = df_raw.iloc[data_start_row:].copy()
    df_intermediate = pd.DataFrame()

    temp_series = pd.to_numeric(data_df.iloc[:, temp_idx], errors='coerce')
    is_kelvin = 'k' in possible_headers.iloc[temp_idx].lower()
    df_intermediate['Temp_K'] = temp_series if is_kelvin else temp_series + 273.15
    
    mass_present = tg_idx is not None
    if mass_present:
        df_intermediate['Mass_mg'] = pd.to_numeric(data_df.iloc[:, tg_idx], errors='coerce')
    else:
        df_intermediate['Mass_mg'] = 1.0

    dsc_present = dsc_idx is not None
    if dsc_present:
        df_intermediate['DSC_raw'] = pd.to_numeric(data_df.iloc[:, dsc_idx], errors='coerce')
        dsc_header = possible_headers.iloc[dsc_idx].lower()
    else:
        df_intermediate['DSC_raw'] = 0.0
        dsc_header = ""

    if time_idx is not None:
        time_series = pd.to_numeric(data_df.iloc[:, time_idx], errors='coerce')
        is_seconds = any(s in possible_headers.iloc[time_idx].lower() for s in ['sec', '(s)'])
        df_intermediate['Time_s'] = time_series if is_seconds else time_series * 60
    else:
        df_intermediate['Time_s'] = np.nan

    df_intermediate.dropna(subset=['Temp_K'], inplace=True)
    start_k, end_k = temp_range_c[0] + 273.15, temp_range_c[1] + 273.15
    df_truncated = df_intermediate[(df_intermediate['Temp_K'] >= start_k) & (df_intermediate['Temp_K'] <= end_k)].copy()
    if df_truncated.empty: raise ValueError(f"No data found in the specified temperature range {temp_range_c}°C.")

    df_truncated.sort_values(by='Temp_K', ascending=True, inplace=True)
    df_truncated.reset_index(drop=True, inplace=True)

    # --- 5. Standardize Data & Calculate Derivatives ---
    initial_mass = df_truncated['Mass_mg'].iloc[0] if mass_present else 1.0
    if mass_present and initial_mass > 0:
        df_truncated['TG_pct'] = (df_truncated['Mass_mg'] / initial_mass) * 100
    else:
        df_truncated['TG_pct'] = 100.0

    dsc_signal = df_truncated['DSC_raw'].copy()
    if 'uw' in dsc_header or 'μw' in dsc_header:
        dsc_signal /= 1000.0
    if not any(s in dsc_header for s in ['/mg', '/g']) and mass_present and initial_mass > 0:
        dsc_signal /= initial_mass
    df_truncated['DSC'] = dsc_signal - dsc_signal.iloc[0]

    # --- 6. NEW: Robust Heating Rate & Time Calculation ---
    beta_K_per_s = 0
    time_col_is_valid = 'Time_s' in df_truncated.columns and df_truncated['Time_s'].notna().all()

    if time_col_is_valid:
        corr_matrix = np.corrcoef(df_truncated['Time_s'], df_truncated['Temp_K'])
        if abs(corr_matrix[0, 1]) > 0.95:  # Check for strong correlation
            fit = np.polyfit(df_truncated['Time_s'], df_truncated['Temp_K'], 1)
            if fit[0] > 0:  # Check for positive heating rate
                beta_K_per_s = fit[0]

    # Fallback to filename if time column is invalid or yields a bad fit
    if beta_K_per_s <= 0:
        beta_match = re.search(r"(\d+(?:\.\d+)?)\s*(k|c)", Path(path).stem, re.IGNORECASE)
        if not beta_match: raise ValueError(f"Time column is invalid/missing, and heating rate could not be parsed from filename: {Path(path).name}.")
        beta_K_per_min = float(beta_match.group(1))
        beta_K_per_s = beta_K_per_min / 60.0
        # Recalculate time axis based on filename heating rate
        if beta_K_per_s > 0:
            df_truncated['Time_s'] = (df_truncated['Temp_K'] - df_truncated['Temp_K'].iloc[0]) / beta_K_per_s
        else:
            df_truncated['Time_s'] = 0.0

    # --- 7. NEW: Final Formatting with Time Zero-Point Reset ---
    time_in_minutes = df_truncated['Time_s'] / 60
    
    final_df = pd.DataFrame({
        "Time_min": time_in_minutes - time_in_minutes.iloc[0], # Ensures time starts at 0
        "Temp_C": df_truncated['Temp_K'] - 273.15,
        "TG_pct": df_truncated['TG_pct'],
        "DSC": df_truncated['DSC']
    })
    
    final_df.attrs['kinetic_analysis_possible'] = mass_present
    
    return beta_K_per_s * 60, final_df

# ###############################################################################
# #################### END: MODIFIED DATA PREPROCESSING SECTION #################
# ###############################################################################


def build_friedman_tables(dfs, alphas):
    betas = sorted(dfs.keys())
    T_at_alpha = {a: {β: np.interp(a, dfs[β]["alpha"], dfs[β]["Temp_K"]) for β in betas} for a in alphas}
    df_aT = pd.DataFrame([{ "alpha": a, **{f"T_{β:.2f}K/min": T_at_alpha[a][β] for β in betas} } for a in alphas])

    rows_xy = []
    for a in alphas:
        for β in betas:
            T_val = T_at_alpha[a][β]
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            dAdT_val = np.interp(T_val, dfs[β]["Temp_K"], dfs[β]["dAdT"])
            if dAdT_val <= 0 or not np.isfinite(dAdT_val): continue
            # Y value is ln(d(alpha)/dt) which is ln(beta * d(alpha)/dT)
            rows_xy.append({"alpha": a, "beta": β, "X": 1.0/T_val, "Y": np.log(β * dAdT_val)})
            ### MODIFICATION END ###
    df_xy = pd.DataFrame(rows_xy)

    ea_rows = []
    for a in alphas:
        s = df_xy[df_xy["alpha"] == a]
        if len(s) < 2: continue
        x, y = s["X"], s["Y"]
        slope, _ = np.polyfit(x, y, 1)
        residuals = y - (np.polyval(np.polyfit(x,y,1), x))
        SSE = (residuals**2).sum(); SST = ((y - y.mean())**2).sum()
        R2 = 1 - SSE/SST if SST > 0 else 0
        se_slope = np.sqrt(SSE / (len(x)-2)) / np.sqrt(((x-x.mean())**2).sum() + EPS) if len(x)>2 else 0
        ea_rows.append({"alpha": a, "Ea_kJ_per_mol": -slope*R/1000, "R2": R2, "StdErr_kJ": se_slope*R/1000})
    df_ea = pd.DataFrame(ea_rows)
    return df_ea, df_aT, df_xy

def build_kas_tables(dfs, alphas):
    betas = sorted(dfs.keys())
    T_at_alpha = {a: {β: np.interp(a, dfs[β]["alpha"], dfs[β]["Temp_K"]) for β in betas} for a in alphas}
    df_aT = pd.DataFrame([{ "alpha": a, **{f"T_{β:.2f}K/min": T_at_alpha[a][β] for β in betas} } for a in alphas])

    rows_xy = []
    for a in alphas:
        for β in betas:
            T_val = T_at_alpha[a][β]
            rows_xy.append({"alpha": a, "beta": β, "X": 1.0/T_val, "Y": np.log(β / T_val**2)})
    df_xy = pd.DataFrame(rows_xy)

    ea_rows = []
    for a in alphas:
        s = df_xy[df_xy["alpha"] == a]
        if len(s) < 2: continue
        x, y = s["X"], s["Y"]
        slope, _ = np.polyfit(x, y, 1)
        residuals = y - (np.polyval(np.polyfit(x,y,1), x))
        SSE = (residuals**2).sum(); SST = ((y - y.mean())**2).sum()
        R2 = 1 - SSE/SST if SST > 0 else 0
        se_slope = np.sqrt(SSE / (len(x)-2)) / np.sqrt(((x-x.mean())**2).sum() + EPS) if len(x)>2 else 0
        ea_rows.append({"alpha": a, "Ea_kJ_per_mol": -slope*R/1000, "R2": R2, "StdErr_kJ": se_slope*R/1000})
    df_ea = pd.DataFrame(ea_rows)
    return df_ea, df_aT, df_xy

def build_ofw_tables(dfs, alphas):
    betas = sorted(dfs.keys())
    T_at_alpha = {a: {β: np.interp(a, dfs[β]["alpha"], dfs[β]["Temp_K"]) for β in betas} for a in alphas}
    df_aT = pd.DataFrame([{ "alpha": a, **{f"T_{β:.2f}K/min": T_at_alpha[a][β] for β in betas} } for a in alphas])

    rows_xy = []
    for a in alphas:
        for β in betas:
            T_val = T_at_alpha[a][β]
            rows_xy.append({"alpha": a, "beta": β, "X": 1.0/T_val, "Y": np.log(β)})
    df_xy = pd.DataFrame(rows_xy)

    ea_rows = []
    for a in alphas:
        s = df_xy[df_xy["alpha"] == a]
        if len(s) < 2: continue
        x, y = s["X"], s["Y"]
        slope, _ = np.polyfit(x, y, 1)
        residuals = y - (np.polyval(np.polyfit(x,y,1), x))
        SSE = (residuals**2).sum(); SST = ((y - y.mean())**2).sum()
        R2 = 1 - SSE/SST if SST > 0 else 0
        se_slope = np.sqrt(SSE / (len(x)-2)) / np.sqrt(((x-x.mean())**2).sum() + EPS) if len(x)>2 else 0
        ea_rows.append({"alpha": a, "Ea_kJ_per_mol": -slope*R/1000 * 0.921, "R2": R2, "StdErr_kJ": se_slope*R/1000 * 0.921})
    df_ea = pd.DataFrame(ea_rows)
    return df_ea, df_aT, df_xy

def interp_Tα(dfs, alphas):
    out = {a: {} for a in alphas}
    for β, df in dfs.items():
        for a in alphas:
            idx = (df["alpha"] >= a).idxmax()
            
            if df.loc[idx, "alpha"] == a or idx == 0:
                out[a][β] = df.loc[idx, "Temp_K"]
            else:
                i0, i1 = idx - 1, idx
                a0, a1_val = df.loc[i0, "alpha"], df.loc[i1, "alpha"]
                T0, T1 = df.loc[i0, "Temp_K"], df.loc[i1, "Temp_K"]
                
                denominator = a1_val - a0
                if abs(denominator) < EPS:
                    out[a][β] = T0
                else:
                    out[a][β] = T0 + (a - a0) * (T1 - T0) / denominator
    return out

def temp_integral(Ea, T_alpha, T_prev_alpha):
    """Calculates the temperature integral between T_prev_alpha and T_alpha."""
    integrand = lambda T: np.exp(-Ea / (R * T))
    result, _ = quad(integrand, T_prev_alpha, T_alpha)
    return result

def objective_function(Ea, T_alpha_data, T_prev_alpha_data):
    """Objective function for the advanced Vyazovkin method."""
    betas = list(T_alpha_data.keys())
    n = len(betas)
    phi = 0.0
    for i in range(n):
        for j in range(n):
            if i == j: continue
            β_i, β_j = betas[i], betas[j]
            T_i, T_j = T_alpha_data[β_i], T_alpha_data[β_j]
            T_prev_i, T_prev_j = T_prev_alpha_data[β_i], T_prev_alpha_data[β_j]
            
            if T_i <= T_prev_i or T_j <= T_prev_j: continue

            I_i = temp_integral(Ea, T_i, T_prev_i)
            I_j = temp_integral(Ea, T_j, T_prev_j)
            
            if abs(I_j) < 1e-100 or abs(β_i) < 1e-100: continue
            
            phi += abs((I_i * β_j) / (I_j * β_i) - 1.0)
    return phi

def calculate_vyazovkin_ea(Tα_dict, alphas, Ea_bounds=(1_000, 500_000)):
    """Calculates Ea for each alpha using the advanced Vyazovkin method."""
    results = []
    for idx, α in enumerate(alphas):
        if idx == 0:
            continue

        prev_α = alphas[idx - 1]
        T_alpha_data = Tα_dict[α]
        T_prev_alpha_data = Tα_dict[prev_α]

        res = minimize_scalar(
            objective_function,
            args=(T_alpha_data, T_prev_alpha_data),
            bounds=Ea_bounds,
            method='bounded'
        )
        results.append({
            "alpha": α,
            "Ea_kJ_per_mol": res.x / 1000.0
        })
    df = pd.DataFrame(results)
    df['R2'] = np.nan
    df['StdErr_kJ'] = np.nan
    return df

def build_vyazovkin_tables(dfs, alphas):
    """Builds all tables required for the Vyazovkin analysis."""
    Tα = interp_Tα(dfs, alphas)
    df_ea = calculate_vyazovkin_ea(Tα, alphas)
    
    valid_alphas = df_ea['alpha'].unique()
    
    betas = sorted(next(iter(Tα.values())).keys())
    df_aT = pd.DataFrame([{ "alpha": a, **{f"T_{β:.2f}K/min": Tα[a][β] for β in betas} } for a in valid_alphas])
    
    df_xy = pd.DataFrame() # No X-Y plot for Vyazovkin
    return df_ea, df_aT, df_xy

def _smooth_ea_series(ea_df, window: int = 7, poly: int = 3):
    if window % 2 == 0: window += 1
    ea_arr = ea_df["Ea_kJ_per_mol"].to_numpy(float)
    if ea_arr.size >= window:
        return savgol_filter(ea_arr, window_length=window, polyorder=poly, mode="interp")
    return ea_arr

def prep_arrays(dfs, ea_ser, alphas):
    ### MODIFICATION START: UNIFICATION OF GLOBAL FIT ###
    # This function now prepares data for d(alpha)/dt (time) domain fitting
    # 'y' is now d(alpha)/dt (in 1/min)
    # 'βs' (beta) is no longer needed in the data tuple for the residual function
    
    y_dAdt, a, a1, T, Ea = [], [], [], [], []
    for β, df in dfs.items():
        alpha = np.clip(df["alpha"].values, EPS, 1-EPS)
        
        ### MODIFICATION START: Renamed dadt -> dAdT ###
        mask = (alpha >= alphas.min()) & (alpha <= alphas.max()) & np.isfinite(df["dAdT"].values)
        alpha_masked = alpha[mask]
        
        # This is the key change: y is now d(alpha)/dt (1/min)
        y_dAdt.extend(df["dAdT"].values[mask] * β) 
        ### MODIFICATION END ###
        
        a.extend(alpha_masked)
        a1.extend(np.clip(1 - alpha_masked, EPS, None)); T.extend(df["Temp_K"].values[mask])
        Ea.extend(np.interp(alpha_masked, ea_ser.index, ea_ser.values))
        
    arrs = [np.asarray(v, float) for v in (y_dAdt, a, a1, T, np.asarray(Ea))]
    good = np.all(np.isfinite(arrs), axis=0)
    # Return tuple without beta
    return tuple(ar[good] for ar in arrs)
    ### MODIFICATION END ###

def resid(par, data):
    ### MODIFICATION START: UNIFICATION OF GLOBAL FIT ###
    # This residual function now fits in the d(alpha)/dt (time) domain
    # It now compares experimental d(alpha)/dt to model d(alpha)/dt
    
    m, n, p_, logA = par
    # Data tuple no longer contains beta
    y_dAdt, a, a1, T, Ea = data
    
    f = a**m * a1**n * (-np.log(a1))**p_
    
    # Model d(alpha)/dt = A * exp(-Ea/RT) * f(alpha)
    # A is A_per_min (since data was prepared with beta in K/min)
    A_per_min = math.exp(logA)
    y_calc = A_per_min * np.exp(-Ea / (R * T)) * f
    
    return y_calc - y_dAdt
    ### MODIFICATION END ###

def _df_to_latex_table(df, caption, label):
    lines = []
    col_format = "c" * len(df.columns)
    lines.append("\\begin{table}[h!]")
    lines.append("\\centering")
    lines.append(f"\\caption{{{caption}}}")
    lines.append(f"\\label{{{label}}}")
    lines.append(f"\\begin{{tabular}}{{{col_format}}}")
    lines.append("\\toprule")
    safe_columns = [str(c).replace('%', r'\%').replace('&', r'\&') for c in df.columns]
    lines.append(" & ".join(safe_columns) + r" \\ \midrule")
    for _, row in df.iterrows():
        def format_cell(x):
            if isinstance(x, (int, float)): return f"{x:.4g}"
            return str(x)
        lines.append(" & ".join(map(format_cell, row)) + r" \\")
    lines.append("\\bottomrule")
    lines.append("\\end{tabular}")
    lines.append("\\end{table}")
    return "\n".join(lines)

class App(tk.Tk):
    BG = "#f8f9fa"
    BG_SECONDARY = "#ffffff"
    FG = "#212529"
    FG_SECONDARY = "#6c757d"
    BRAND = "#6366f1"
    BRAND_ACCENT = "#8b5cf6"

    def resource_path(self, relative_path):
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def __init__(self):
        super().__init__()
        self.title("TGAX Kinetics v1.0 Advanced Thermal Analysis")
        self.configure(bg=self.BG)

        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            self.iconbitmap(icon_path)
        except tk.TclError:
            pass

        self.config_path = Path(__file__).parent / "tgax_config.json"
        self._load_and_set_dpi()

        try:
            self.state('zoomed')
        except tk.TclError:
            self.attributes('-zoomed', True)
        
        self.task_queue = Queue()
        self.is_task_running = False
        self.menu_states = {}

        ### MODIFICATION START: Added m0/m_inf attributes ###
        self.dfs, self.plots, self.fit_results, self.cka_results, self.preprocessed_files, self.sample_masses = {}, {}, None, None, {}, {}
        self.mass_loss_parameters = {} # Stores {beta: (m0, m_inf)}
        ### MODIFICATION END ###
        
        self.autocatalytic_results = None
        self.npa_results = None # <--- ### MODIFICATION START ###
        # --- MODIFICATION START: New attribute for storing the last fitted model ---
        
        self.last_fitted_model = None 
        # --- MODIFICATION END ---
        self.autocatalytic_model_definitions = {
            "SB": ("Generalized Sestak-Berggren (SB)", "dα/dt = A·exp(-E/RT)·αᵐ(1-α)ⁿ[-ln(1-α)]ᵖ"),
            "KS": ("Kamal–Sourour (KS)", "dα/dt = [A₁exp(-E₁/RT) + A₂exp(-E₂/RT)·αᵐ](1-α)ⁿ"),
            "GAI": ("Initiation + Autocatalysis (GAI)", "dα/dt = A·exp(-E/RT)·(1-α)ⁿ¹[z₀+αⁿ²]"),
            "PAR": ("Parallel System (PAR)", "dα/dt = A·exp(-E₁/RT)f₁(α) + A₂exp(-E₂/RT)f₂(α)")
        }
        self.alphas = np.arange(0.15, 0.86, 0.05).round(2)
        self.ea = self.aT = self.xy = self.conversion_time_df = None
        self.current_method = "Vyazovkin"
        
        self._setup_headers()
        self._setup_modern_style()
        self._setup_menu()
        self._setup_status_bar()
        self._setup_main_interface()

    def _load_and_set_dpi(self):
        dpi = 100
        try:
            with open(self.config_path, 'r') as f:
                config = json.load(f)
                if isinstance(config.get('screen_dpi'), int):
                    dpi = config['screen_dpi']
        except (FileNotFoundError, json.JSONDecodeError):
            pass 
        
        plt.rcParams['figure.dpi'] = dpi
        self.dpi_var = tk.IntVar(value=dpi)

    def on_dpi_change(self):
        new_dpi = self.dpi_var.get()
        plt.rcParams['figure.dpi'] = new_dpi
        self.status_var.set(f"Screen plot resolution set to {new_dpi} DPI. New plots will use this setting.")
        try:
            with open(self.config_path, 'w') as f:
                json.dump({'screen_dpi': new_dpi}, f)
        except Exception as e:
            self.status_var.set(f"Could not save DPI setting: {e}")

    def _setup_headers(self):
        self.display_header_map = {
            "alpha": "α", "Ea_kJ_per_mol": "Eₐ (kJ/mol)", "R2": "R²", "StdErr_kJ": "Std. Err. (kJ)",
            "beta": "β (K/min)", "Time_min": "Time (min)", "Temp_K": "T (K)", "TG_pct": "TG (%)",
            "DSC":"DSC (mW/mg)", "DTG": "DTG", "DTA":"DTA", "t_alpha_days": "tₐ (days)",
            "Parameter": "Parameter", "Value": "Value", "Std. Error": "Std. Error", 'index': 'Parameter',
            "A_per_s": "A (s⁻¹)"
        }
        self.latex_header_map = {
            "alpha": r"$\alpha$", "Ea_kJ_per_mol": r"$E_\mathrm{a}$ (kJ/mol)", "R2": r"$R^2$",
            "StdErr_kJ": "Std. Err. (kJ)", "t_alpha_days": r"$t_\mathrm{a}$ (days)", "Parameter": "Parameter",
            "Value": "Value", "Std. Error": "Std. Error", 'index': 'Parameter',
            "A_per_s": r"$A$ (s$^{-1}$)"
        }

    def _setup_modern_style(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure('.', background=self.BG, foreground=self.FG, fieldbackground=self.BG_SECONDARY, borderwidth=0, focuscolor='none')

        style.configure("Primary.TButton", background=self.BRAND, foreground="white", font=("Segoe UI", 11, "bold"), padding=(18, 11), borderwidth=0)
        style.map("Primary.TButton", background=[('active', self.BRAND_ACCENT), ('pressed', '#4f46e5')])
        style.configure("Secondary.TButton", background="#e5e7eb", foreground=self.FG, font=("Segoe UI", 10), padding=(14, 9), borderwidth=0)
        style.map("Secondary.TButton", background=[('active', '#d1d5db'), ('pressed', '#9ca3af')])

        style.configure("TFrame", background=self.BG)
        style.configure("TLabel", background=self.BG, foreground=self.FG, font=("Segoe UI", 10))

        style.configure("Modern.TEntry", fieldbackground=self.BG_SECONDARY, borderwidth=1, insertcolor=self.FG, font=("Segoe UI", 10))
        style.map('TCombobox', fieldbackground=[('readonly', self.BG_SECONDARY)])
        style.configure("TCombobox", font=("Segoe UI", 10), borderwidth=1)

        style.configure("Modern.Treeview", background=self.BG_SECONDARY, foreground=self.FG, fieldbackground=self.BG_SECONDARY, borderwidth=0, rowheight=28, font=("Segoe UI", 10))
        style.configure("Modern.Treeview.Heading", background=self.BRAND, foreground="white", font=("Segoe UI", 11, "bold"), borderwidth=0, padding=8)
        style.map("Modern.Treeview.Heading", background=[('active', self.BRAND_ACCENT)])

        style.configure("Modern.TLabelframe", background=self.BG, bordercolor="#d1d5db")
        style.configure("Modern.TLabelframe.Label", background=self.BG, foreground=self.FG, font=("Segoe UI", 11, "bold"))
        style.configure("Modern.TCheckbutton", background=self.BG, foreground=self.FG, font=("Segoe UI", 10), focuscolor=self.BG)

    def _setup_menu(self):
        self.menubar = tk.Menu(self, bg=self.BG_SECONDARY, fg=self.FG, activebackground=self.BRAND, activeforeground="white", borderwidth=0, font=("Segoe UI", 10))
        self.config(menu=self.menubar)
        menu_configs = [("File", "mF"), ("View", "mV"), ("Analysis", "mP"), ("Export", "mE"), ("Theme", "mTheme"), ("Help", "mH")]
        for label, attr in menu_configs:
            menu = tk.Menu(self.menubar, tearoff=False, bg=self.BG_SECONDARY, fg=self.FG, activebackground=self.BRAND, activeforeground="white", font=("Segoe UI", 10))
            setattr(self, attr, menu)
            self.menubar.add_cascade(label=label, menu=menu)

        self.mF.add_command(label="📂 Import Data Files...", command=self.import_data)
        self.mF.add_separator()
        self.mF.add_command(label="🔄 Refresh Analysis (Isoconversional)", command=self._recalculate)

        self.mV.add_command(label="📊 Activation Energy (Ea)", command=lambda: self.show_table(self.ea, "Ea"))
        self.mV.add_command(label="🌡️ Conversion-Temperature (α-T)", command=lambda: self.show_table(self.aT, "α-T"))
        self.mV.add_command(label="📈 Temperature-Conversion (T-α)", command=lambda: self.show_table(self.aT, "T-α"))
        self.mV.add_command(label="📉 Linear Regression (X-Y)", command=lambda: self.show_table(self.xy, "X-Y"))
        self.mV.add_command(label="⏱️ Time-Conversion (t-α)", command=self.show_t_alpha)
        self.mV.add_separator()
        self.mV.add_command(label="🔥 Raw Thermograms", command=self.show_thermograms)
        self.mV.add_separator()
        self.mV.add_command(label="📊 TG/DTG Parameter Analysis", command=self.run_tg_dtg_analysis)
        self.mV.add_command(label="🔥 DSC Parameter Analysis", command=self.run_dsc_analysis)

        # --- MODIFICATION START: Removed old Lifetime Prediction menu item ---
        self.mP.add_command(label="[MODEL-FIT] Combined Kinetic Analysis (CKA)...", command=self.run_cka_analysis)
        self.mP.add_command(label="[MODEL-FIT] Autocatalytic Model Fit...", command=self.run_autocatalytic_fit)
        ### MODIFICATION START: Add Non-Parametric Analysis (NPA) ###
        self.mP.add_separator()
        self.mP.add_command(label="[MODEL-FREE] Non-Parametric Analysis (NPA)...", command=self.run_npa_analysis)
        ### MODIFICATION END ###

        self.mE.add_command(label="📝 Generate Word Report...", command=self.save_report)
        self.mE.add_command(label="📝 Generate LaTeX Report...", command=self.save_latex_report)
        self.mE.add_command(label="🗂️Raw Data Preprocessing...", command=self.export_preprocessed_data)
        self.mE.add_separator()
        self.mE.add_command(label="💾 Export All Tables...", command=self.save_tables)

        bg_colors = {"Light": "#F8F9FA", "Dark": "#2E3440", "White": "#FFFFFF"}
        accent_colors = {"Indigo": "#6366f1", "Purple": "#9B59B6", "Coral": "#FF6B6B", "Green": "#27AE60", "Blue": "#3498DB"}
        mBg = tk.Menu(self.mTheme, tearoff=0)
        self.mTheme.add_cascade(label="Background Color", menu=mBg)
        for name, color in bg_colors.items(): mBg.add_command(label=name, command=lambda c=color: self.update_theme(bg=c))
        mAccent = tk.Menu(self.mTheme, tearoff=0)
        self.mTheme.add_cascade(label="Accent Color", menu=mAccent)
        for name, color in accent_colors.items(): mAccent.add_command(label=name, command=lambda c=color: self.update_theme(accent=c))

        self.mTheme.add_separator()
        mDpi = tk.Menu(self.mTheme, tearoff=0)
        self.mTheme.add_cascade(label="Screen Resolution (DPI)", menu=mDpi)
        for dpi_val in [100, 150, 200, 300, 600]:
            mDpi.add_radiobutton(label=f"{dpi_val} DPI", variable=self.dpi_var,
                                 value=dpi_val, command=self.on_dpi_change)

        self.mH.add_command(label="📋 Environment Requirements", command=self.show_requirements)
        self.mH.add_command(label="📖 Instructions", command=self.show_instructions)
        self.mH.add_separator()
        self.mH.add_command(label="📧 Contact the Author", command=self.show_contact)

    def _setup_status_bar(self):
        self.status_var = tk.StringVar(value="Ready. Import data to begin.")
        self.status_frame = tk.Frame(self, bg="#343a40")
        self.status_frame.pack(side="bottom", fill="x")
        self.status_label = tk.Label(self.status_frame, textvariable=self.status_var, bg="#343a40", fg="white", anchor="w", padx=10, pady=5, font=("Segoe UI", 9))
        self.status_label.pack(side="left")

    def _setup_main_interface(self):
        self.main_container = ttk.Frame(self, style="TFrame")
        self.main_container.pack(fill="both", expand=True, padx=20, pady=20)
        self.show_welcome_screen()

    def show_welcome_screen(self):
        for widget in self.main_container.winfo_children():
            widget.destroy()
        self.main_container.grid_rowconfigure(0, weight=1)
        self.main_container.grid_columnconfigure(0, weight=1)
        welcome_frame = ttk.Frame(self.main_container, style="TFrame")
        welcome_frame.grid(row=0, column=0)
        ttk.Label(welcome_frame, text="TGAX Kinetics", font=("Segoe UI", 36, "bold"), foreground=self.BRAND, anchor="center").pack(pady=(0, 5))
        ttk.Label(welcome_frame, text="Comprehensive Kinetic Analysis & Lifetime Prediction", font=("Segoe UI", 16), anchor="center", foreground=self.FG_SECONDARY).pack(pady=(0, 40))
        ttk.Button(welcome_frame, text="Import Data Files to Begin", style="Primary.TButton", command=self.import_data).pack()
        ttk.Label(welcome_frame, text="Or use the File menu", font=("Segoe UI", 10), foreground=self.FG_SECONDARY, anchor="center").pack(pady=(10, 0))

    def update_theme(self, bg=None, accent=None):
        if bg:
            self.BG = bg
            if bg == "#2E3440": # Dark mode
                self.FG = "#ECEFF4"
                self.BG_SECONDARY = "#3B4252"
                self.FG_SECONDARY = "#D8DEE9"
            else: # Light modes
                self.FG = "#212529"
                self.BG_SECONDARY = "#FFFFFF"
                self.FG_SECONDARY = "#6c757d"

        if accent:
            self.BRAND = accent
            hex_color = accent.lstrip('#')
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            light_rgb = tuple(min(255, c + 40) for c in rgb)
            self.BRAND_ACCENT = f"#{light_rgb[0]:02x}{light_rgb[1]:02x}{light_rgb[2]:02x}"

        self.configure(bg=self.BG)
        self._setup_modern_style()
        self.menubar.configure(bg=self.BG_SECONDARY, fg=self.FG, activebackground=self.BRAND, activeforeground="white")
        for i in range(self.menubar.index('end') + 1):
             self.menubar.entryconfig(i, background=self.BG_SECONDARY, foreground=self.FG, activebackground=self.BRAND, activeforeground='white')

        self.status_label.config(bg="#343a40" if self.BG != "#2E3440" else "#4C566A")
        self.show_welcome_screen() # Redraw main screen with new colors
        self.update_idletasks()

    def _start_task(self, worker_func, *args, on_success=None, on_error=None, **kwargs):
        if self.is_task_running:
            messagebox.showwarning("Busy", "Another task is currently running. Please wait.")
            return

        self.is_task_running = True
        self._disable_ui()
        self.status_var.set("Processing... please wait.")

        thread = threading.Thread(target=self._worker_wrapper, 
                                  args=(worker_func, args, kwargs, on_success, on_error))
        thread.daemon = True
        thread.start()
        self._check_queue()

    def _worker_wrapper(self, worker_func, args, kwargs, on_success, on_error):
        try:
            result = worker_func(*args, **kwargs)
            self.task_queue.put(('success', result, on_success))
        except Exception as e:
            error_info = (e, traceback.format_exc())
            self.task_queue.put(('error', error_info, on_error))

    def _check_queue(self):
        try:
            status, data, callback = self.task_queue.get_nowait()
            
            self.is_task_running = False
            self._enable_ui()

            if status == 'success':
                if callback:
                    callback(data)
            elif status == 'error':
                if callback:
                    callback(data)
                else:
                    self._on_task_error(data)
        except Empty:
            self.after(100, self._check_queue)
    
    def _on_task_error(self, error_info):
        exception, tb_str = error_info
        self.status_var.set("An error occurred.")
        messagebox.showerror("Error", f"An error occurred in the background task:\n\n{exception}")
        print(tb_str)

    def _disable_ui(self):
        self.menu_states.clear()
        for i in range(self.menubar.index('end') + 1):
            state = self.menubar.entrycget(i, "state")
            self.menu_states[i] = state
            self.menubar.entryconfig(i, state="disabled")
    
    def _enable_ui(self):
        for i in range(self.menubar.index('end') + 1):
            original_state = self.menu_states.get(i, "normal")
            self.menubar.entryconfig(i, state=original_state)

    def show_requirements(self):
        win = tk.Toplevel(self)
        win.title("Environment Requirements")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception:
            pass
        win.configure(bg=self.BG_SECONDARY)
        frame = ttk.Frame(win, padding=20)
        frame.pack(fill="both", expand=True)
        text_content = textwrap.dedent("""
            To run TGAX Kinetics successfully, please ensure you have a Python environment with the following libraries installed. You can install them using pip.

            1.  Python: Version 3.8 or newer.

            2.  Required Libraries:
                - pandas: For data manipulation and analysis.
                - numpy: For numerical operations.
                - matplotlib: For creating plots and graphs.
                - scipy: For scientific computing and optimization.
                - openpyxl: To read .xlsx files.
                - python-docx: To export reports to .docx format.

            Example installation command for all libraries:
            pip install pandas numpy matplotlib scipy openpyxl python-docx
            """)
        text_widget = tk.Text(frame, wrap=tk.WORD, bg=self.BG_SECONDARY, fg=self.FG, relief=tk.FLAT, padx=10, pady=10, font=("Segoe UI", 10), highlightthickness=0, borderwidth=0)
        text_widget.insert(tk.END, text_content); text_widget.config(state=tk.DISABLED)
        text_widget.pack(expand=True, fill=tk.BOTH)
        
        win.update_idletasks()
        win.geometry(f"{win.winfo_reqwidth()+20}x{win.winfo_reqheight()+20}")
        win.resizable(False, False)

    def show_instructions(self):
        win = tk.Toplevel(self)
        win.title("Instructions")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception:
            pass
        win.configure(bg=self.BG_SECONDARY)
        frame = ttk.Frame(win, padding=20)
        frame.pack(fill="both", expand=True)
        text_content = textwrap.dedent("""
            Welcome to TGAX Kinetics! Follow these steps to perform a kinetic analysis:

            1.  Import Data:
                - Go to File > Import Data... or click the main button.
                - Select one or more TGA data files (.xlsx, .csv, .txt).

            2.  Set Calculation Parameters:
                - A settings dialog will appear for the isoconversional analysis.
                - Choose the kinetic method (Vyazovkin is recommended) and define the conversion (α) range.
                - Click "Start Calculation". The program will process in the background.

            3.  View Results:
                - When calculations are done, the main results (Ea plot/table) will appear. Use the "View" menu to open other tables and plots.
                - The "View" menu now also contains independent "TG/DTG Parameter Analysis" and "DSC Parameter Analysis" options. These functions have their own dialogs for settings.

            4.  Analyze & Predict:
                - Use the "Analysis" menu for additional, independent analyses like CKA, Global Fit, and Lifetime Prediction. These also run in the background.
                - The "Autocatalytic Model Fit" dialog now includes a fully functional "cross-validation" checkbox for rigorous model testing.
                - All fit result windows (CKA, Global Fit, etc.) now include a "Plot Model TG vs. Exp" button to check the integral fit quality.

            5.  Export Your Work:
                - The "Export" menu provides multiple options to save reports, tables, and plots.
                - All plot windows now have a navigation toolbar and a right-click context menu to save the image.
            """)
        text_widget = tk.Text(frame, wrap=tk.WORD, bg=self.BG_SECONDARY, fg=self.FG, relief=tk.FLAT, padx=10, pady=10, font=("Segoe UI", 10), highlightthickness=0, borderwidth=0)
        text_widget.insert(tk.END, text_content); text_widget.config(state=tk.DISABLED)
        text_widget.pack(expand=True, fill=tk.BOTH)
        
        win.update_idletasks()
        win.geometry(f"{win.winfo_reqwidth()+40}x{win.winfo_reqheight()+20}")
        win.resizable(False, False)


    def show_contact(self):
        win = tk.Toplevel(self)
        win.title("Contact the Author")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception:
            pass
        win.configure(bg=self.BG_SECONDARY)
        frame = ttk.Frame(win, padding="20"); frame.pack(expand=True, fill=tk.BOTH)
        ttk.Label(frame, text="Author: Xing Xiaohui", font=("Segoe UI", 11)).pack(pady=4)
        ttk.Label(frame, text="Affiliation: Beijing Institute of Technology", font=("Segoe UI", 11)).pack(pady=4)
        ttk.Label(frame, text="Email:", font=("Segoe UI", 11)).pack(pady=(12,0))
        ttk.Label(frame, text="361143458@163.com", font=("Segoe UI", 11, "bold")).pack()
        
        win.update_idletasks()
        win.geometry(f"{win.winfo_reqwidth()+20}x{win.winfo_reqheight()+20}")
        win.resizable(False, False)

    def simple_load_one(self, path):
        beta_match = re.search(r"(\d+(?:\.\d+)?)\s*[kK]", Path(path).stem)
        if not beta_match:
            raise ValueError(f"Could not extract heating rate (β) from filename: {Path(path).name}.\nEnable advanced preprocessing if format is non-standard.")
        beta = float(beta_match.group(1))

        names = ["Time_min", "Temp_C", "TG_pct", "DSC"]
        try:
            if str(path).lower().endswith(".xlsx"):
                df = pd.read_excel(path, skiprows=1, names=names)
            else:
                df = pd.read_csv(path, skiprows=1, sep=None, engine="python", names=names)
        except Exception as e:
            raise ValueError(f"Simple loading failed for {Path(path).name}. Enable advanced preprocessing. Error: {e}")

        for col in names:
            if col not in df.columns:
                df[col] = 0.0
        
        df = df.astype(float)
        return beta, df

    def import_data(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        paths = filedialog.askopenfilenames(filetypes=[("Data files", "*.xlsx *.csv *.txt")])
        if not paths: return

        def pre_scan_for_temp(path):
            try:
                df_raw = pd.read_csv(path, sep=None, engine='python', header=None, on_bad_lines='skip', usecols=lambda x: isinstance(x, int) and x < 5 or x is None, encoding_errors='ignore')
            except Exception:
                try:
                    df_raw = pd.read_excel(path, header=None, usecols=lambda x: isinstance(x, int) and x < 5 or x is None)
                except Exception:
                    return 35.0
            max_val = 0
            for _, row in df_raw.head(100).iterrows():
                for item in row:
                    try:
                        val = float(item)
                        if 20 < val < 1500:
                            max_val = max(max_val, val)
                    except (ValueError, TypeError):
                        continue
            return max_val if max_val > 0 else 800.0

        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            max_temps = [pre_scan_for_temp(p) for p in paths]
        end_temp_suggestion = max(max_temps) - 10 if max_temps else 800
        temp_suggestion = (35.0, end_temp_suggestion)

        self.status_var.set(f"{len(paths)} file(s) selected. Opening settings..."); self.update_idletasks()
        
        dlg = CalculationSettingsDialog(self, temp_range_suggestion=temp_suggestion)
        self.wait_window(dlg)
        
        if dlg.result:
            self._start_task(self._import_data_worker, paths, dlg.result, on_success=self._on_import_data_success)
        else:
            self.status_var.set("Calculation cancelled. Ready.")

    def _import_data_worker(self, paths, settings):
        local_dfs, local_preprocessed = {}, {}
        local_dfs, local_preprocessed, local_mass_params = {}, {}, {}
        processed_count, errors = 0, []
        
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            for p in paths:
                try:
                    if settings["preprocess_enabled"]:
                        beta, df_processed = preprocess_file(p, settings["temp_range_c"])
                    else:
                        beta, df_processed = self.simple_load_one(p)
                    
                    df_processed["Temp_K"] = df_processed["Temp_C"] + 273.15
                    m0, m_inf = df_processed["TG_pct"].iloc[0], df_processed["TG_pct"].min()
                    
                    ### MODIFICATION START: Store m0 and m_inf ###
                    local_mass_params[beta] = (m0, m_inf)
                    ### MODIFICATION END ###
                    
                    if abs(m0 - m_inf) < EPS:
                        if df_processed.attrs.get('kinetic_analysis_possible', False):
                            errors.append(f"- {Path(p).name}: No significant mass loss.")
                    
                    df_processed["alpha"] = (m0 - df_processed["TG_pct"]) / (m0 - m_inf) if abs(m0 - m_inf) > EPS else 0
                    df_processed.sort_values("Temp_K", inplace=True)
                    df_processed['alpha'] = df_processed['alpha'].clip(0, 1).cummax()

                    df_processed["DTG_min"] = np.gradient(df_processed["TG_pct"], df_processed["Time_min"])
                    df_processed["DTG"]  = np.gradient(df_processed["TG_pct"], df_processed["Temp_K"])
                    
                    ### MODIFICATION START: Renamed dadt -> dAdT ###
                    df_processed["dAdT"] = np.gradient(df_processed["alpha"], df_processed["Temp_K"])
                    ### MODIFICATION END ###
                    
                    df_processed["DTA"]  = np.gradient(df_processed["DSC"], df_processed["Temp_K"])
                    
                    if not df_processed["DSC"].empty:
                        df_processed["DSC"] -= df_processed["DSC"].iloc[0]
                    if not df_processed["DTA"].empty:
                        df_processed["DTA"] -= df_processed["DTA"].iloc[0]

                    local_dfs[beta] = df_processed
                    local_preprocessed[f"{beta:.2f}K_per_min"] = df_processed
                    processed_count += 1
                except Exception as e:
                    errors.append(f"- {Path(p).name}: {e}")
            
            warning_messages = "\n".join(str(warn.message) for warn in w)

        if not local_dfs:
            raise ValueError("No files were successfully processed.")

        kinetic_analysis_possible = all(df.attrs.get('kinetic_analysis_possible', True) for df in local_dfs.values())
        
        ea, aT, xy = pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
        if kinetic_analysis_possible:
            a_min, a_max, a_step = settings["alpha_range"]
            alphas = np.round(np.arange(a_min, a_max + EPS, a_step), 4)
            method_map = {"Friedman": build_friedman_tables, "KAS": build_kas_tables, "OFW": build_ofw_tables, "Vyazovkin": build_vyazovkin_tables}
            calculation_func = method_map[settings["method"]]
            ea, aT, xy = calculation_func(local_dfs, alphas)
            
        return {
            "dfs": local_dfs, "preprocessed_files": local_preprocessed, "ea": ea, "aT": aT, "xy": xy,
            "settings": settings, "errors": errors, "warnings": warning_messages, "processed_count": processed_count,
            "kinetic_analysis_possible": kinetic_analysis_possible,
            "mass_loss_parameters": local_mass_params
        }

    def _on_import_data_success(self, data):
        self.dfs.clear(); self.preprocessed_files.clear(); self.sample_masses.clear()
        self.dfs.clear(); self.preprocessed_files.clear(); self.sample_masses.clear(); self.mass_loss_parameters.clear()
        self.ea = self.aT = self.xy = self.conversion_time_df = self.fit_results = self.cka_results = None

        self.dfs = data['dfs']
        self.preprocessed_files = data['preprocessed_files']
        self.ea, self.aT, self.xy = data['ea'], data['aT'], data['xy']
        self.current_method = data['settings']['method']
        a_min, a_max, a_step = data['settings']["alpha_range"]
        self.alphas = np.round(np.arange(a_min, a_max + EPS, a_step), 4)
        
        ### MODIFICATION START: Store m0/m_inf ###
        self.mass_loss_parameters = data['mass_loss_parameters']
        ### MODIFICATION END ###

        if data['warnings']:
            messagebox.showwarning("Preprocessing Warnings", data['warnings'], parent=self)
        if data['errors']:
            messagebox.showwarning("Processing Issues", "Some files could not be processed:\n\n" + "\n".join(data['errors']))

        for widget in self.main_container.winfo_children(): widget.destroy()

        if not data['kinetic_analysis_possible']:
            messagebox.showwarning("Kinetic Analysis Disabled",
                                   "One or more files were missing a Mass/Weight column. All TGA-based kinetic analyses have been disabled.")
            self.mF.entryconfig(2, state="disabled")
            # --- MODIFICATION START: Check for item type before changing state ---
            for i in range(self.mP.index('end') + 1):
                try:
                    item_type = self.mP.type(i)
                    if item_type == "command":
                        self.mP.entryconfig(i, state="disabled")
                except tk.TclError:
                    pass
            # --- MODIFICATION END ---
            self.mV.entryconfig(8, state="disabled")
        else:
            self.mF.entryconfig(2, state="normal")
            # --- MODIFICATION START: Check for item type before changing state ---
            for i in range(self.mP.index('end') + 1):
                try:
                    item_type = self.mP.type(i)
                    if item_type == "command":
                        self.mP.entryconfig(i, state="normal")
                except tk.TclError:
                    pass
            # --- MODIFICATION END ---
            self.mV.entryconfig(8, state="normal")
        
        self.status_var.set(f"{data['processed_count']} files processed. View thermograms or run analysis.")
        
        if data['kinetic_analysis_possible']:
            self._on_recalculate_success({
                "ea": self.ea, "aT": self.aT, "xy": self.xy, "method": self.current_method
            })
        else:
            self.show_thermograms()
    def export_preprocessed_data(self):
        if not self.preprocessed_files:
            return messagebox.showwarning("No Data", "No preprocessed data available to export. Please import and process files first.")
        
        file_path = filedialog.asksaveasfilename(
            title="Export Preprocessed Data",
            initialfile="Preprocessed_Data.xlsx",
            filetypes=[("Excel file", "*.xlsx"), ("CSV files", "*.csv"), ("Text files", "*.txt")]
        )
        if not file_path: return

        p = Path(file_path)
        try:
            if p.suffix == '.xlsx':
                with pd.ExcelWriter(p, engine='openpyxl') as writer:
                    for basename, df in self.preprocessed_files.items():
                        export_df = df[['Time_min', 'Temp_C', 'TG_pct', 'DSC']].copy()
                        export_df['Heating Rate (K/min)'] = float(basename.split('K')[0])
                        export_df.to_excel(writer, sheet_name=basename, index=False)
                messagebox.showinfo("Success", f"All preprocessed data saved to:\n{p}")
            else:
                directory, stem, suffix = p.parent, p.stem, p.suffix
                sep = ',' if suffix == '.csv' else '\t'
                for basename, df in self.preprocessed_files.items():
                    export_df = df[['Time_min', 'Temp_C', 'TG_pct', 'DSC']].copy()
                    export_df['Heating Rate (K/min)'] = float(basename.split('K')[0])
                    new_path = directory / f"{stem}_{basename}{suffix}"
                    export_df.to_csv(new_path, index=False, sep=sep, float_format='%.6g')
                messagebox.showinfo("Success", f"All preprocessed files saved in folder:\n{directory}")
        except Exception as e:
            messagebox.showerror("Error Saving", f"Could not save files. Error: {e}")

    def _recalculate(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if not self.dfs: return messagebox.showwarning("No Data", "Please import data before running an analysis.")
        
        kinetic_analysis_possible = all(df.attrs.get('kinetic_analysis_possible', True) for df in self.dfs.values())
        if not kinetic_analysis_possible:
            return messagebox.showwarning("Analysis Disabled", "Cannot run kinetic analysis because one or more files are missing mass data.")

        self._start_task(self._recalculate_worker, on_success=self._on_recalculate_success)

    def _recalculate_worker(self):
        method_map = {"Friedman": build_friedman_tables, "KAS": build_kas_tables, "OFW": build_ofw_tables, "Vyazovkin": build_vyazovkin_tables}
        calculation_func = method_map[self.current_method]
        ea, aT, xy = calculation_func(self.dfs, self.alphas)
        return {"ea": ea, "aT": aT, "xy": xy, "method": self.current_method}

    def _on_recalculate_success(self, data):
        self.ea, self.aT, self.xy = data["ea"], data["aT"], data["xy"]
        if data["method"] == "Vyazovkin":
            self.mV.entryconfig("📉 Linear Regression (X-Y)", state="disabled")
        else:
            self.mV.entryconfig("📉 Linear Regression (X-Y)", state="normal")
        self.status_var.set(f"✓ Calculation complete using {data['method']}. View results.")
        self.show_table(self.ea, "Ea")

    def show_table(self, df, title):
        if df is None:
            return messagebox.showwarning("No Data", "Please import and process data first.")

        if df.empty and title == "X-Y":
            return messagebox.showinfo("Not Applicable", f"The {self.current_method} method does not have an X-Y plot.")

        win = tk.Toplevel(self); win.configure(bg=self.BG); win.title(f"{title} Table & Plot")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("900x600")

        header_map = self.display_header_map.copy()
        y_col_names = {"Friedman": "ln(β·dα/dt)", "KAS": "ln(β/T²)", "OFW": "ln(β)"}
        if title == "X-Y": header_map["Y"] = y_col_names.get(self.current_method, "Y")
        new_cols = {c: f"T at {c.split('_')[1]}" for c in df.columns if c.startswith("T_")}
        header_map.update(new_cols)

        display_df = df.rename(columns=header_map)
        cols = list(map(str, display_df.columns))

        frame = ttk.Frame(win, padding=10)
        frame.pack(fill="both", expand=True)

        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Label(btn_frame, text=f"{title} Data", font=("Segoe UI", 12, "bold")).pack(side="left")
        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(df, title), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Pop-out Plot", command=lambda: self.plot_data(title), style="Secondary.TButton").pack(side="right", padx=5)

        tv = ttk.Treeview(frame, columns=cols, show="headings", style="Modern.Treeview")
        for c in cols: tv.heading(c, text=c); tv.column(c, width=110, anchor="e")
        for _, r in display_df.iterrows(): tv.insert("", "end", values=[f"{v:.6g}" if isinstance(v, (int, float)) else v for v in r])

        ysb = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=tv.yview)
        xsb = ttk.Scrollbar(frame, orient=tk.HORIZONTAL, command=tv.xview)
        tv.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        tv.pack(side="left", fill="both", expand=True)
        ysb.pack(side="right", fill="y")
        xsb.pack(side="bottom", fill="x")

    def _export_table(self, df, title):
        p = filedialog.asksaveasfilename(
            defaultextension=".csv",
            initialfile=f"{title.replace('-','_')}.csv",
            filetypes=[("CSV file", "*.csv"), ("Text file", "*.txt"), ("Excel file", "*.xlsx"), ("All files", "*.*")]
        )
        if not p: return
        
        try:
            p = Path(p)
            if p.suffix == '.csv':
                df.to_csv(p, index=False, float_format='%.6g')
            elif p.suffix == '.txt':
                df.to_csv(p, index=False, sep='\t', float_format='%.6g')
            elif p.suffix == '.xlsx':
                df.to_excel(p, index=False)
            messagebox.showinfo("Success", f"Table saved to {p.name}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Could not save the file.\nError: {e}")
    
    def add_smart_legend(self, fig, ax, num_items):
        if num_items <= 0:
            return

        handles, labels = ax.get_legend_handles_labels()
        if hasattr(ax, 'twinx_ax'):
            h2, l2 = ax.twinx_ax.get_legend_handles_labels()
            handles.extend(h2)
            labels.extend(l2)

        if not handles: return

        xlim = ax.get_xlim()
        ylim = ax.get_ylim()
        x_mid = (xlim[0] + xlim[1]) / 2
        y_mid = (ylim[0] + ylim[1]) / 2

        quadrants = {
            'upper right': 0,
            'upper left': 0,
            'lower left': 0,
            'lower right': 0,
        }

        lines = ax.get_lines()
        if hasattr(ax, 'twinx_ax'):
            lines.extend(ax.twinx_ax.get_lines())

        for line in lines:
            xdata, ydata = line.get_xdata(), line.get_ydata()
            for x, y in zip(xdata, ydata):
                if x > x_mid and y > y_mid:
                    quadrants['upper right'] += 1
                elif x <= x_mid and y > y_mid:
                    quadrants['upper left'] += 1
                elif x <= x_mid and y <= y_mid:
                    quadrants['lower left'] += 1
                elif x > x_mid and y <= y_mid:
                    quadrants['lower right'] += 1
        
        best_loc = min(quadrants, key=quadrants.get)
        
        # --- MODIFICATION START: Legend Size Coordination ---
        # Increased font size for better readability in the application window.
        # The SaveOptionsDialog will handle reverting to a smaller size for export.
        ax.legend(handles, labels, loc=best_loc, fontsize=9, frameon=False)
        # --- MODIFICATION END: Legend Size Coordination ---

    def _create_plot_window(self, fig, title, is_thermo_plot=False):
        if title in self.plots and self.plots[title].winfo_exists():
            return self.plots[title].focus_force()

        win = tk.Toplevel(self)
        win.title(title)
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception:
            pass

        # --- MODIFICATION START: Hide Edge Ticks ---
        # Loop through all axes in the figure and apply the hiding function
        for ax in fig.get_axes():
            self._hide_edge_tick_labels(ax)
        fig.tight_layout() # Re-run tight_layout to adjust for removed labels
        # --- MODIFICATION END ---

        canvas_frame = ttk.Frame(win)
        canvas_frame.pack(fill="both", expand=True)
        
        canvas = FigureCanvasTkAgg(fig, master=canvas_frame)
        canvas.draw()
        
        toolbar = NavigationToolbar2Tk(canvas, canvas_frame)
        toolbar.update()
        canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        canvas.mpl_connect("button_press_event", lambda ev: self.show_save_dialog(ev, fig, is_thermo_plot))
        self.plots[title] = win
        plt.close(fig)

    def _hide_edge_tick_labels(self, ax):
        """Hides the first and last tick labels on both x and y axes."""
        for axis in [ax.xaxis, ax.yaxis]:
            ticks = axis.get_major_ticks()
            # Ensure there are enough ticks to hide the edges
            if len(ticks) > 1:
                ticks[0].label1.set_visible(False)
                ticks[-1].label1.set_visible(False)

    def plot_data(self, title):
        fig, ax = plt.subplots(figsize=(6, 4.5))
        
        if title == "Ea":
            ax.errorbar(self.ea["alpha"], self.ea["Ea_kJ_per_mol"], yerr=self.ea["StdErr_kJ"], fmt='-o', capsize=3, label=f"$E_\\mathrm{{a}}$ ({self.current_method})")
            ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel(r"Activation Energy, $E_\mathrm{a}$ (kJ mol$^{-1}$)")
        elif title == "α-T":
            num_curves = len([c for c in self.aT.columns if c.startswith("T_")])
            for col in [c for c in self.aT.columns if c.startswith("T_")]:
                β = col.split('_')[1][:-5]
                ax.plot(self.aT["alpha"], self.aT[col], "-o", label=f"$\\beta$ = {β} K/min")
            ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel("Temperature (K)")
            self.add_smart_legend(fig, ax, num_curves)
        elif title == "T-α":
            num_curves = len([c for c in self.aT.columns if c.startswith("T_")])
            for col in [c for c in self.aT.columns if c.startswith("T_")]:
                β = col.split('_')[1][:-5]
                ax.plot(self.aT[col], self.aT["alpha"], "-o", label=f"$\\beta$ = {β} K/min")
            ax.set_xlabel("Temperature (K)"); ax.set_ylabel(r"Conversion, $\alpha$")
            self.add_smart_legend(fig, ax, num_curves)
        elif title == "X-Y":
            if self.xy.empty: messagebox.showinfo("Not Applicable", f"The {self.current_method} method does not have a linear X-Y plot."); plt.close(fig); return
            y_labels = {"Friedman": r"$\ln(\beta \cdot \mathrm{d}\alpha/\mathrm{d}t)$","KAS": r"$\ln(\beta/T^2)$","OFW": r"$\ln(\beta)$"}
            ax.set_ylabel(y_labels.get(self.current_method, "Y-Value"))
            
            alphas_with_data = self.xy['alpha'].unique()
            
            for a in alphas_with_data:
                s = self.xy[self.xy["alpha"] == a]
                if s.empty: continue
                slope, inter = np.polyfit(s["X"], s["Y"], 1)
                p = ax.scatter(s["X"], s["Y"], s=15, label=f"$\\alpha$ = {a:.2f}")
                ax.plot(s["X"], slope * s["X"] + inter, color=p.get_facecolor()[0], lw=1.0)
            ax.set_xlabel(r"1/T (K$^{-1}$)")
            self.add_smart_legend(fig, ax, len(alphas_with_data))
        elif title == "t-α":
            for β, df in self.dfs.items(): ax.plot(df["Time_min"], df["alpha"], label=f"$\\beta$ = {β:.2f} K/min")
            ax.set_xlabel("Time (min)"); ax.set_ylabel(r"Conversion, $\alpha$")
            self.add_smart_legend(fig, ax, len(self.dfs))
        
        ax.margins(x=0)
        fig.tight_layout()
             
        self._create_plot_window(fig, f"{title} Plot")

    def show_save_dialog(self, event, fig, is_thermo_plot=False):
        if event.button == 3: SaveOptionsDialog(self, fig, is_thermo_plot)

    def _thermo_plot(self, kind, label):
        tag = f"{label} Plot"
        fig, ax = plt.subplots(figsize=(5.5, 4.5))
        num_curves = len(self.dfs)
        for i, (β, df) in enumerate(self.dfs.items()): ax.plot(df["Temp_K"], df[kind], label=f"$\\beta$ = {β:.2f} K/min")
        ax.set_xlabel("T (K)"); ax.margins(x=0.0)

        if kind == "TG_pct": 
            ax.set_ylabel("TG (%)")
        elif kind == 'DSC':
            ax.set_ylabel("Heat flow (a.u.)")
            ax.set_yticks([])
        else: 
            ax.set_ylabel(f"{label} (a.u.)")
            ax.set_yticks([])
        
        self.add_smart_legend(fig, ax, num_curves)
        fig.tight_layout()

        is_fillable = kind in ["DSC", "DTA", "DTG"]
        self._create_plot_window(fig, tag, is_thermo_plot=is_fillable)

    def show_thermograms(self):
        if not self.dfs: return messagebox.showwarning("No Data", "Please import data first.")
        win = tk.Toplevel(self); win.title("Thermograms")
        win.geometry("800x600")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass
        
        merge = pd.concat(self.dfs.values(), keys=self.dfs.keys(), names=["beta", "idx"]).reset_index(level=0)

        display_df = merge.rename(columns=self.display_header_map)
        cols = list(display_df.columns.intersection(["β (K/min)", "T (K)", "TG (%)", "DSC (mW/mg)", "DTG", "DTA"]))

        frame = ttk.Frame(win, padding=10)
        frame.pack(fill="both", expand=True)

        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        for k, label in {"TG_pct": "TG", "DSC": "DSC", "DTG": "DTG", "DTA": "DTA"}.items():
            ttk.Button(btn_frame, text=f"Plot {label}", command=lambda k=k, l=label: self._thermo_plot(k, l), style="Secondary.TButton").pack(side="left", padx=4)
        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(merge, "thermo_raw_data"), style="Secondary.TButton").pack(side="right", padx=4)

        tv = ttk.Treeview(frame, columns=cols, show="headings", height=18, style="Modern.Treeview")
        for c in cols: tv.heading(c, text=c); tv.column(c, width=110, anchor="e")
        for _, r in display_df[cols].head(1000).iterrows(): tv.insert("", "end", values=[f"{val:.6g}" for val in r])
        tv.pack(fill="both", expand=True, pady=(5,0))


    def show_t_alpha(self):
        if not self.dfs: return messagebox.showwarning("No Data", "Please import data first.")
        win = tk.Toplevel(self); win.title("Time-Conversion (t-α) Table & Plot")
        win.geometry("800x600")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass

        all_dfs = pd.concat(self.dfs.values(), keys=self.dfs.keys(), names=["beta", "idx"]).reset_index()

        display_df = all_dfs.rename(columns=self.display_header_map)
        cols = ["β (K/min)", "Time (min)", "α", "T (K)"]

        frame = ttk.Frame(win, padding=10)
        frame.pack(fill="both", expand=True)

        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(all_dfs[["beta", "Time_min", "alpha", "Temp_K"]], "t_alpha_data"), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Pop-out Plot", command=lambda: self.plot_data("t-α"), style="Secondary.TButton").pack(side="right", padx=5)

        tv = ttk.Treeview(frame, columns=cols, show="headings", style="Modern.Treeview")
        for c in cols: tv.heading(c, text=c); tv.column(c, anchor="e", width=130)
        for _, r in display_df[cols].iterrows(): tv.insert("", "end", values=[f"{v:.4g}" for v in r])
        tv.pack(fill="both", expand=True, pady=(5,0))

    def global_fit(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if not self.dfs or self.ea is None: 
            return messagebox.showwarning("No Data", "Please run an isoconversional analysis first to get the Ea curve.")
        prompt = "Enter α range for Global Fit (e.g., 0.15,0.85):"; initial = f"{self.alphas.min()},{self.alphas.max()}"
        rng = simpledialog.askstring("Global Fit Settings", prompt, initialvalue=initial, parent=self)
        if not rng: return
        try:
            a_min, a_max = map(float, rng.split(','))
            if not (0 < a_min < a_max < 1): raise ValueError
            fit_alphas = np.round(np.arange(a_min, a_max + EPS, 0.01), 4)
        except (ValueError, IndexError): return messagebox.showerror("Error", "Invalid α range format.")
        
        # We pass self.dfs to the worker, which will then call the refactored logic function
        self._start_task(self._global_fit_worker, self.dfs, fit_alphas, on_success=self._on_global_fit_success)

    def _global_fit_logic(self, dfs, alphas):
        """
        Refactored logic for global fitting that can be called by any workflow.
        ### MODIFICATION: This function now fits in the d(alpha)/dt (time) domain ###
        """
        ea_ser = self.ea.set_index("alpha")["Ea_kJ_per_mol"] * 1000
        data = prep_arrays(dfs, ea_ser, alphas) # prep_arrays now returns data for d(alpha)/dt
        if len(data[0]) < 10: 
            raise ValueError("Too few data points for a reliable fit in the selected range.")

        # resid function now expects d(alpha)/dt, so initial guess for logA is adjusted (e.g., 30 for 1/min)
        res = least_squares(resid, [1, 1, 0, 30], args=(data,), bounds=([0, 0, 0, 10], [5, 5, 5, 50]), loss="soft_l1", x_scale='jac', method='trf')
        if not res.success: 
            raise RuntimeError("The optimization did not converge.")

        m, n, p_, logA = res.x
        try:
            n_samples, n_params = len(data[0]), len(res.x)
            cov = np.linalg.pinv(res.jac.T @ res.jac) * res.cost * 2 / (n_samples - n_params)
            se = np.sqrt(np.diag(cov))
        except: 
            se = np.full(4, np.nan)
            cov = np.full((4, 4), np.nan)

        return {
            "model_name": "GlobalFit",
            ### MODIFICATION: A is now A_per_min, converted to A_per_s for display consistency ###
            "params": {"m": m, "n": n, "p": p_, "A_per_s": math.exp(logA) / 60, "Ea_kJ_per_mol": self.ea['Ea_kJ_per_mol'].mean()}, 
            "errors": {"m": se[0], "n": se[1], "p": se[2], "A_per_s": se[3] * math.exp(logA) / 60, "Ea_kJ_per_mol": self.ea['Ea_kJ_per_mol'].std()},
            "cov": cov, "raw_params": res.x, "param_order": ['m', 'n', 'p', 'logA']
        }

    def _global_fit_worker(self, dfs, fit_alphas):
        """Worker wrapper for the global fit logic."""
        return self._global_fit_logic(dfs, fit_alphas)

    def _on_global_fit_success(self, results):
        self.fit_results = results
        self.last_fitted_model = {
            "model_name": "GlobalFit", "params": self.fit_results['params'],
            "raw_params": self.fit_results['raw_params'], "cov": self.fit_results['cov'],
            "param_order": self.fit_results['param_order'], "source": "Global Fit",
            "Ea_source": "isoconversional" # Mark that Ea is from iso
        }
        self.status_var.set("✓ Global Fit complete. Showing results.")
        
        win = tk.Toplevel(self, bg=self.BG); win.title("Global Fit Results")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("800x550")
        
        frame = ttk.Frame(win, padding="20"); frame.pack(fill="both", expand=True)
        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Label(btn_frame, text="Global Fit Results (Sestak-Berggren Model)", font=("Segoe UI", 12, "bold")).pack(side="left")
        
        fit_df = pd.DataFrame([self.fit_results['params'], self.fit_results['errors']]).T
        fit_df.columns = ['Value', 'Std. Error']; fit_df.index.name = 'Parameter'
        fit_df = fit_df.reset_index()

        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(fit_df, "global_fit_results"), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Predict Conversion Time...", command=self.launch_prediction_dialog, style="Secondary.TButton").pack(side="right", padx=5)
        
        ### MODIFICATION START: Added new buttons ###
        ttk.Button(btn_frame, text="Export TG Fit Table", command=self.export_tg_fit_table, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Model TG vs. Exp", command=self.plot_tg_fit_comparison, style="Secondary.TButton").pack(side="right", padx=5)
        ### MODIFICATION END ###
        
        ttk.Button(btn_frame, text="Plot Fit vs. Experimental", command=self.plot_global_fit_comparison, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Compare Mechanism Shape...", command=lambda: self.launch_mechanism_comparison(fit_type='global_fit'), style="Secondary.TButton").pack(side="right", padx=5)
        
        p = self.fit_results['params']
        eq_text = f"$f(\\alpha) = \\alpha^{{{p['m']:.2f}}} (1-\\alpha)^{{{p['n']:.2f}}} [-\\ln(1-\\alpha)]^{{{p['p']:.2f}}}$"
        fig_eq = plt.figure(figsize=(6, 1.5), facecolor=self.BG)
        fig_eq.text(0.5, 0.5, eq_text, va='center', ha='center', fontsize=12, color=self.FG)
        canvas = FigureCanvasTkAgg(fig_eq, master=frame); canvas.draw()
        canvas.get_tk_widget().pack(fill='x', expand=False, pady=(10,15)); plt.close(fig_eq)

        tv_cols = ("Parameter", "Value", "Std. Error")
        tv = ttk.Treeview(frame, columns=tv_cols, show="headings", height=5, style="Modern.Treeview"); tv.pack(fill="x", expand=True)
        for col_id in tv_cols: tv.heading(col_id, text=col_id); tv.column(col_id, anchor="center")
        
        ### MODIFICATION START: Standardized param display ###
        param_order = ["Ea_kJ_per_mol", "A_per_s", "m", "n", "p"]
        for name in param_order:
            val, err = self.fit_results["params"].get(name, 0), self.fit_results["errors"].get(name, 0)
            display_name = self.display_header_map.get(name, name)
            if name == "Ea_kJ_per_mol":
                tv.insert("", "end", values=(display_name, f"{val:.6g}", f"{err:.2e} (Mean/StdDev from Iso)"))
            else:
                tv.insert("", "end", values=(display_name, f"{val:.6g}", f"{err:.2e}"))
        ### MODIFICATION END ###

    def plot_global_fit_comparison(self):
        if not self.fit_results: return
        
        ### MODIFICATION START: UNIFICATION OF GLOBAL FIT ###
        # This plot now shows d(alpha)/dt (1/min) vs. Temperature
        
        fig, ax = plt.subplots(figsize=(6, 5))
        params = self.fit_results['params']
        m, n, p_ = params['m'], params['n'], params['p']
        
        # Get A_per_min from logA in raw_params
        logA = self.fit_results['raw_params'][self.fit_results['param_order'].index('logA')]
        A_min = math.exp(logA)
        
        ea_ser = self.ea.set_index("alpha")["Ea_kJ_per_mol"] * 1000

        num_curves = 0
        for i, (beta_k_min, df) in enumerate(self.dfs.items()):
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            # Experimental rate is d(alpha)/dt = dAdT * beta
            dAdt_exp_per_min = df['dAdT'] * beta_k_min
            ax.plot(df['Temp_K'], dAdt_exp_per_min, '--', alpha=0.7, label=f"β={beta_k_min:.1f} (Exp.)")
            ### MODIFICATION END ###
            num_curves += 1

            alpha = np.clip(df["alpha"].values, EPS, 1 - EPS)
            Ea_interp = np.interp(alpha, ea_ser.index, ea_ser.values)
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", RuntimeWarning)
                f_alpha = alpha**m * (1-alpha)**n * (-np.log(1-alpha))**p_
            
            # Model rate is d(alpha)/dt = A_min * exp(-Ea/RT) * f(alpha)
            dAdt_model = A_min * np.exp(-Ea_interp / (R * df['Temp_K'])) * f_alpha
            ax.plot(df['Temp_K'], dAdt_model, '-', label=f"β={beta_k_min:.1f} (Fit)")
            num_curves += 1

        ax.set_xlabel("Temperature (K)")
        ax.set_ylabel(r"Reaction Rate, d$\alpha$/dt (min$^{-1}$)")
        ax.margins(x=0)
        self.add_smart_legend(fig, ax, num_curves)
        fig.tight_layout()
        self._create_plot_window(fig, "Global Fit vs. Experimental Data")
        ### MODIFICATION END ###

    # ##################################################################################
    # ############# START: NEW/MODIFIED CROSS-VALIDATION & FITTING SECTION #############
    # ##################################################################################

    def run_autocatalytic_fit(self):
        """MODIFIED: Main entry point for complex model fitting, now branches to CV."""
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if self.ea is None:
            return messagebox.showwarning("Prerequisite Missing", "Please run an Isoconversional Analysis first.", parent=self)
        
        dlg = AutocatalyticModelDialog(self, self.autocatalytic_model_definitions)
        self.wait_window(dlg)
        if not dlg.result: return

        model_key = dlg.result["model"]
        use_cross_val = dlg.result["cross_validation"]

        prompt = f"Enter α range for {model_key} model fit (e.g., 0.1,0.9):"; initial = f"{self.alphas.min()},{self.alphas.max()}"
        rng = simpledialog.askstring("Fit Settings", prompt, initialvalue=initial, parent=self)
        if not rng: return
        try:
            a_min, a_max = map(float, rng.split(','))
            if not (0 < a_min < a_max < 1): raise ValueError
            fit_alphas = np.round(np.arange(a_min, a_max + EPS, 0.01), 4)
        except (ValueError, IndexError): return messagebox.showerror("Error", "Invalid α range format.")
        
        # --- NEW: Branching logic for Cross-Validation ---
        if use_cross_val:
            self.status_var.set(f"Starting Cross-Validation for {model_key} model...")
            self._start_task(self._run_cross_validation_worker, model_key, fit_alphas, on_success=self._on_cross_validation_success)
        else:
            # --- Standard fitting workflow ---
            self.status_var.set(f"Starting single fit for {model_key} model...")
            if model_key == "SB": self._start_task(self._global_fit_logic, self.dfs, fit_alphas, on_success=self._on_global_fit_success)
            elif model_key == "KS": self._start_task(self._fit_kamal_sourour_logic, self.dfs, fit_alphas, on_success=self._on_autocatalytic_fit_success)
            elif model_key == "GAI": self._start_task(self._fit_gai_logic, self.dfs, fit_alphas, on_success=self._on_autocatalytic_fit_success)
            elif model_key == "PAR": self._start_task(self._fit_par_logic, self.dfs, fit_alphas, on_success=self._on_autocatalytic_fit_success)

    def _run_cross_validation_worker(self, model_key, fit_alphas):
        """NEW: Core worker for Leave-One-Out Cross-Validation."""
        all_betas = sorted(self.dfs.keys())
        if len(all_betas) < 3:
            raise ValueError("Cross-validation requires at least 3 heating rates.")

        model_logic_map = {
            "SB": self._calculate_cka_logic, "KS": self._fit_kamal_sourour_logic,
            "GAI": self._fit_gai_logic, "PAR": self._fit_par_logic
        }
        if model_key not in model_logic_map:
            raise ValueError(f"Unknown model key for cross-validation: {model_key}")
        
        fit_function = model_logic_map[model_key]
        cv_results = []

        for i, beta_to_validate in enumerate(all_betas):
            self.status_var.set(f"CV Iteration {i+1}/{len(all_betas)}: Validating β = {beta_to_validate:.1f} K/min")
            
            # 1. Data Partition
            training_dfs = {b: df for b, df in self.dfs.items() if b != beta_to_validate}
            validation_df = self.dfs[beta_to_validate]

            # 2. Model Fitting on Training Set
            try:
                fit_result = fit_function(training_dfs, fit_alphas)
            except Exception as e:
                cv_results.append({"beta_validated": beta_to_validate, "SSE": np.nan, "params": {}, "error": str(e)})
                continue

            # 3. Prediction on Validation Set
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            dAdt_exp_per_min = validation_df['dAdT'] * beta_to_validate
            ### MODIFICATION END ###
            
            dAdt_pred_per_min = self._predict_rate_from_fit(fit_result, validation_df)
            
            # Ensure arrays are aligned and finite for SSE calculation
            valid_indices = np.isfinite(dAdt_exp_per_min) & np.isfinite(dAdt_pred_per_min)
            dAdt_exp_valid = dAdt_exp_per_min[valid_indices]
            dAdt_pred_valid = dAdt_pred_per_min[valid_indices]

            # 4. Comparison (SSE)
            sse = np.sum((dAdt_pred_valid - dAdt_exp_valid)**2)
            cv_results.append({"beta_validated": beta_to_validate, "SSE": sse, "params": fit_result['params'], "raw_params": fit_result['raw_params'], "error": None})
            
        return {"model_name": model_key, "cv_results": cv_results}

    def _predict_rate_from_fit(self, fit_result, df_to_predict):
        """NEW: Helper to predict reaction rate for a given dataframe and fitted model."""
        model_name = fit_result.get('model_name', '').upper()
        raw_params = fit_result.get('raw_params')
        if raw_params is None: return np.full_like(df_to_predict['Temp_K'].values, np.nan)
        
        T_K = df_to_predict['Temp_K'].values
        alpha = np.clip(df_to_predict["alpha"].values, EPS, 1 - EPS)
        a1 = np.clip(1 - alpha, EPS, None)
        dAdt_model_per_min = np.zeros_like(T_K)

        with np.errstate(all='ignore'):
            if model_name in ["CKA", "SB", "CKA (FIXED EA)"]:
                m, n, p_, logA, Ea_J = raw_params
                A_min = np.exp(logA)
                f_alpha = (alpha**m) * (a1**n) * ((-np.log(a1))**p_)
                dAdt_model_per_min = A_min * np.exp(-Ea_J / (R * T_K)) * f_alpha
            
            ### MODIFICATION START: GlobalFit Prediction ###
            elif model_name == "GLOBALFIT":
                m, n, p_, logA = raw_params
                A_min = np.exp(logA)
                f_alpha = (alpha**m) * (a1**n) * ((-np.log(a1))**p_)
                # Get Ea from iso-conversional
                ea_ser = self.ea.set_index("alpha")["Ea_kJ_per_mol"] * 1000
                Ea_interp = np.interp(alpha, ea_ser.index, ea_ser.values)
                dAdt_model_per_min = A_min * np.exp(-Ea_interp / (R * T_K)) * f_alpha
            ### MODIFICATION END ###
            
            elif model_name == "KAMAL-SOUROUR":
                logA1, E1_J, logA2, E2_J, m, n = raw_params
                k1 = np.exp(logA1) * np.exp(-E1_J / (R * T_K))
                k2 = np.exp(logA2) * np.exp(-E2_J / (R * T_K))
                dAdt_model_per_min = (k1 + k2 * (alpha**m)) * (a1**n)
            elif model_name == "GAI":
                logA, E_J, n1, z0, n2 = raw_params
                k = np.exp(logA) * np.exp(-E_J / (R * T_K))
                f_alpha = (a1**n1) * (z0 + alpha**n2)
                dAdt_model_per_min = k * f_alpha
            elif model_name == "PAR":
                logA1, E1_J, n1, logA2, E2_J, n2 = raw_params
                k1 = np.exp(logA1) * np.exp(-E1_J / (R * T_K))
                k2 = np.exp(logA2) * np.exp(-E2_J / (R * T_K))
                dAdt_model_per_min = k1 * (a1**n1) + k2 * (a1**n2)
                
        return dAdt_model_per_min

    def _on_cross_validation_success(self, data):
        """NEW: Callback to display the CV results window."""
        self.status_var.set(f"✓ Cross-Validation for {data['model_name']} complete.")
        self.show_cv_results_window(data)

    def show_cv_results_window(self, data):
        """NEW: Creates the Toplevel window to display all CV results."""
        model_name = data['model_name']
        win = tk.Toplevel(self); win.title(f"Cross-Validation Results ({model_name})")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico"); win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("1000x800")
        main_frame = ttk.Frame(win, padding=15); main_frame.pack(fill="both", expand=True)

        # --- Summary Table ---
        summary_frame = ttk.LabelFrame(main_frame, text="Validation Summary", padding=10)
        summary_frame.pack(fill="x", expand=False, pady=(0, 15))
        
        btn_frame = ttk.Frame(summary_frame); btn_frame.pack(fill='x', pady=5)
        ttk.Button(btn_frame, text="Plot Validation Curves", command=lambda: self.plot_cv_results(data), style="Secondary.TButton").pack(side="right")
        ttk.Button(btn_frame, text="Export Plot Data", command=lambda: self._export_cv_plot_data(data), style="Secondary.TButton").pack(side="right", padx=5)
        summary_df = pd.DataFrame(data['cv_results'])[['beta_validated', 'SSE', 'error']]
        ttk.Label(btn_frame, text=f"Mean SSE: {summary_df['SSE'].mean():.3e}", font=("Segoe UI", 10, "bold")).pack(side="left", padx=10)

        tv_summary = ttk.Treeview(summary_frame, columns=('Held-out β (K/min)', 'Prediction SSE', 'Status'), show="headings", height=len(summary_df), style="Modern.Treeview")
        for col in tv_summary['columns']: tv_summary.heading(col, text=col); tv_summary.column(col, width=150, anchor='center')
        for _, r in summary_df.iterrows():
            status = "Success" if r['error'] is None else "Fit Failed"
            tv_summary.insert("", "end", values=(f"{r['beta_validated']:.2f}", f"{r['SSE']:.3e}" if pd.notna(r['SSE']) else "N/A", status))
        tv_summary.pack(fill="x", expand=True, pady=5)

        # --- Parameter Stability Table ---
        params_frame = ttk.LabelFrame(main_frame, text="Parameter Stability Across Training Sets", padding=10)
        params_frame.pack(fill="both", expand=True)
        
        # Unpack parameters into a flat DataFrame
        all_params = []
        for r in data['cv_results']:
            p_row = {'Trained without β': f"{r['beta_validated']:.2f}"}
            p_row.update(r['params'])
            all_params.append(p_row)
        params_df = pd.DataFrame(all_params)

        param_cols = list(params_df.columns)
        tv_params = ttk.Treeview(params_frame, columns=param_cols, show="headings", style="Modern.Treeview")
        for c in param_cols: tv_params.heading(c, text=c); tv_params.column(c, width=100, anchor='e')
        for _, r in params_df.iterrows():
            vals = [f"{v:.4g}" if isinstance(v, (float, int)) else v for v in r]
            tv_params.insert("", "end", values=vals)
        
        ysb = ttk.Scrollbar(params_frame, orient=tk.VERTICAL, command=tv_params.yview)
        xsb = ttk.Scrollbar(params_frame, orient=tk.HORIZONTAL, command=tv_params.xview)
        tv_params.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        tv_params.pack(side="left", fill="both", expand=True, pady=5)
        ysb.pack(side="right", fill="y"); xsb.pack(side="bottom", fill="x")

    def plot_cv_results(self, data):
        """NEW: Plots the experimental vs. predicted curves for each validation fold."""
        fig, axes = plt.subplots(len(data['cv_results']), 1, figsize=(7, 3 * len(data['cv_results'])), sharex=True)
        if len(data['cv_results']) == 1: axes = [axes] # Ensure axes is iterable
        fig.suptitle(f"Cross-Validation Predictions for {data['model_name']} Model", fontsize=14)

        for i, result in enumerate(data['cv_results']):
            ax = axes[i]
            beta_val = result['beta_validated']
            validation_df = self.dfs[beta_val]
            
            # Plot experimental data
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            exp_rate = validation_df['dAdT'] * beta_val
            ### MODIFICATION END ###
            
            ax.plot(validation_df['Temp_K'], exp_rate, 'o', ms=4, alpha=0.7, label=f"Experimental")
            
            # Plot predicted data
            if result['error'] is None:
                # Re-create the fit_result structure for the prediction function
                temp_fit_result = {'model_name': data['model_name'], **result}
                pred_rate = self._predict_rate_from_fit(temp_fit_result, validation_df)
                ax.plot(validation_df['Temp_K'], pred_rate, '-', lw=2.0, label=f"Predicted")
            
            ax.set_title(f"Validation on β = {beta_val:.1f} K/min (SSE: {result['SSE']:.2e})", fontsize=10)
            ax.set_ylabel(r"dα/dt (min$^{-1}$)")
            ax.legend()
            ax.margins(x=0)
        axes[-1].set_xlabel("Temperature (K)")
        fig.tight_layout(rect=[0, 0.03, 1, 0.96])
        self._create_plot_window(fig, f"Cross-Validation Plot ({data['model_name']})")
    
    def _export_cv_plot_data(self, data):
        """NEW: Gathers and exports the data used in the cross-validation plot."""
        if not data or not data.get('cv_results'):
            return messagebox.showwarning("No Data", "No cross-validation plot data to export.")

        all_plot_data = []
        model_name = data['model_name']

        for result in data['cv_results']:
            beta_val = result['beta_validated']
            validation_df = self.dfs[beta_val]

            temp_df = validation_df[['Temp_K', 'alpha']].copy()
            temp_df['Heating_Rate_Validated (K/min)'] = beta_val
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            temp_df['Experimental_Rate (min^-1)'] = validation_df['dAdT'] * beta_val
            ### MODIFICATION END ###

            if result['error'] is None and 'raw_params' in result:
                # Re-create the fit_result structure for the prediction function
                temp_fit_result = {'model_name': model_name, **result}
                pred_rate = self._predict_rate_from_fit(temp_fit_result, validation_df)
                temp_df['Predicted_Rate (min^-1)'] = pred_rate
            else:
                temp_df['Predicted_Rate (min^-1)'] = np.nan
            
            all_plot_data.append(temp_df)

        if not all_plot_data:
            return messagebox.showwarning("No Data", "Could not generate any data to export.")

        export_df = pd.concat(all_plot_data, ignore_index=True)
        self._export_table(export_df, f"{model_name}_Cross_Validation_Plot_Data")
        
    # --- Refactored Fitting Logic Functions ---
    def resid_kamal_sourour(self, params, data):
        logA1, E1_J, logA2, E2_J, m, n = params
        y_dadt_per_min, alpha, a1, T_K = data
        logA1, E1_J, logA2, E2_J, m, n = params
        y_dAdt_per_min, alpha, a1, T_K = data
        with np.errstate(all='ignore'):
            k1 = np.exp(logA1) * np.exp(-E1_J / (R * T_K))
            k2 = np.exp(logA2) * np.exp(-E2_J / (R * T_K))
            y_calc = (k1 + k2 * (alpha**m)) * (a1**n)
            return y_calc - y_dadt_per_min
            return y_calc - y_dAdt_per_min
            
    def resid_gai(self, params, data):
        logA, E_J, n1, z0, n2 = params
        y_dadt_per_min, alpha, a1, T_K = data
        y_dAdt_per_min, alpha, a1, T_K = data
        with np.errstate(all='ignore'):
            k = np.exp(logA) * np.exp(-E_J / (R * T_K))
            f_alpha = (a1**n1) * (z0 + alpha**n2)
            y_calc = k * f_alpha
            return y_calc - y_dadt_per_min
            return y_calc - y_dAdt_per_min

    def resid_par(self, params, data):
        logA1, E1_J, n1, logA2, E2_J, n2 = params
        y_dadt_per_min, alpha, a1, T_K = data
        y_dAdt_per_min, alpha, a1, T_K = data
        with np.errstate(all='ignore'):
            k1 = np.exp(logA1) * np.exp(-E1_J / (R * T_K))
            k2 = np.exp(logA2) * np.exp(-E2_J / (R * T_K))
            y_calc = k1 * (a1**n1) + k2 * (a1**n2)
            return y_calc - y_dadt_per_min
            return y_calc - y_dAdt_per_min

    def _fit_gai_logic(self, dfs, alphas):
        data = self.prep_arrays_for_model_fit(dfs, alphas)
        if len(data[0]) < 20: raise ValueError("Too few data points for a reliable GAI fit.")
        Ea_mean_J = self.ea['Ea_kJ_per_mol'].mean() * 1000
        initial_guess = [15, Ea_mean_J, 1.0, 1e-4, 1.0]
        bounds = ([0, 10000, 0, 0, 0], [70, 400000, 5, 1, 5])
        res = least_squares(self.resid_gai, initial_guess, args=(data,), bounds=bounds, loss="soft_l1", x_scale='jac', method='trf')
        if not res.success: raise RuntimeError("GAI optimization did not converge.")
        logA, E_J, n1, z0, n2 = res.x
        try:
            n_samples, n_params = len(data[0]), len(res.x)
            cov = np.linalg.pinv(res.jac.T @ res.jac) * res.cost * 2 / (n_samples - n_params)
            se = np.sqrt(np.diag(cov))
        except: se = np.full(5, np.nan); cov = np.full((5, 5), np.nan)
        return {"model_name": "GAI", "params": {"Ea_kJ_per_mol": E_J / 1000, "A_per_s": np.exp(logA) / 60, "n1": n1, "z0": z0, "n2": n2},
                "errors": {"Ea_kJ_per_mol": se[1] / 1000, "A_per_s": se[0] * np.exp(logA) / 60, "n1": se[2], "z0": se[3], "n2": se[4]},
                "cov": cov, "raw_params": res.x, "param_order": ['logA', 'E_J', 'n1', 'z0', 'n2']}

    def _fit_par_logic(self, dfs, alphas):
        data = self.prep_arrays_for_model_fit(dfs, alphas)
        if len(data[0]) < 20: raise ValueError("Too few data points for a reliable PAR fit.")
        Ea_mean_J = self.ea['Ea_kJ_per_mol'].mean() * 1000
        initial_guess = [15, Ea_mean_J * 0.9, 1.0, 15, Ea_mean_J * 1.1, 1.0]
        bounds = ([0, 10000, 0, 0, 10000, 0], [70, 400000, 5, 70, 400000, 5])
        res = least_squares(self.resid_par, initial_guess, args=(data,), bounds=bounds, loss="soft_l1", x_scale='jac', method='trf')
        if not res.success: raise RuntimeError("PAR optimization did not converge.")
        logA1, E1_J, n1, logA2, E2_J, n2 = res.x
        try:
            n_samples, n_params = len(data[0]), len(res.x)
            cov = np.linalg.pinv(res.jac.T @ res.jac) * res.cost * 2 / (n_samples - n_params)
            se = np.sqrt(np.diag(cov))
        except: se = np.full(6, np.nan); cov = np.full((6, 6), np.nan)
        return {"model_name": "PAR", "params": {"E1_kJ_per_mol": E1_J / 1000, "A1_per_s": np.exp(logA1) / 60, "n1": n1, "E2_kJ_per_mol": E2_J / 1000, "A2_per_s": np.exp(logA2) / 60, "n2": n2},
                "errors": {"E1_kJ_per_mol": se[1] / 1000, "A1_per_s": se[0] * np.exp(logA1) / 60, "n1": se[2], "E2_kJ_per_mol": se[4] / 1000, "A2_per_s": se[3] * np.exp(logA2) / 60, "n2": se[5]},
                "cov": cov, "raw_params": res.x, "param_order": ['logA1', 'E1_J', 'n1', 'logA2', 'E2_J', 'n2']}

    def _fit_kamal_sourour_logic(self, dfs, alphas):
        data = self.prep_arrays_for_model_fit(dfs, alphas)
        if len(data[0]) < 20: raise ValueError("Too few data points for a reliable Kamal-Sourour fit.")
        Ea_mean_J = self.ea['Ea_kJ_per_mol'].mean() * 1000
        initial_guess = [15, Ea_mean_J * 0.8, 15, Ea_mean_J * 1.2, 0.5, 1.5]
        bounds = ([0, 10000, 0, 10000, 0, 0], [70, 400000, 70, 400000, 5, 5])
        res = least_squares(self.resid_kamal_sourour, initial_guess, args=(data,), bounds=bounds, loss="soft_l1", x_scale='jac', method='trf')
        if not res.success: raise RuntimeError("Kamal-Sourour optimization did not converge.")
        logA1, E1_J, logA2, E2_J, m, n = res.x
        try:
            n_samples, n_params = len(data[0]), len(res.x)
            cov = np.linalg.pinv(res.jac.T @ res.jac) * res.cost * 2 / (n_samples - n_params)
            se = np.sqrt(np.diag(cov))
        except: se = np.full(6, np.nan); cov = np.full((6, 6), np.nan)
        return {"model_name": "Kamal-Sourour", "params": {"E1_kJ_per_mol": E1_J / 1000, "A1_per_s": np.exp(logA1) / 60, "E2_kJ_per_mol": E2_J / 1000, "A2_per_s": np.exp(logA2) / 60, "m": m, "n": n},
                "errors": {"E1_kJ_per_mol": se[1] / 1000, "A1_per_s": se[0] * np.exp(logA1) / 60, "E2_kJ_per_mol": se[3] / 1000, "A2_per_s": se[2] * np.exp(logA2) / 60, "m": se[4], "n": se[5]},
                "cov": cov, "raw_params": res.x, "param_order": ['logA1', 'E1_J', 'logA2', 'E2_J', 'm', 'n']}

    def _on_autocatalytic_fit_success(self, results):
        self.autocatalytic_results = results
        self.last_fitted_model = { "model_name": results['model_name'], "params": self.autocatalytic_results['params'],
            "raw_params": self.autocatalytic_results['raw_params'], "cov": self.autocatalytic_results.get('cov', np.nan),
            "param_order": self.autocatalytic_results.get('param_order', []), "source": "Autocatalytic Fit" }
        self.status_var.set(f"✓ {results['model_name']} Fit complete.")
        self.show_autocatalytic_results()

    def show_autocatalytic_results(self):
        # This function remains largely the same, displaying single-fit results
        if not hasattr(self, 'autocatalytic_results') or not self.autocatalytic_results: return
        results, model_name = self.autocatalytic_results, self.autocatalytic_results['model_name']
        win = tk.Toplevel(self, bg=self.BG); win.title(f"{model_name} Fit Results")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico"); win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("800x600")

        frame = ttk.Frame(win, padding="20"); frame.pack(fill="both", expand=True)
        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Label(btn_frame, text=f"{model_name} Fit Results", font=("Segoe UI", 12, "bold")).pack(side="left")

        params_df = pd.DataFrame([results['params'], results['errors']]).T
        params_df.columns = ['Value', 'Std. Error']; params_df.index.name = 'Parameter'

        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(params_df.reset_index(), f"{model_name}_fit_results"), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Predict Conversion Time...", command=self.launch_prediction_dialog, style="Secondary.TButton").pack(side="right", padx=5)
        
        ### MODIFICATION START: Added new buttons ###
        ttk.Button(btn_frame, text="Export TG Fit Table", command=self.export_tg_fit_table, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Model TG vs. Exp", command=self.plot_tg_fit_comparison, style="Secondary.TButton").pack(side="right", padx=5)
        ### MODIFICATION END ###
        
        ttk.Button(btn_frame, text="Export Plot Data", command=self._export_plot_data_autocatalytic, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Fit vs. Experimental", command=self.plot_autocatalytic_fit, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Compare Mechanism Shape...", command=lambda: self.launch_mechanism_comparison(fit_type='autocatalytic'), style="Secondary.TButton").pack(side="right", padx=5)
        
        # Displaying equation and results table... (code unchanged from original)
        p = results['params']; eq_text = f"Fitted {model_name} Model Equation"
        if model_name == "Kamal-Sourour":
            p_disp = {"E₁ (kJ/mol)": "E1_kJ_per_mol", "A₁ (s⁻¹)": "A1_per_s", "E₂ (kJ/mol)": "E2_kJ_per_mol", "A₂ (s⁻¹)": "A2_per_s", "m": "m", "n": "n"}
            p_vals = {k: v for k, v in p.items() if k in p_disp.values()}
            eq_text = (f"$d\\alpha/dt = [A_1 e^{{-E_1/RT}} + A_2 e^{{-E_2/RT}} \\cdot \\alpha^{{m}}] (1-\\alpha)^{{n}}$\n"
                       f"$E_1={p_vals.get('E1_kJ_per_mol', 0):.1f}$, $A_1={p_vals.get('A1_per_s', 0):.2e}$, "
                       f"$E_2={p_vals.get('E2_kJ_per_mol', 0):.1f}$, $A_2={p_vals.get('A2_per_s', 0):.2e}$, "
                       f"$m={p_vals.get('m', 0):.2f}$, $n={p_vals.get('n', 0):.2f}$")
        elif model_name == "GAI":
            eq_text = (f"$d\\alpha/dt = A e^{{-E/RT}} (1-\\alpha)^{{n_1}} (z_0 + \\alpha^{{n_2}})$\n"
                       f"$E_a={p['Ea_kJ_per_mol']:.1f}$, $A={p['A_per_s']:.2e}$, "
                       f"$n_1={p['n1']:.2f}$, $z_0={p['z0']:.2e}$, $n_2={p['n2']:.2f}$")
        elif model_name == "PAR":
            p_disp = {"E₁ (kJ/mol)": "E1_kJ_per_mol", "A₁ (s⁻¹)": "A1_per_s", "n₁": "n1", "E₂ (kJ/mol)": "E2_kJ_per_mol", "A₂ (s⁻¹)": "A2_per_s", "n₂": "n2"}
            p_vals = {k: v for k, v in p.items() if k in p_disp.values()}
            eq_text = (f"$d\\alpha/dt = A_1 e^{{-E_1/RT}}(1-\\alpha)^{{n_1}} + A_2 e^{{-E_2/RT}}(1-\\alpha)^{{n_2}}$\n"
                       f"$E_1={p_vals.get('E1_kJ_per_mol', 0):.1f}$, $A_1={p_vals.get('A1_per_s', 0):.2e}$, $n_1={p_vals.get('n1', 0):.2f}$\n"
                       f"$E_2={p_vals.get('E2_kJ_per_mol', 0):.1f}$, $A_2={p_vals.get('A2_per_s', 0):.2e}$, $n_2={p_vals.get('n2', 0):.2f}$")
        fig_eq = plt.figure(figsize=(8, 3), facecolor=self.BG)
        fig_eq.text(0.5, 0.5, eq_text, va='center', ha='center', fontsize=11, color=self.FG, multialignment='center')
        canvas = FigureCanvasTkAgg(fig_eq, master=frame); canvas.draw()
        canvas.get_tk_widget().pack(fill='x', expand=False, pady=(10, 15)); plt.close(fig_eq)
        tv_cols = ("Parameter", "Value", "Std. Error")
        tv = ttk.Treeview(frame, columns=tv_cols, show="headings", height=len(p)+1, style="Modern.Treeview"); tv.pack(fill="x", expand=True)
        for col_id in tv_cols: tv.heading(col_id, text=col_id); tv.column(col_id, anchor="center")
        for param_name, val in p.items():
            display_name = next((k for k, v in self.display_header_map.items() if v == param_name), param_name)
            err = results["errors"].get(display_name, 0)
            tv.insert("", "end", values=(param_name, f"{val:.4g}", f"{err:.2e}"))

    def plot_autocatalytic_fit(self):
        if not hasattr(self, 'autocatalytic_results') or not self.autocatalytic_results: return
        fig, ax = plt.subplots(figsize=(7, 5.5))
        num_curves = 0
        for beta_k_min, df in self.dfs.items():
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            dAdt_exp_per_min = df['dAdT'] * beta_k_min 
            ### MODIFICATION END ###
            
            p = ax.plot(df['Temp_K'], dAdt_exp_per_min, 'o', ms=4, alpha=0.6, label=f"β={beta_k_min:.1f} (Exp.)")
            exp_color = p[0].get_color()
            num_curves += 1
            dAdt_model_per_min = self._predict_rate_from_fit(self.autocatalytic_results, df)
            ax.plot(df['Temp_K'], dAdt_model_per_min, '-', color=exp_color, lw=2.0, label=f"β={beta_k_min:.1f} (Fit)")
            num_curves += 1
        ax.set_xlabel("Temperature (K)"); ax.set_ylabel(r"Reaction Rate, d$\alpha$/dt (min$^{-1}$)")
        ax.margins(x=0); self.add_smart_legend(fig, ax, num_curves); fig.tight_layout()
        self._create_plot_window(fig, f"{self.autocatalytic_results['model_name']} Fit Plot")
        
    def run_cka_analysis(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if not self.dfs: return messagebox.showwarning("No Data", "Please import data before running a CKA analysis.")
        prompt = "Enter α range for CKA (e.g., 0.1,0.9):"; initial = f"{self.alphas.min()},{self.alphas.max()}"
        rng = simpledialog.askstring("CKA Settings", prompt, initialvalue=initial, parent=self)
        if not rng: return
        try:
            a_min, a_max = map(float, rng.split(','))
            if not (0 < a_min < a_max < 1): raise ValueError
            fit_alphas = np.round(np.arange(a_min, a_max + EPS, 0.01), 4)
        except (ValueError, IndexError): return messagebox.showerror("Error", "Invalid α range format.")
        self._start_task(self._calculate_cka_logic, self.dfs, fit_alphas, on_success=self._on_cka_success)

    def prep_arrays_for_model_fit(self, dfs, alphas):
        y_dadt_per_min, a, a1, T = [], [], [], []
        y_dAdt_per_min, a, a1, T = [], [], [], []
        for β, df in dfs.items():
            alpha_vals = np.clip(df["alpha"].values, EPS, 1 - EPS)
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            mask = (alpha_vals >= alphas.min()) & (alpha_vals <= alphas.max()) & np.isfinite(df["dAdT"].values)
            ### MODIFICATION END ###
            
            alpha_masked = alpha_vals[mask]
            if len(alpha_masked) == 0: continue
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            y_dAdt_per_min.extend(df["dAdT"].values[mask] * β) # Convert to rate per minute
            ### MODIFICATION END ###
            
            a.extend(alpha_masked); a1.extend(np.clip(1 - alpha_masked, EPS, None)); T.extend(df["Temp_K"].values[mask])
        arrs = [np.asarray(v, float) for v in (y_dadt_per_min, a, a1, T)]
        arrs = [np.asarray(v, float) for v in (y_dAdt_per_min, a, a1, T)]
        good = np.all(np.isfinite(arrs), axis=0)
        return tuple(ar[good] for ar in arrs)

    def resid_cka(self, par, data):
        # --- MODIFICATION: Ea is now fixed and passed in data ---
        m, n, p_, logA = par
        y_dAdt_per_min, a, a1, T, Ea_J = data # Ea_J is now fixed
        # --- END MODIFICATION ---
        
        with np.errstate(all='ignore'):
            f_alpha = (a**m) * (a1**n) * ((-np.log(a1))**p_)
        A_per_min = math.exp(logA)
        y_calc = A_per_min * np.exp(-Ea_J / (R * T)) * f_alpha
        return y_calc - y_dAdt_per_min

    def _calculate_cka_logic(self, dfs, alphas):
        """
        Refactored CKA fitting logic.
        --- MODIFICATION: Implements fixed Ea from Iso-Mean to avoid KCE ---
        """
        
        # --- 1. Get Ea data from isoconversional method ---
        if self.ea is None or self.ea.empty:
            raise ValueError("Isoconversional Ea results (self.ea) are missing or empty.")
            
        # --- 2. Filter Ea values based on the fitting alpha range ---
        ea_in_range = self.ea[
            (self.ea['alpha'] >= alphas.min()) & 
            (self.ea['alpha'] <= alphas.max())
        ]['Ea_kJ_per_mol'].dropna()
        
        if ea_in_range.empty:
            raise ValueError(f"No valid Ea values found in the selected alpha range [{alphas.min()}, {alphas.max()}]")

        # --- 3. Apply IQR filter to remove outliers ---
        Q1 = ea_in_range.quantile(0.25)
        Q3 = ea_in_range.quantile(0.75)
        IQR_val = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR_val
        upper_bound = Q3 + 1.5 * IQR_val
        
        ea_filtered = ea_in_range[(ea_in_range >= lower_bound) & (ea_in_range <= upper_bound)]
        
        if ea_filtered.empty:
            self.status_var.set("Warning: IQR filter removed all Ea values. Using unfiltered mean.")
            ea_filtered = ea_in_range # Fallback to unfiltered data
        
        # --- 4. Calculate the fixed Ea value in J/mol ---
        Ea_fixed_J = ea_filtered.mean() * 1000.0
        Ea_std_dev_kJ = ea_filtered.std() # For reference, not used in fit
        
        # --- 5. Prepare data arrays, now passing Ea_fixed_J ---
        y_dAdt_per_min, a, a1, T = [], [], [], []
        for β, df in dfs.items():
            alpha_vals = np.clip(df["alpha"].values, EPS, 1 - EPS)
            mask = (alpha_vals >= alphas.min()) & (alpha_vals <= alphas.max()) & np.isfinite(df["dAdT"].values)
            alpha_masked = alpha_vals[mask]
            if len(alpha_masked) == 0: continue
            
            y_dAdt_per_min.extend(df["dAdT"].values[mask] * β) 
            a.extend(alpha_masked); a1.extend(np.clip(1 - alpha_masked, EPS, None)); T.extend(df["Temp_K"].values[mask])
        
        arrs = [np.asarray(v, float) for v in (y_dAdt_per_min, a, a1, T)]
        good = np.all(np.isfinite(arrs), axis=0)
        
        # Add the fixed Ea to the data tuple
        data = tuple(ar[good] for ar in arrs) + (np.full(np.sum(good), Ea_fixed_J),)
        
        if len(data[0]) < 10: raise ValueError("Too few data points for a reliable CKA fit.")
        
        # --- 6. Perform 4-parameter optimization ---
        initial_guess = [1, 1, 0, 30] # m, n, p, logA
        bounds = ([0, 0, 0, -30], [5, 5, 5, 70]) # Bounds for m, n, p, logA
        
        res = least_squares(self.resid_cka, initial_guess, args=(data,), bounds=bounds, loss="soft_l1", x_scale='jac', method='trf')
        if not res.success: raise RuntimeError("The CKA (Fixed Ea) optimization did not converge.")
        
        m, n, p_, logA = res.x
        
        # --- 7. Calculate Std. Errors for the 4 fitted parameters ---
        try:
            n_samples, n_params = len(data[0]), len(res.x) # n_params is 4
            cov = np.linalg.pinv(res.jac.T @ res.jac) * res.cost * 2 / (n_samples - n_params)
            se = np.sqrt(np.diag(cov)) # 4-element array
        except: 
            se = np.full(4, np.nan); cov = np.full((4, 4), np.nan)
        
        # --- 8. Re-assemble the full 5-parameter result set ---
        return {
            "model_name": "CKA (Fixed Ea)", 
            "params": {
                "Ea_kJ_per_mol": Ea_fixed_J / 1000, 
                "A_per_s": math.exp(logA) / 60, 
                "m": m, 
                "n": n, 
                "p": p_
            },
            "errors": {
                "Ea_kJ_per_mol": Ea_std_dev_kJ, # Store the stdev of the iso-data
                "A_per_s": se[3] * math.exp(logA) / 60, 
                "m": se[0], 
                "n": se[1], 
                "p": se[2]
            },
            "cov": cov, # Note: This is the 4x4 covariance matrix
            "raw_params": [m, n, p_, logA, Ea_fixed_J], # Store all 5 raw values
            "param_order": ['m', 'n', 'p', 'logA', 'Ea_J']
        }
        
    def _on_cka_success(self, results):
        self.cka_results = results
        self.last_fitted_model = { "model_name": self.cka_results['model_name'], "params": self.cka_results['params'],
            "raw_params": self.cka_results['raw_params'], "cov": self.cka_results['cov'],
            "param_order": self.cka_results['param_order'], "source": "Combined Kinetic Analysis" }
        self.status_var.set("✓ CKA complete. Showing results.")
        self.show_cka_results()
    def show_cka_results(self):
        if not self.cka_results: return
        
        # --- MODIFICATION: Title updated to reflect fixed Ea ---
        win_title = "CKA Results (Fixed Ea from Iso-Mean)"
        if self.cka_results.get("model_name"):
             win_title = f"{self.cka_results['model_name']} Fit Results"
        # --- END MODIFICATION ---

        win = tk.Toplevel(self, bg=self.BG); win.title(win_title)
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.co")
            win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("800x550")

        frame = ttk.Frame(win, padding="20"); frame.pack(fill="both", expand=True)
        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Label(btn_frame, text=win_title, font=("Segoe UI", 12, "bold")).pack(side="left")
        
        cka_params_df = pd.DataFrame([self.cka_results['params'], self.cka_results['errors']]).T
        cka_params_df.columns = ['Value', 'Std. Error']
        cka_params_df.index.name = 'Parameter'
        cka_params_df = cka_params_df.reset_index()

        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(cka_params_df, "cka_results_fixed_ea"), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Predict Conversion Time...", command=self.launch_prediction_dialog, style="Secondary.TButton").pack(side="right", padx=5)
        
        ttk.Button(btn_frame, text="Export TG Fit Table", command=self.export_tg_fit_table, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Model TG vs. Exp", command=self.plot_tg_fit_comparison, style="Secondary.TButton").pack(side="right", padx=5)
        
        ttk.Button(btn_frame, text="Export Plot Data", command=self._export_plot_data_cka, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Fit vs. Experimental", command=self.plot_cka_fit, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Compare Mechanism Shape...", command=lambda: self.launch_mechanism_comparison(fit_type='cka'), style="Secondary.TButton").pack(side="right", padx=5)
        p = self.cka_results['params']
        eq_text = f"$f(\\alpha) = \\alpha^{{{p['m']:.2f}}} (1-\\alpha)^{{{p['n']:.2f}}} [-\\ln(1-\\alpha)]^{{{p['p']:.2f}}}$"
        fig_eq = plt.figure(figsize=(6, 1.5), facecolor=self.BG)
        fig_eq.text(0.5, 0.5, eq_text, va='center', ha='center', fontsize=12, color=self.FG)
        canvas = FigureCanvasTkAgg(fig_eq, master=frame); canvas.draw()
        canvas.get_tk_widget().pack(fill='x', expand=False, pady=(10,15)); plt.close(fig_eq)

        tv_cols = ("Parameter", "Value", "Std. Error")
        tv = ttk.Treeview(frame, columns=tv_cols, show="headings", height=6, style="Modern.Treeview"); tv.pack(fill="x", expand=True)
        for col_id in tv_cols: tv.heading(col_id, text=col_id); tv.column(col_id, anchor="center")
        
        param_order = ["Ea_kJ_per_mol", "A_per_s", "m", "n", "p"]
        for name in param_order:
            val, err = self.cka_results["params"].get(name, 0), self.cka_results["errors"].get(name, 0)
            display_name = self.display_header_map.get(name, name)
            
            # --- MODIFICATION: Add (Fixed) label as requested ---
            if name == "Ea_kJ_per_mol":
                tv.insert("", "end", values=(display_name, f"{val:.6g}", "(Fixed from Iso-Mean)"))
            else:
                tv.insert("", "end", values=(display_name, f"{val:.6g}", f"{err:.2e}"))
            # --- END MODIFICATION ---

    def plot_cka_fit(self):
        if not self.cka_results: return
        fig, ax = plt.subplots(figsize=(6, 5))
        
        num_curves = 0
        for i, (beta_k_min, df) in enumerate(self.dfs.items()):
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            dAdt_exp_per_min = df['dAdT'] * beta_k_min 
            ### MODIFICATION END ###
            
            ax.plot(df['Temp_K'], dAdt_exp_per_min, '--', alpha=0.7, label=f"β={beta_k_min:.1f} (Exp.)")
            num_curves += 1

            dAdt_model_per_min = self._predict_rate_from_fit(self.cka_results, df)
            ax.plot(df['Temp_K'], dAdt_model_per_min, '-', label=f"β={beta_k_min:.1f} (CKA Fit)")
            num_curves += 1

        ax.set_xlabel("Temperature (K)"); ax.set_ylabel(r"Reaction Rate, d$\alpha$/dt (min$^{-1}$)")
        ax.margins(x=0)
        self.add_smart_legend(fig, ax, num_curves)
        fig.tight_layout()
        self._create_plot_window(fig, "CKA Fit Plot")

    def launch_mechanism_comparison(self, fit_type):
        if fit_type == 'cka':
            fit_results = self.cka_results
        elif fit_type == 'global_fit':
            fit_results = self.fit_results
        elif fit_type == 'autocatalytic':
            ac_results = self.autocatalytic_results
            if not ac_results:
                 return messagebox.showwarning("No Data", "Please run an Autocatalytic Fit analysis first.")
            
            alphas_test = np.linspace(0.01, 0.99, 200)
            rates = self._get_autocatalytic_rate(alphas_test, 300, ac_results) 
            if rates is None or not np.any(rates):
                alpha_max_rate = 0.5
            else:
                alpha_max_rate = alphas_test[np.nanargmax(rates)]

            m_eff = alpha_max_rate
            n_eff = 1 - alpha_max_rate
            
            fit_results = {
                'params': {'m': m_eff, 'n': n_eff, 'p': 0}
            }
        else:
            fit_results = None

        if not fit_results:
            return messagebox.showwarning("No Data", f"Please run a relevant analysis first.")
        
        dlg = ModelSelectionDialog(self, IDEAL_MODELS)
        self.wait_window(dlg)
        
        if dlg.result:
            selected_models = dlg.result
            self.show_mechanism_comparison_results(selected_models, fit_results, fit_type.upper())

    def show_mechanism_comparison_results(self, selected_models, fit_results, fit_name):
        win = tk.Toplevel(self)
        win.title(f"Mechanism Comparison ({fit_name})")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("900x750")

        main_frame = ttk.Frame(win, padding=15)
        main_frame.pack(fill="both", expand=True)

        alphas_table = np.linspace(0.05, 0.95, 19)
        comparison_df = pd.DataFrame({'alpha': alphas_table})
        params = fit_results['params']
        m, n, p_ = params.get('m', 0), params.get('n', 0), params.get('p', 0)
        
        exp_col_name = f'Experimental ({fit_name})'

        with np.errstate(divide='ignore', invalid='ignore'):
            f_exp = alphas_table**m * (1 - alphas_table)**n * (-np.log(1 - alphas_table))**p_
            f_exp_05 = 0.5**m * (1 - 0.5)**n * (-np.log(1 - 0.5))**p_
            y_exp_norm = f_exp / f_exp_05 if np.abs(f_exp_05) > EPS else np.full_like(f_exp, np.nan)
            comparison_df[exp_col_name] = y_exp_norm

            ranking_results = []
            for model_code in selected_models:
                model_func = IDEAL_MODELS[model_code][1]
                f_ideal = model_func(alphas_table)
                f_ideal_05 = model_func(np.array(0.5))
                y_model_norm = f_ideal / f_ideal_05 if np.abs(f_ideal_05) > EPS else np.full_like(f_ideal, np.nan)
                comparison_df[model_code] = y_model_norm

                temp_df = pd.DataFrame({'exp': y_exp_norm, 'model': y_model_norm}).dropna()
                if not temp_df.empty:
                    y_true, y_pred = temp_df['exp'], temp_df['model']
                    ss_res, ss_tot = np.sum((y_true - y_pred) ** 2), np.sum((y_true - np.mean(y_true)) ** 2)
                    r2 = 1 - (ss_res / ss_tot) if ss_tot > 0 else 0.0
                else:
                    r2 = np.nan
                ranking_results.append({'Model Code': model_code, 'Model Name': IDEAL_MODELS[model_code][0], 'R-squared': r2})
            
            ranking_df = pd.DataFrame(ranking_results).sort_values(by='R-squared', ascending=False).reset_index(drop=True)

        ranking_frame = ttk.LabelFrame(main_frame, text="Goodness-of-Fit Ranking", padding="10", style="Modern.TLabelframe")
        ranking_frame.pack(fill="x", expand=False, pady=(0, 15))
        
        ranking_btn_frame = ttk.Frame(ranking_frame)
        ranking_btn_frame.pack(fill='x', pady=5)
        ttk.Button(ranking_btn_frame, text="Export Ranking", command=lambda: self._export_table(ranking_df, "mechanism_ranking_data"), style="Secondary.TButton").pack(side="right")
        ttk.Button(ranking_btn_frame, text="Plot Comparison", command=lambda: self.plot_mechanism_comparison_graph(comparison_df, exp_col_name), style="Secondary.TButton").pack(side="right", padx=5)

        ranking_cols = list(ranking_df.columns)
        tv_rank = ttk.Treeview(ranking_frame, columns=ranking_cols, show="headings", height=5, style="Modern.Treeview")
        tv_rank.column("Model Code", width=100, anchor="w"); tv_rank.column("Model Name", width=400, anchor="w"); tv_rank.column("R-squared", width=120, anchor="e")
        for c in ranking_cols: tv_rank.heading(c, text=c)
        for _, r in ranking_df.iterrows():
            tv_rank.insert("", "end", values=[r['Model Code'], r['Model Name'], f"{r['R-squared']:.4f}" if pd.notna(r['R-squared']) else "N/A"])
        tv_rank.pack(fill="x", expand=True, pady=(5, 0))

        data_frame = ttk.LabelFrame(main_frame, text="Normalized f(α)/f(0.5) Data", padding="10", style="Modern.TLabelframe")
        data_frame.pack(fill="both", expand=True)

        data_btn_frame = ttk.Frame(data_frame)
        data_btn_frame.pack(fill='x', pady=5)
        ttk.Button(data_btn_frame, text="Export Raw Data", command=lambda: self._export_table(comparison_df, "mechanism_comparison_data"), style="Secondary.TButton").pack(side="right")

        data_cols = list(comparison_df.columns)
        tv_data = ttk.Treeview(data_frame, columns=data_cols, show="headings", style="Modern.Treeview")
        for c in data_cols: tv_data.heading(c, text=c); tv_data.column(c, width=100, anchor="e")
        for _, r in comparison_df.iterrows():
            tv_data.insert("", "end", values=[f"{v:.4f}" if isinstance(v, (int, float)) else v for v in r.fillna('N/A')])
        
        ysb = ttk.Scrollbar(data_frame, orient=tk.VERTICAL, command=tv_data.yview)
        xsb = ttk.Scrollbar(data_frame, orient=tk.HORIZONTAL, command=tv_data.xview)
        tv_data.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        tv_data.pack(side="left", fill="both", expand=True, pady=(5,0))
        ysb.pack(side="right", fill="y"); xsb.pack(side="bottom", fill="x")

    def plot_mechanism_comparison_graph(self, comparison_df, exp_col_name):
        fig, ax = plt.subplots(figsize=(6, 5))
        num_items = 0
        for column in comparison_df.columns:
            if column == 'alpha': continue
            num_items += 1
            if 'Experimental' in column:
                ax.plot(comparison_df['alpha'], comparison_df[column], 'p', color='black', markersize=6, label=column, zorder=10)
            else:
                alphas_dense = np.linspace(0.01, 0.99, 200)
                model_func = IDEAL_MODELS[column][1]
                with np.errstate(divide='ignore', invalid='ignore'):
                    f_ideal = model_func(alphas_dense)
                    f_ideal_05 = model_func(np.array(0.5))
                    if np.abs(f_ideal_05) > EPS:
                        ax.plot(alphas_dense, f_ideal / f_ideal_05, label=column)

        ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel(r"$f(\alpha) / f(0.5)$")
        ax.set_xlim(0, 1.0); ax.set_ylim(0, 2.1)
        self.add_smart_legend(fig, ax, num_items)
        fig.tight_layout()
        self._create_plot_window(fig, "Mechanism Shape Comparison Plot")
    
    def launch_prediction_dialog(self):
        if not self.last_fitted_model:
            return messagebox.showwarning("No Model", "A kinetic model must be fitted first.\nPlease use a function from the 'Analysis' menu (e.g., CKA).")
        
        dlg = ConversionTimeSettingsDialog(self)
        self.wait_window(dlg)
        if not dlg.result: return

        self._start_task(self._predict_conversion_time_worker, dlg.result, on_success=self._on_prediction_success)

    def _predict_conversion_time_worker(self, settings):
        T0_K = settings["T0"]
        model_info = self.last_fitted_model
        
        ### MODIFICATION START ###
        model_name = model_info.get("model_name", "")
        
        # Check if the model is Non-Parametric (NPA) or a Fixed Ea CKA.
        # - NPA: No covariance matrix.
        # - CKA (Fixed Ea): A 4x4 covariance matrix that is incompatible with the 5-param error logic.
        # In both cases, we must skip the error propagation.
        if model_name == "NPA" or model_name == "CKA (Fixed Ea)":
            self.status_var.set(f"Running Prediction for {model_name} (no CI)...")
            conversion_time_df = self.predict_cumulative_conversion_time(model_info, T0_K)
        ### MODIFICATION END ###
        else:
            # Original path for parametric models (GAI, original CKA, etc.)
            self.status_var.set(f"Running Prediction for {model_name} (with CI)...")
            conversion_time_df = self.predict_cumulative_conversion_time_with_error(model_info, T0_K)
        
        if conversion_time_df is None or conversion_time_df.empty:
            raise ValueError("Conversion time prediction failed. Check the fitted model and input temperature.")
            
        return {"conversion_time_df": conversion_time_df, "T0": T0_K, "model_source": model_info.get("source", "Unknown Model")}

    def _on_prediction_success(self, results):
        self.conversion_time_df = results['conversion_time_df']
        T0 = results['T0']
        model_source = results['model_source']
        
        win = tk.Toplevel(self); win.configure(bg=self.BG); win.title(f"Conversion Time Prediction @ {T0} K (from {model_source})")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("900x600")

        frame = ttk.Frame(win, padding=10); frame.pack(fill="both", expand=True)
        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(self.conversion_time_df, f"conversion_time_at_{T0}K"), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Pop-out Plot", command=self.plot_conversion_time, style="Secondary.TButton").pack(side="right", padx=5)

        cols = ("α", "tₐ (days)", "95% CI Lower", "95% CI Upper"); tv = ttk.Treeview(frame, columns=cols, show="headings", style="Modern.Treeview")
        for c in cols: tv.heading(c, text=c)
        tv.column("α", anchor="center", width=120); tv.column("tₐ (days)", anchor="e", width=180)
        tv.column("95% CI Lower", anchor="e", width=180); tv.column("95% CI Upper", anchor="e", width=180)

        for _, r in self.conversion_time_df.iterrows(): 
            tv.insert("", "end", values=(f"{r.alpha:.4g}", f"{r.t_alpha_days:.6g}", f"{r.ci_lower_days:.6g}", f"{r.ci_upper_days:.6g}"))
        tv.pack(fill="both", expand=True, pady=(5,0))

    def plot_conversion_time(self):
        fig, ax = plt.subplots(figsize=(6, 5));
        
        days_to_years = 365.25
        predicted_years = self.conversion_time_df["t_alpha_days"] / days_to_years
        lower_ci_years = self.conversion_time_df["ci_lower_days"] / days_to_years
        upper_ci_years = self.conversion_time_df["ci_upper_days"] / days_to_years

        ax.plot(self.conversion_time_df["alpha"], predicted_years, "o-", label="Predicted Time")
        ax.fill_between(self.conversion_time_df["alpha"], lower_ci_years, upper_ci_years, color='dodgerblue', alpha=0.2, label="95% Confidence Interval")
        
        ax.set_xlabel(r"Conversion, $\alpha$")
        ax.set_ylabel("Conversion Time, $t_\\mathrm{a}$ (years)") 
        
        self.add_smart_legend(fig, ax, 2)
        fig.tight_layout(pad=0.5)
        ax.margins(x=0)
        self._create_plot_window(fig, "Conversion Time Plot with Confidence Interval")

    def _get_autocatalytic_rate(self, alpha, T0_K, model_info):
        model_name = model_info['model_name']
        raw_params = model_info['raw_params']
        
        alpha_clipped = np.clip(alpha, EPS, 1 - EPS)
        a1 = 1 - alpha_clipped
        
        with np.errstate(all='ignore'):
            if model_name == "Kamal-Sourour":
                logA1, E1_J, logA2, E2_J, m, n = raw_params
                k1, k2 = np.exp(logA1) * np.exp(-E1_J / (R * T0_K)), np.exp(logA2) * np.exp(-E2_J / (R * T0_K))
                rate_per_min = (k1 + k2 * (alpha_clipped**m)) * (a1**n)
                return rate_per_min / 60.0
            
            elif model_name == "GAI":
                logA, E_J, n1, z0, n2 = raw_params
                k = np.exp(logA) * np.exp(-E_J / (R * T0_K)); f_alpha = (a1**n1) * (z0 + alpha_clipped**n2)
                return (k * f_alpha) / 60.0

            elif model_name == "PAR":
                logA1, E1_J, n1, logA2, E2_J, n2 = raw_params
                k1, k2 = np.exp(logA1) * np.exp(-E1_J / (R * T0_K)), np.exp(logA2) * np.exp(-E2_J / (R * T0_K))
                return (k1 * (a1**n1) + k2 * (a1**n2)) / 60.0
        return None

    def predict_cumulative_conversion_time_with_error(self, model_info, T0_K):
        raw_params = model_info.get('raw_params')
        cov_matrix = model_info.get('cov')
        param_order = model_info.get('param_order')

        if raw_params is None or cov_matrix is None or param_order is None or np.isnan(cov_matrix).any():
             messagebox.showwarning("Error Propagation Disabled", 
                                  "Covariance matrix is missing or invalid for this model fit. "
                                  "Prediction will proceed without confidence intervals.", parent=self)
             return self.predict_cumulative_conversion_time(model_info, T0_K)

        def get_rate_and_derivatives(alpha, params_vec):
            rate = 0; grad = np.zeros_like(params_vec)
            p_map = {name: val for name, val in zip(param_order, params_vec)}
            
            alpha = np.clip(alpha, EPS, 1 - EPS); a1 = 1 - alpha

            if 'p' in p_map and ('Ea_J' in p_map or 'logA' in p_map):
                m, n, p_ = p_map.get('m',0), p_map.get('n',0), p_map.get('p',0)
                logA = p_map.get('logA',0)
                A = np.exp(logA)
                
                if model_info.get('Ea_source') == 'isoconversional':
                    ea_ser = self.ea.set_index("alpha")["Ea_kJ_per_mol"] * 1000
                    Ea_J = np.interp(alpha, ea_ser.index, ea_ser.values)
                else:
                    Ea_J = p_map.get('Ea_J',0)

                exp_term = np.exp(-Ea_J / (R * T0_K))
                f_alpha = (alpha**m) * (a1**n) * ((-np.log(a1))**p_)
                rate = (A/60) * exp_term * f_alpha # rate in 1/sec
                if abs(rate) > 1e-20:
                    if 'm' in param_order: grad[param_order.index('m')] = rate * np.log(alpha)
                    if 'n' in param_order: grad[param_order.index('n')] = rate * np.log(a1)
                    if 'p' in param_order: grad[param_order.index('p')] = rate * np.log(-np.log(a1))
                    if 'logA' in param_order: grad[param_order.index('logA')] = rate
                    if 'Ea_J' in param_order: grad[param_order.index('Ea_J')] = rate * (-1 / (R * T0_K))
            else:
                 rate = self._get_autocatalytic_rate(alpha, T0_K, model_info) or 0
            return rate, grad

        def integrand_builder(param_idx_to_diff):
            def integrand(alpha):
                rate, grad = get_rate_and_derivatives(alpha, raw_params)
                if abs(rate) < 1e-30: return 0.0
                partial_deriv_rate = grad[param_idx_to_diff]
                return -1.0 / (rate**2) * partial_deriv_rate
            return integrand

        alpha_points = np.linspace(0.01, 0.99, 50)
        time_points_sec, time_variance = [], []
        
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", integrate.IntegrationWarning)
            for alpha_target in alpha_points:
                integrand_t = lambda a: 1.0 / (get_rate_and_derivatives(a, raw_params)[0] + 1e-30)
                time_sec, _ = quad(integrand_t, EPS, alpha_target)
                time_points_sec.append(time_sec)
                grad_t_alpha = np.zeros_like(raw_params)
                for i in range(len(raw_params)):
                    integrand_func = integrand_builder(i)
                    partial_t_alpha, _ = quad(integrand_func, EPS, alpha_target)
                    grad_t_alpha[i] = partial_t_alpha
                J = grad_t_alpha
                var = J.T @ cov_matrix @ J
                time_variance.append(var)

        df = pd.DataFrame({"alpha": alpha_points, "t_alpha_sec": np.array(time_points_sec), "t_alpha_var": np.array(time_variance)})
        df['t_alpha_std_err'] = np.sqrt(df['t_alpha_var'])
        df['t_alpha_days'] = df['t_alpha_sec'] / 86400.0
        df['ci_lower_days'] = (df['t_alpha_sec'] - 1.96 * df['t_alpha_std_err']) / 86400.0
        df['ci_upper_days'] = (df['t_alpha_sec'] + 1.96 * df['t_alpha_std_err']) / 86400.0
        df['ci_lower_days'] = df['ci_lower_days'].clip(lower=0)
        return df

    def predict_cumulative_conversion_time(self, model_info, T0_K):
        params, model_name = model_info.get('params', {}), model_info.get('model_name')
        
        ### MODIFICATION START: Create smooth interpolators to avoid integration errors ###
        npa_Ea_func = None
        npa_Z_func = None
        
        if model_name == "NPA":
            try:
                # 1. Create a smooth, shape-preserving spline function for Ea(α)
                Ea_curve = model_info.get('Ea_curve')
                if Ea_curve is None or Ea_curve.empty: raise ValueError("NPA Ea_curve is missing")
                Ea_J_data = Ea_curve.values * 1000 # Convert kJ/mol to J/mol
                Ea_alpha_data = Ea_curve.index.values
                # Use PchipInterpolator for a stable, monotonic spline
                npa_Ea_func = PchipInterpolator(Ea_alpha_data, Ea_J_data, extrapolate=True)

                # 2. Create a smooth, shape-preserving spline function for log(Z(α))
                # We use log(Z) because Z itself can span many orders of magnitude
                Z_curve = model_info.get('Z_curve')
                if Z_curve is None or Z_curve.empty: raise ValueError("NPA Z_curve is missing")
                Z_alpha_data = Z_curve.index.values
                # Filter out non-positive values before taking log
                valid_Z_mask = Z_curve.values > 0
                log_Z_data = np.log(Z_curve.values[valid_Z_mask])
                log_Z_alpha_data = Z_alpha_data[valid_Z_mask]
                if len(log_Z_alpha_data) < 2: raise ValueError("Not enough valid Z(α) points for spline.")
                npa_Z_func = PchipInterpolator(log_Z_alpha_data, log_Z_data, extrapolate=True)
            except Exception as e:
                messagebox.showerror("NPA Interpolation Error", f"Failed to build spline functions for prediction. Error: {e}", parent=self)
                return pd.DataFrame() # Return empty dataframe
        ### MODIFICATION END ###

        def get_rate_at_alpha(alpha):
            alpha_clipped = np.clip(alpha, EPS, 1 - EPS)
            
            # --- SYNTAX ERROR FIX: Restructured if/elif logic ---
            # We must define 'rate' and only return at the end
            
            rate_per_sec = 1e-30 # Default safety value
            
            if model_name == "NPA":
                if npa_Ea_func is None or npa_Z_func is None:
                    rate_per_sec = 1e-30 # Safety check if interpolation failed
                else:
                    # Use the smooth spline functions instead of np.interp
                    Ea_J = npa_Ea_func(alpha_clipped)
                    log_Z = npa_Z_func(alpha_clipped)
                    
                    # Calculate rate. np.exp(log_Z) is Z_alpha
                    rate_per_sec = np.exp(log_Z) * np.exp(-Ea_J / (R * T0_K)) # Returns rate in 1/sec
            
            # This elif now correctly follows the if block
            elif model_name in ["CKA", "CKA (Fixed Ea)", "GlobalFit"]:
                if model_info.get('Ea_source') == 'isoconversional':
                    ea_ser = self.ea.set_index("alpha")["Ea_kJ_per_mol"] * 1000
                    Ea_J = np.interp(alpha_clipped, ea_ser.index, ea_ser.values)
                else:
                    Ea_J = params.get('Ea_kJ_per_mol', self.ea['Ea_kJ_per_mol'].mean()) * 1000
                
                A_s = params.get('A_per_s', math.exp(params.get('A', 30)) / 60)
                m, n, p_ = params['m'], params['n'], params['p']
                k = A_s * math.exp(-Ea_J / (R * T0_K))
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore", RuntimeWarning)
                    log_term = (-np.log(1 - alpha_clipped))**p_ if p_ > EPS else 1.0
                    f_alpha = (alpha_clipped**m) * (1-alpha_clipped)**n * log_term
                rate_per_sec = k * f_alpha # Returns rate in 1/sec
            else:
                rate = self._get_autocatalytic_rate(alpha, T0_K, model_info) # This helper already returns 1/sec
                rate_per_sec = rate if rate is not None else 1e-30
            
            return rate_per_sec
            # --- END SYNTAX ERROR FIX ---

        def integrand(alpha): return 1.0 / (get_rate_at_alpha(alpha) + 1e-30)
        
        alpha_points = np.linspace(0.01, 0.99, 99)
        # --- MODIFICATION: Increase integration limit to handle "spiky" functions ---
        time_points_sec = []
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always", integrate.IntegrationWarning)
            for a in alpha_points:
                try:
                    # We also add epsabs/epsrel to handle tiny numbers
                    time_val, _ = quad(integrand, EPS, a, limit=200, epsabs=1e-30, epsrel=1e-6)
                    time_points_sec.append(time_val)
                except Exception:
                    time_points_sec.append(np.nan) # Append NaN if quad fails catastrophically
            
            if w:
                self.status_var.set("Integration complete with warnings. Results may be imprecise.")
        
        df = pd.DataFrame({ "alpha": alpha_points, "t_alpha_days": np.array(time_points_sec) / 86400.0 })
        df['ci_lower_days'] = np.nan; df['ci_upper_days'] = np.nan
        return df
    
    def _add_plot_to_doc(self, doc, plot_type):
        fig, ax = plt.subplots(figsize=(12/2.54, 9/2.54))
        if plot_type == 'Ea':
            ax.errorbar(self.ea["alpha"], self.ea["Ea_kJ_per_mol"], yerr=self.ea["StdErr_kJ"], fmt='-o', capsize=3)
            ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel(r"Activation Energy, $E_\mathrm{a}$ (kJ mol$^{-1}$)")
        elif plot_type == 'ConversionTime':
            ax.plot(self.conversion_time_df["alpha"], self.conversion_time_df["t_alpha_days"], "o-", label="Prediction")
            ax.fill_between(self.conversion_time_df["alpha"], self.conversion_time_df["ci_lower_days"], self.conversion_time_df["ci_upper_days"], alpha=0.2, label="95% CI")
            ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel("Conversion Time, $t_\\mathrm{a}$ (days)")
            min_val, max_val = self.conversion_time_df["t_alpha_days"].min(), self.conversion_time_df["t_alpha_days"].max()
            if min_val > 0 and max_val / min_val > 100: ax.set_yscale('log'); ax.set_ylabel("Conversion Time, $t_\\mathrm{a}$ (days) [log scale]")
            ax.legend()
        fig.tight_layout(pad=0.5); memfile = io.BytesIO(); fig.savefig(memfile, format='png', dpi=300); memfile.seek(0)
        doc.add_picture(memfile, width=Cm(12)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        plt.close(fig)

    def save_report(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if self.ea is None: return messagebox.showwarning("No Data", "No data to export.")
        p = filedialog.asksaveasfilename(defaultextension=".docx", initialfile="TGA_Kinetic_Report_v1.0.docx", filetypes=[("Word Document", "*.docx")])
        if not p: return
        self._start_task(self._save_report_worker, p, on_success=lambda p: messagebox.showinfo("Success", f"Report saved to {p}"))

    def _save_report_worker(self, path):
        doc = Document(); doc.add_heading("Kinetic Analysis Report", 0)
        doc.add_heading("Activation Energy Calculation", level=1)
        equations = {"Friedman": "ln(β · dα/dT) = ln[A · f(α)] - Eₐ / (R · T)","KAS": "ln(β / T²) = ln(A·R / (Eₐ · g(α))) - Eₐ / (R · T)","OFW": "ln(β) = ln(A · Eₐ / (R · g(α))) - 5.331 - 1.052 · Eₐ / (R · T)","Vyazovkin": "Advanced non-linear isoconversional method"}
        equations = {"Friedman": "ln(β · dα/dt) = ln[A · f(α)] - Eₐ / (R · T)","KAS": "ln(β / T²) = ln(A·R / (Eₐ · g(α))) - Eₐ / (R · T)","OFW": "ln(β) = ln(A · Eₐ / (R · g(α))) - 5.331 - 1.052 · Eₐ / (R · T)","Vyazovkin": "Advanced non-linear isoconversional method"}
        doc.add_paragraph(f"The activation energy (Eₐ) was calculated using the {self.current_method} method.")
        doc.add_paragraph(f"Principle: {equations.get(self.current_method, 'N/A')}")
        doc.add_heading("Activation Energy vs. Conversion", level=2); self._add_plot_to_doc(doc, 'Ea')
        ea_display = self.ea.rename(columns=self.display_header_map)
        tbl = doc.add_table(rows=1, cols=len(ea_display.columns), style='Table Grid')
        for j, c in enumerate(ea_display.columns): tbl.cell(0, j).text = str(c)
        for _, r in ea_display.iterrows():
            row_cells = tbl.add_row().cells
            for j, v in enumerate(r): row_cells[j].text = f"{v:.5g}" if isinstance(v, (float, int)) else str(v)
        if self.conversion_time_df is not None and not self.conversion_time_df.empty:
            doc.add_heading("Conversion Time Prediction Results", level=1)
            model_source = self.last_fitted_model.get("source", "the fitted model")
            doc.add_paragraph(f"The isothermal conversion time (tₐ) was predicted by numerically integrating the rate equation derived from {model_source}. Uncertainty was propagated using a first-order Taylor expansion of the parameter covariance matrix to establish a 95% confidence interval (CI).")
            doc.add_heading("Predicted Conversion Time vs. Conversion (with 95% CI)", level=2); self._add_plot_to_doc(doc, 'ConversionTime')
        doc.save(path); return Path(path).name

    def save_latex_report(self):
        # This function remains largely unchanged but is provided for completeness.
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if self.ea is None: return messagebox.showwarning("No Data", "No data to export.")
        p = filedialog.asksaveasfilename(defaultextension=".tex", initialfile="TGA_Kinetic_Report_v1.0.tex", filetypes=[("LaTeX File", "*.tex")])
        p = filedialog.asksaveasfilename(defaultextension=".tex", initialfile="TGA_Kinetic_Report_v1.0.tex", filetypes=[("LaTeX File", "*.tex")])
        if not p: return
        self._start_task(self._save_latex_report_worker, p, on_success=lambda p: messagebox.showinfo("Success", f"LaTeX report saved to {p}"))

    def _save_latex_report_worker(self, path):
        # This function remains largely unchanged but is provided for completeness.
        p = Path(path)
        doc_parts = [ textwrap.dedent(r"""
                \documentclass{article}
                \usepackage{amsmath}\usepackage{booktabs}\usepackage{geometry}\usepackage{graphicx}
                \geometry{a4paper, margin=1in}
                \title{Kinetic Analysis Report}\author{TGA Kinetics Software v1.0}\date{\today}
                \title{Kinetic Analysis Report}\author{TGA Kinetics Software v1.0}\date{\today}
                \begin{document}\maketitle """), r"\section{Activation Energy Calculation}" ]
        equations = {"Friedman": r"$\ln(\beta \frac{\mathrm{d}\alpha}{\mathrm{d}T}) = \ln[A f(\alpha)] - \frac{E_\mathrm{a}}{RT}$", "KAS": r"$\ln(\frac{\beta}{T^2}) = \ln(\frac{AR}{E_\mathrm{a} g(\alpha)}) - \frac{E_\mathrm{a}}{RT}$", "OFW": r"$\ln(\beta) = \ln(\frac{A E_\mathrm{a}}{R g(\alpha)}) - 5.331 - 1.052 \frac{E_\mathrm{a}}{RT}$", "Vyazovkin": "Advanced non-linear isoconversional method"}
        equations = {"Friedman": r"$\ln(\frac{\mathrm{d}\alpha}{\mathrm{d}t}) = \ln[A f(\alpha)] - \frac{E_\mathrm{a}}{RT}$", "KAS": r"$\ln(\frac{\beta}{T^2}) = \ln(\frac{AR}{E_\mathrm{a} g(\alpha)}) - \frac{E_\mathrm{a}}{RT}$", "OFW": r"$\ln(\beta) = \ln(\frac{A E_\mathrm{a}}{R g(\alpha)}) - 5.331 - 1.052 \frac{E_\mathrm{a}}{RT}$", "Vyazovkin": "Advanced non-linear isoconversional method"}
        doc_parts.append(f"The activation energy ($E_\\mathrm{{a}}$) was calculated using the {self.current_method} method.")
        doc_parts.append(f"Principle: {equations.get(self.current_method, 'N/A')}")
        fig_path_ea = p.parent / "ea_plot.png"; fig, ax = plt.subplots(figsize=(12/2.54, 9/2.54))
        ax.errorbar(self.ea["alpha"], self.ea["Ea_kJ_per_mol"], yerr=self.ea["StdErr_kJ"], fmt='-o', capsize=3)
        ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel(r"Activation Energy, $E_\mathrm{a}$ (kJ mol$^{-1}$)")
        fig.tight_layout(pad=0.5); fig.savefig(fig_path_ea, dpi=300); plt.close(fig)
        abs_fig_path_ea_str = str(fig_path_ea.resolve()).replace('\\', '/')
        doc_parts.append(textwrap.dedent(f"""\n\\begin{{figure}}[h!]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{{{abs_fig_path_ea_str}}}\n\\caption{{Activation energy as a function of conversion ({self.current_method}).}}\n\\label{{fig:ea}}\n\\end{{figure}}"""))
        ea_latex = self.ea.rename(columns=self.latex_header_map)
        doc_parts.append(_df_to_latex_table(ea_latex.head(15), f"Activation energy data (first 15 points).", "tab:ea"))
        if self.conversion_time_df is not None and not self.conversion_time_df.empty:
            doc_parts.append(r"\section{Conversion Time Prediction Results}")
            model_source = self.last_fitted_model.get("source", "the fitted model")
            doc_parts.append(f"The isothermal conversion time ($t_\\mathrm{{a}}$) was predicted by integrating the rate equation from {model_source}. Uncertainty was propagated using the parameter covariance matrix to establish a 95\\% confidence interval.")
            fig_path_life = p.parent / "conversion_time_plot.png"; fig, ax = plt.subplots(figsize=(12/2.54, 9/2.54))
            ax.plot(self.conversion_time_df["alpha"], self.conversion_time_df["t_alpha_days"], "o-", label="Prediction")
            ax.fill_between(self.conversion_time_df["alpha"], self.conversion_time_df["ci_lower_days"], self.conversion_time_df["ci_upper_days"], alpha=0.2, label="95\\% CI")
            ax.set_xlabel(r"Conversion, $\alpha$"); ax.set_ylabel("Conversion Time, $t_\\mathrm{a}$ (days)")
            min_val, max_val = self.conversion_time_df["t_alpha_days"].min(), self.conversion_time_df["t_alpha_days"].max()
            if min_val > 0 and max_val / min_val > 100: ax.set_yscale('log'); ax.set_ylabel("Conversion Time, $t_\\mathrm{a}$ (days) [log scale]")
            ax.legend(); fig.tight_layout(pad=0.5); fig.savefig(fig_path_life, dpi=300); plt.close(fig)
            abs_fig_path_life_str = str(fig_path_life.resolve()).replace('\\', '/')
            doc_parts.append(textwrap.dedent(f"""\n\\begin{{figure}}[h!]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{{{abs_fig_path_life_str}}}\n\\caption{{Predicted conversion time with 95\\% confidence interval.}}\n\\label{{fig:life}}\n\\end{{figure}}"""))
        doc_parts.append(r"\end{document}")
        with open(p, 'w', encoding='utf-8') as f: f.write("\n\n".join(doc_parts))
        return p.name
    
    def save_tables(self):
        if not self.dfs: return messagebox.showwarning("No Data", "No data available to export. Please import and process files first.")
        file_path = filedialog.asksaveasfilename(title="Export All Tables", initialfile="Analysis_Results.xlsx", filetypes=[("Excel file", "*.xlsx")])
        if not file_path: return
        p = Path(file_path)
        try:
            with pd.ExcelWriter(p, engine='openpyxl') as writer:
                if self.ea is not None: self.ea.to_excel(writer, sheet_name="Ea_results", index=False)
                if self.aT is not None: self.aT.to_excel(writer, sheet_name="T_alpha_results", index=False)
                if self.xy is not None: self.xy.to_excel(writer, sheet_name="XY_linear_data", index=False)
                if self.conversion_time_df is not None: self.conversion_time_df.to_excel(writer, sheet_name="conversion_time_results", index=False)
            messagebox.showinfo("Success", f"All tables saved to:\n{p}")
        except Exception as e: messagebox.showerror("Error Saving", f"Could not save file. Error: {e}")

    def _calculate_tangent_intersection(self, x, y, peak_idx, direction):
        try:
            first_deriv = np.gradient(y, x)
            if direction == 'onset': inflection_idx = np.argmin(first_deriv[:peak_idx]) if peak_idx > 0 else 0
            else: inflection_idx = np.argmax(first_deriv[peak_idx:]) + peak_idx if peak_idx < len(first_deriv) -1 else len(first_deriv) -1
            slope = first_deriv[inflection_idx]; x1, y1 = x[inflection_idx], y[inflection_idx]
            if direction == 'onset': base_y_vals = y[:len(y)//20]; base_y = np.mean(base_y_vals) if len(base_y_vals) > 0 else y[0]
            else: base_y_vals = y[-len(y)//20:]; base_y = np.mean(base_y_vals) if len(base_y_vals) > 0 else y[-1]
            if abs(slope) < EPS: return x1
            return x1 + (base_y - y1) / slope
        except (ValueError, IndexError): return x[peak_idx]

    def run_tg_dtg_analysis(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if not self.dfs: return messagebox.showwarning("No Data", "Please import and process data first.")
        self._start_task(self._run_tg_dtg_analysis_worker, on_success=self._on_tg_dtg_analysis_success)

    def _run_tg_dtg_analysis_worker(self):
        results = {}
        for beta, df in self.dfs.items():
            try:
                peak_idx, T_peak = df['DTG_min'].idxmin(), df['Temp_K'].loc[df['DTG_min'].idxmin()]
                L_max = -df['DTG_min'].loc[peak_idx] 
                T_onset = self._calculate_tangent_intersection(df['Temp_K'].values, df['TG_pct'].values, peak_idx, 'onset')
                T_end = self._calculate_tangent_intersection(df['Temp_K'].values, df['TG_pct'].values, peak_idx, 'endset')
                tg_onset, tg_end = np.interp(T_onset, df['Temp_K'], df['TG_pct']), np.interp(T_end, df['Temp_K'], df['TG_pct'])
                results[f"{beta:.2f} K/min"] = {"T_onset (K)": T_onset, "T_peak (K)": T_peak, "T_end (K)": T_end, "ML (%)": tg_onset - tg_end, "L_max (%/min)": L_max}
            except Exception as e: print(f"Could not process beta={beta} for TG/DTG: {e}")
        if not results: raise ValueError("TG/DTG parameter calculation failed for all files.")
        return pd.DataFrame(results).T

    def _on_tg_dtg_analysis_success(self, results_df):
        self.status_var.set("✓ TG/DTG Parameter Analysis complete.")
        self.show_parameter_table(results_df, "TG/DTG Parameters", lambda: self.plot_tg_dtg(results_df))

    def run_dsc_analysis(self):
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if not self.dfs: return messagebox.showwarning("No Data", "Please import and process data first.")
        goal_dlg = DSCGoalDialog(self); self.wait_window(goal_dlg)
        if not goal_dlg.result: return
        goal, use_mass_norm = goal_dlg.result
        if use_mass_norm:
            self.sample_masses.clear()
            for beta in self.dfs.keys():
                mass = simpledialog.askfloat("Initial Sample Mass", f"Enter initial mass (mg) for experiment at {beta:.2f} K/min:", parent=self, initialvalue=10.0, minvalue=1e-6)
                if mass is None: return messagebox.showwarning("Cancelled", "Mass input cancelled. Aborting DSC Analysis.")
                self.sample_masses[beta] = mass if mass > 0 else 10.0
        baseline_method = "standard"
        if goal == 'enthalpy':
            baseline_dlg = DSCChoiceDialog(self); self.wait_window(baseline_dlg)
            if not baseline_dlg.result: return
            baseline_method = baseline_dlg.result
        if baseline_method == "interactive" and goal == 'enthalpy': self.run_dsc_interactive_workflow(goal)
        else: self._start_task(self._run_dsc_analysis_worker, goal, 'standard', on_success=self._on_dsc_analysis_success)

    def run_dsc_interactive_workflow(self, goal):
        interactive_baselines = {}
        try:
            for beta, df in self.dfs.items():
                mass = self.sample_masses.get(beta, 1.0)
                df_norm = df.copy(); df_norm['DSC_norm'] = df_norm['DSC'] / mass
                plotter = InteractivePlotter(self, df_norm, beta, mass); self.wait_window(plotter)
                if plotter.result and plotter.baseline_points: interactive_baselines[beta] = plotter.baseline_points
                else: messagebox.showwarning("Cancelled", "Interactive baseline selection was cancelled. Aborting analysis.", parent=self); return
            self._start_task(self._run_dsc_analysis_worker, goal, 'interactive', on_success=self._on_dsc_analysis_success, predefined_baselines=interactive_baselines)
        except Exception as e: messagebox.showerror("Error", f"An error occurred during the interactive process: {e}", parent=self)

    def _run_dsc_analysis_worker(self, goal, baseline_method, predefined_baselines=None):
        results, baselines = {}, {}; baselines = predefined_baselines if predefined_baselines else {}
        for beta, df in self.dfs.items():
            mass = self.sample_masses.get(beta, 1.0)
            df_norm = df.copy(); df_norm['DSC_norm'] = df_norm['DSC'] / mass
            peak_idx = df_norm['DSC_norm'].idxmin() if df_norm['DSC_norm'].min() < -df_norm['DSC_norm'].max() else df_norm['DSC_norm'].idxmax()
            params = {"T_peak (K)": df_norm['Temp_K'].loc[peak_idx]}
            params["T_onset (K)"] = self._calculate_tangent_intersection(df_norm['Temp_K'].values, df_norm['DSC_norm'].values, peak_idx, 'onset')
            params["T_end (K)"] = self._calculate_tangent_intersection(df_norm['Temp_K'].values, df_norm['DSC_norm'].values, peak_idx, 'endset')
            if goal == 'enthalpy':
                if baseline_method == "interactive":
                    if beta in baselines: (T_start, y_start), (T_end, y_end) = baselines[beta]
                    else: continue
                else:
                    y_start, y_end = np.interp(params["T_onset (K)"], df_norm['Temp_K'], df_norm['DSC_norm']), np.interp(params["T_end (K)"], df_norm['Temp_K'], df_norm['DSC_norm'])
                    T_start, T_end = params["T_onset (K)"], params["T_end (K)"]
                    baselines[beta] = [(T_start, y_start), (T_end, y_end)]
                mask = (df_norm['Temp_K'] >= T_start) & (df_norm['Temp_K'] <= T_end)
                sub_df = df_norm[mask]; baseline_vals = np.linspace(y_start, y_end, len(sub_df))
                dsc_corrected = sub_df['DSC_norm'] - baseline_vals
                params["ΔH (J/g)"] = integrate.trapezoid(dsc_corrected, sub_df['Time_min'] * 60)
            results[f"{beta:.2f} K/min"] = params
        if not results: raise ValueError("DSC parameter calculation failed.")
        return {"results_df": pd.DataFrame(results).T, "baselines": baselines}

    def _on_dsc_analysis_success(self, data):
        self.status_var.set("✓ DSC Parameter Analysis complete.")
        self.show_parameter_table(data["results_df"], "DSC Parameters", lambda: self.plot_dsc(data["results_df"], data["baselines"]))

    def show_parameter_table(self, results_df, title, plot_callback):
        win = tk.Toplevel(self); win.title(title); win.geometry("900x500")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico"); win.iconbitmap(icon_path)
        except Exception: pass
        frame = ttk.Frame(win, padding="20"); frame.pack(fill="both", expand=True)
        btn_frame = ttk.Frame(frame); btn_frame.pack(fill="x", pady=5)
        ttk.Label(btn_frame, text=title, font=("Segoe UI", 12, "bold")).pack(side="left")
        export_command = self._export_plot_data_tgdtg if title == "TG/DTG Parameters" else self._export_plot_data_dsc if title == "DSC Parameters" else None
        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(results_df.reset_index().rename(columns={'index':'Experiment'}), title), style="Secondary.TButton").pack(side="right", padx=5)
        if export_command: ttk.Button(btn_frame, text="Export Plot Data", command=export_command, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Results", command=plot_callback, style="Secondary.TButton").pack(side="right", padx=5)
        df_display = results_df.reset_index().rename(columns={'index': 'Experiment'}); cols = list(df_display.columns)
        tv = ttk.Treeview(frame, columns=cols, show="headings", style="Modern.Treeview")
        for c in cols: tv.heading(c, text=c); tv.column(c, width=120, anchor='w' if c == 'Experiment' else 'e')
        for _, row in df_display.iterrows():
            values = [row.iloc[0]] + [f"{v:.4f}" if isinstance(v, (float, int)) else v for v in row.iloc[1:]]
            tv.insert("", "end", values=values)
        tv.pack(fill="both", expand=True, pady=10)

    def _export_plot_data_tgdtg(self):
        if not self.dfs: return messagebox.showwarning("No Data", "No data to export.")
        all_data = [df[['Temp_K', 'TG_pct', 'DTG_min']].assign(**{'Heating_Rate (K/min)': beta}) for beta, df in self.dfs.items()]
        export_df = pd.concat(all_data, ignore_index=True).rename(columns={'DTG_min': 'Mass_Loss_Rate (%/min)'})
        self._export_table(export_df, "TG_DTG_Plot_Data")
        
    def _export_plot_data_dsc(self):
        if not self.dfs: return messagebox.showwarning("No Data", "No data to export.")
        all_data = []
        for beta, df in self.dfs.items():
            mass = self.sample_masses.get(beta, 1.0); temp_df = df[['Temp_K', 'DSC']].copy()
            temp_df['DSC_Normalized'] = temp_df['DSC'] / mass; temp_df['Heating_Rate (K/min)'] = beta; temp_df['Initial_Mass (mg)'] = mass
            all_data.append(temp_df)
        self._export_table(pd.concat(all_data, ignore_index=True), "DSC_Plot_Data")

    def _export_plot_data_cka(self):
        if not self.cka_results or not self.dfs: return messagebox.showwarning("No Data", "No CKA results to export.")
        all_data = []
        for beta_k_min, df in self.dfs.items():
            temp_df = df[['Temp_K', 'alpha']].copy(); temp_df['Heating_Rate (K/min)'] = beta_k_min
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            temp_df['Experimental_Rate (min^-1)'] = df['dAdT'] * beta_k_min
            ### MODIFICATION END ###
            
            temp_df['CKA_Fit_Rate (min^-1)'] = self._predict_rate_from_fit(self.cka_results, df)
            all_data.append(temp_df)
        self._export_table(pd.concat(all_data, ignore_index=True), "CKA_Fit_Comparison_Data")
    
    def _export_plot_data_autocatalytic(self):
        if not self.autocatalytic_results or not self.dfs: return messagebox.showwarning("No Data", "No Autocatalytic Fit results to export.")
        all_data = []
        for beta_k_min, df in self.dfs.items():
            temp_df = df[['Temp_K', 'alpha']].copy(); temp_df['Heating_Rate (K/min)'] = beta_k_min
            
            ### MODIFICATION START: Renamed dadt -> dAdT ###
            temp_df['Experimental_Rate (min^-1)'] = df['dAdT'] * beta_k_min
            ### MODIFICATION END ###
            
            dadt_model_per_sec = self._get_autocatalytic_rate(df['alpha'].values, df['Temp_K'].values, self.autocatalytic_results)
            temp_df['Fit_Rate (min^-1)'] = dadt_model_per_sec * 60 if dadt_model_per_sec is not None else np.nan
            dAdt_model_per_sec = self._get_autocatalytic_rate(df['alpha'].values, df['Temp_K'].values, self.autocatalytic_results)
            temp_df['Fit_Rate (min^-1)'] = dAdt_model_per_sec * 60 if dAdt_model_per_sec is not None else np.nan
            all_data.append(temp_df)
        self._export_table(pd.concat(all_data, ignore_index=True), f"{self.autocatalytic_results['model_name']}_Fit_Comparison_Data")
        
    def plot_tg_dtg(self, results_df):
        fig, ax1 = plt.subplots(figsize=(7, 5)); ax2 = ax1.twinx(); ax1.twinx_ax = ax2
        for i, beta in enumerate(self.dfs.keys()):
            df, label = self.dfs[beta], f"{beta:.1f} K/min"
            ax1.plot(df['Temp_K'], df['TG_pct'], '-', label=f"TG ({label})")
            ax2.plot(df['Temp_K'], -df['DTG_min'], '--', label=f"DTG ({label})")
        ax1.set_xlabel("Temperature (K)"); ax1.set_ylabel("TG (%)", color=NATURE_COLORS[0]); ax2.set_ylabel("Mass Loss Rate (%/min)", color=NATURE_COLORS[1])
        ax1.tick_params(axis='y', labelcolor=NATURE_COLORS[0]); ax2.tick_params(axis='y', labelcolor=NATURE_COLORS[1])
        ax1.margins(x=0); self.add_smart_legend(fig, ax1, num_items=len(self.dfs) * 2); fig.tight_layout()
        self._create_plot_window(fig, "TG/DTG Analysis Plot")

    def plot_dsc(self, results_df, baselines):
        fig, ax = plt.subplots(figsize=(7, 5)); num_curves = 0
        for i, (beta_str, row) in enumerate(results_df.iterrows()):
            target_beta = float(beta_str.split(' ')[0])
            actual_beta_key = min(self.dfs.keys(), key=lambda k: abs(k - target_beta))
            if abs(actual_beta_key - target_beta) / target_beta > 0.001: continue
            df, mass = self.dfs[actual_beta_key], self.sample_masses.get(actual_beta_key, 1.0)
            ax.plot(df['Temp_K'], df['DSC'] / mass, '-', label=beta_str); num_curves += 1
            if actual_beta_key in baselines and baselines[actual_beta_key]:
                (T_start, y_start), (T_end, y_end) = baselines[actual_beta_key]
                ax.plot([T_start, T_end], [y_start, y_end], 'k--', lw=1)
                mask = (df['Temp_K'] >= T_start) & (df['Temp_K'] <= T_end)
                sub_df = df[mask]; baseline_vals = np.linspace(y_start, y_end, len(sub_df))
                ax.fill_between(sub_df['Temp_K'], sub_df['DSC'] / mass, baseline_vals, alpha=0.3)
        ax.set_xlabel("Temperature (K)"); ax.set_ylabel("Heat flow (a.u.)"); ax.set_yticks([]); ax.margins(x=0)
        self.add_smart_legend(fig, ax, num_curves); fig.tight_layout()
        self._create_plot_window(fig, "DSC Analysis Plot", is_thermo_plot=True)

    ### NEW FUNCTION START: TG vs TG PLOT ###
    def plot_tg_fit_comparison(self):
        if not self.last_fitted_model:
            return messagebox.showwarning("No Model", "A kinetic model must be fitted first.")
        
        fig, ax = plt.subplots(figsize=(6, 5))
        num_curves = 0
        
        for beta_k_min, df_exp in self.dfs.items():
            # Plot Experimental TG
            p = ax.plot(df_exp['Temp_K'], df_exp['TG_pct'], '--', alpha=0.7, label=f"β={beta_k_min:.1f} (Exp.)")
            exp_color = p[0].get_color()
            num_curves += 1
            
            # Predict Model TG
            try:
                df_model = self.predict_tg_curve_from_model(df_exp, beta_k_min, self.last_fitted_model)
                ax.plot(df_model['Temp_K'], df_model['TG_pct_model'], '-', color=exp_color, lw=2.0, label=f"β={beta_k_min:.1f} (Fit)")
                num_curves += 1
            except Exception as e:
                print(f"Could not predict TG curve for beta={beta_k_min}: {e}")
        
        ax.set_xlabel("Temperature (K)"); ax.set_ylabel(r"TG (%)")
        ax.margins(x=0); self.add_smart_legend(fig, ax, num_curves); fig.tight_layout()
        self._create_plot_window(fig, f"Model TG vs. Experimental TG ({self.last_fitted_model['source']})")

    def export_tg_fit_table(self):
        if not self.last_fitted_model:
            return messagebox.showwarning("No Model", "A kinetic model must be fitted first.")

        all_data = []
        for beta_k_min, df_exp in self.dfs.items():
            try:
                df_model = self.predict_tg_curve_from_model(df_exp, beta_k_min, self.last_fitted_model)
                
                # Merge experimental and model data
                df_exp_mini = df_exp[['Temp_K', 'alpha', 'TG_pct']].rename(columns={'alpha': 'alpha_exp', 'TG_pct': 'TG_pct_exp'})
                df_model_mini = df_model[['Temp_K', 'alpha_model', 'TG_pct_model']]
                
                df_merged = pd.merge_asof(df_exp_mini.sort_values('Temp_K'), 
                                          df_model_mini.sort_values('Temp_K'), 
                                          on='Temp_K', 
                                          direction='nearest',
                                          suffixes=('_exp', '_model'))
                
                df_merged['Heating_Rate (K/min)'] = beta_k_min
                df_merged['Model_Source'] = self.last_fitted_model['source']
                all_data.append(df_merged)
            except Exception as e:
                print(f"Could not export TG curve for beta={beta_k_min}: {e}")
        
        if not all_data:
            return messagebox.showerror("Error", "Could not generate any data to export.")
            
        export_df = pd.concat(all_data, ignore_index=True)
        self._export_table(export_df, f"{self.last_fitted_model['model_name']}_TG_Fit_Comparison_Data")

    def _get_model_dAdT(self, T_K, alpha, model_info, beta_K_per_s):
        """
        Internal helper for the ODE solver.
        Calculates d(alpha)/dT for a given model, T, and alpha.
        """
        # We need d(alpha)/dt in (1/sec) to match beta_K_per_s
        dAdt_per_sec = 0.0
        
        model_name = model_info.get('model_name', '').upper()
        raw_params = model_info.get('raw_params')
        if raw_params is None: return 0.0

        alpha = np.clip(alpha, EPS, 1 - EPS)
        a1 = 1 - alpha
        
        with np.errstate(all='ignore'):
            if model_name in ["CKA", "SB", "CKA (FIXED EA)"]:
                m, n, p_, logA, Ea_J = raw_params
                A_min = np.exp(logA)
                f_alpha = (alpha**m) * (a1**n) * ((-np.log(a1))**p_)
                dAdt_per_min = A_min * np.exp(-Ea_J / (R * T_K)) * f_alpha
                dAdt_per_sec = dAdt_per_min / 60.0
            
            elif model_name == "GLOBALFIT":
                m, n, p_, logA = raw_params
                A_min = np.exp(logA)
                f_alpha = (alpha**m) * (a1**n) * ((-np.log(a1))**p_)
                ea_ser = self.ea.set_index("alpha")["Ea_kJ_per_mol"] * 1000
                Ea_interp = np.interp(alpha, ea_ser.index, ea_ser.values)
                dAdt_per_min = A_min * np.exp(-Ea_interp / (R * T_K)) * f_alpha
                dAdt_per_sec = dAdt_per_min / 60.0
            
            elif model_name == "KAMAL-SOUROUR":
                logA1, E1_J, logA2, E2_J, m, n = raw_params
                k1 = np.exp(logA1) * np.exp(-E1_J / (R * T_K))
                k2 = np.exp(logA2) * np.exp(-E2_J / (R * T_K))
                dAdt_per_min = (k1 + k2 * (alpha**m)) * (a1**n)
                dAdt_per_sec = dAdt_per_min / 60.0
            elif model_name == "GAI":
                logA, E_J, n1, z0, n2 = raw_params
                k = np.exp(logA) * np.exp(-E_J / (R * T_K))
                f_alpha = (a1**n1) * (z0 + alpha**n2)
                dAdt_per_min = k * f_alpha
                dAdt_per_sec = dAdt_per_min / 60.0
            elif model_name == "PAR":
                logA1, E1_J, n1, logA2, E2_J, n2 = raw_params
                k1 = np.exp(logA1) * np.exp(-E1_J / (R * T_K))
                k2 = np.exp(logA2) * np.exp(-E2_J / (R * T_K))
                dAdt_per_min = k1 * (a1**n1) + k2 * (a1**n2)
                dAdt_per_sec = dAdt_per_min / 60.0
        
        # d(alpha)/dT = d(alpha)/dt / (dT/dt)
        if beta_K_per_s <= 0: return 0.0
        return dAdt_per_sec / beta_K_per_s

    def predict_tg_curve_from_model(self, df_experimental, beta_K_min, model_info):
        """
        Uses an ODE solver to predict the TG curve from a fitted model.
        """
        beta_K_per_s = beta_K_min / 60.0
        
        # Get initial conditions
        T_span = (df_experimental['Temp_K'].iloc[0], df_experimental['Temp_K'].iloc[-1])
        T_eval = df_experimental['Temp_K'].values
        alpha_initial = df_experimental['alpha'].iloc[0]
        
        m0, m_inf = self.mass_loss_parameters.get(beta_K_min, (100.0, 0.0))
        if (m0 - m_inf) == 0:
            m0, m_inf = (100.0, 0.0)

        # Define the function for the ODE solver
        # The solver's function must take (t, y) which for us is (T, alpha)
        ode_func = lambda T, alpha: self._get_model_dAdT(T, alpha, model_info, beta_K_per_s)

        # Solve the differential equation
        sol = solve_ivp(
            ode_func, 
            T_span, 
            [alpha_initial], 
            t_eval=T_eval,
            method='RK45', 
            vectorized=False
        )
        
        if not sol.success:
            raise RuntimeError(f"ODE solver failed for beta={beta_K_min}: {sol.message}")

        alpha_model = sol.y[0].clip(0, 1)
        
        # Reconstruct the TG curve from the predicted alpha
        tg_pct_model = m0 - (alpha_model * (m0 - m_inf))
        
        return pd.DataFrame({
            'Temp_K': sol.t,
            'alpha_model': alpha_model,
            'TG_pct_model': tg_pct_model
        })
    
    ### MODIFICATION START: NEW SECTION FOR NON-PARAMETRIC ANALYSIS (NPA) ###
    
    def run_npa_analysis(self):
        """
        Launches the Non-Parametric Analysis (NPA) task.
        This calculates Z(α) = A(α)f(α) based on the isoconversional Ea(α).
        """
        if self.is_task_running: return messagebox.showwarning("Busy", "Please wait for the current task to finish.")
        if self.ea is None or self.ea.empty: 
            return messagebox.showwarning("No Data", "Please run an isoconversional analysis (e.g., Vyazovkin) first to get the Ea curve.")
        
        prompt = "Enter α range for NPA (e.g., 0.1,0.9):"; initial = f"{self.alphas.min()},{self.alphas.max()}"
        rng = simpledialog.askstring("NPA Settings", prompt, initialvalue=initial, parent=self)
        if not rng: return
        try:
            a_min, a_max = map(float, rng.split(','))
            if not (0 < a_min < a_max < 1): raise ValueError
            fit_alphas = np.round(np.arange(a_min, a_max + EPS, 0.01), 4)
        except (ValueError, IndexError): return messagebox.showerror("Error", "Invalid α range format.")
        
        self._start_task(self._run_npa_worker, self.dfs, self.ea, fit_alphas, on_success=self._on_npa_success)

    def _run_npa_worker(self, dfs, ea_df, alphas):
        """
        Worker function to calculate Z(α) = (dα/dt) / exp(-Ea(α)/RT).
        """
        ea_curve = ea_df.set_index('alpha')['Ea_kJ_per_mol'] * 1000 # Ea in J/mol
        
        all_results = []
        for beta_k_min, df in dfs.items():
            beta_k_s = beta_k_min / 60.0
            
            # Interpolate Ea(α) for the full alpha range of this df
            alpha_clipped = np.clip(df["alpha"].values, EPS, 1 - EPS)
            Ea_interp = np.interp(alpha_clipped, ea_curve.index, ea_curve.values)
            
            # Calculate d(alpha)/dt in (1/sec)
            dadt_per_sec = df['dAdT'] * beta_k_s
            
            # Calculate Z(α)
            exp_term = np.exp(-Ea_interp / (R * df['Temp_K']))
            Z_alpha = dadt_per_sec / (exp_term + EPS) # Z(α) in 1/sec
            
            # Filter for the selected alpha range and valid numbers
            mask = (alpha_clipped >= alphas.min()) & (alpha_clipped <= alphas.max()) & \
                   np.isfinite(Z_alpha) & (Z_alpha > 0)
            
            all_results.append(pd.DataFrame({
                'alpha': alpha_clipped[mask],
                'Z_alpha_per_s': Z_alpha[mask]  
            }))
        
        if not all_results:
            raise ValueError("Could not calculate Z(α). Check data and Ea curve.")
            
        # Combine all data and average Z(α) for each alpha
        combined_df = pd.concat(all_results)
        # Bin alphas to get a clean average curve
        bins = np.linspace(alphas.min(), alphas.max(), 51) # 50 bins
        combined_df['alpha_bin'] = pd.cut(combined_df['alpha'], bins=bins, labels=bins[:-1] + (bins[1]-bins[0])/2)
        
        avg_df = combined_df.groupby('alpha_bin', observed=False, as_index=False).agg( # <--- ### MODIFICATION ###
            Z_alpha_per_s=('Z_alpha_per_s', 'mean')
        ).dropna()
        
        avg_df.rename(columns={'alpha_bin': 'alpha'}, inplace=True)
        avg_df['log_Z_alpha'] = np.log10(avg_df['Z_alpha_per_s'])
        
        return avg_df.astype(float) # Ensure 'alpha' is float

    def _on_npa_success(self, results_df):
        """
        Callback when NPA calculation is complete.
        Stores results and populates 'last_fitted_model' for prediction.
        """
        self.npa_results = results_df
        self.status_var.set("✓ Non-Parametric Analysis (NPA) complete. Showing results.")
        
        # This is the "Holy Grail" model: a non-parametric model
        # based on the isoconversional Ea(α) and the calculated Z(α).
        self.last_fitted_model = {
            "model_name": "NPA",
            "source": "Non-Parametric Analysis",
            "Ea_curve": self.ea.set_index('alpha')['Ea_kJ_per_mol'], # Ea in kJ/mol
            "Z_curve": self.npa_results.set_index('alpha')['Z_alpha_per_s'] # Z(α) in 1/sec
        }
        
        self.show_npa_results_window()

    def show_npa_results_window(self):
        """
        Displays the results of the Non-Parametric Analysis (NPA).
        """
        if self.npa_results is None or self.npa_results.empty: return # <--- ### MODIFICATION ###
        win = tk.Toplevel(self, bg=self.BG); win.title("Non-Parametric Analysis (NPA) Results")
        try:
            icon_path = self.resource_path("BIT_Kinetics_Icon_Tight.ico")
            win.iconbitmap(icon_path)
        except Exception: pass
        win.geometry("800x550")

        frame = ttk.Frame(win, padding="20"); frame.pack(fill="both", expand=True)
        btn_frame = ttk.Frame(frame); btn_frame.pack(pady=5, fill='x')
        ttk.Label(btn_frame, text="NPA Results: Z(α) = A(α)f(α)", font=("Segoe UI", 12, "bold")).pack(side="left")
        
        ttk.Button(btn_frame, text="Export Table", command=lambda: self._export_table(self.npa_results, "npa_results"), style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Predict Conversion Time...", command=self.launch_prediction_dialog, style="Secondary.TButton").pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Plot Results", command=self.plot_npa_results, style="Secondary.TButton").pack(side="right", padx=5)

        fig_eq = plt.figure(figsize=(6, 1.5), facecolor=self.BG)
        fig_eq.text(0.5, 0.5, r"$Z(\alpha) = (d\alpha/dt) / \exp(-E_a(\alpha)/RT)$", va='center', ha='center', fontsize=12, color=self.FG)
        canvas = FigureCanvasTkAgg(fig_eq, master=frame); canvas.draw()
        canvas.get_tk_widget().pack(fill='x', expand=False, pady=(10,15)); plt.close(fig_eq)

        tv_cols = ("α", "Z(α) (s⁻¹)", "log₁₀(Z)")
        tv = ttk.Treeview(frame, columns=tv_cols, show="headings", style="Modern.Treeview"); tv.pack(fill="both", expand=True)
        for col_id in tv_cols: tv.heading(col_id, text=col_id); tv.column(col_id, anchor="center")
        
        for _, row in self.npa_results.iterrows():
            tv.insert("", "end", values=(f"{row['alpha']:.4f}", f"{row['Z_alpha_per_s']:.3e}", f"{row['log_Z_alpha']:.3f}"))

    def plot_npa_results(self):
        """
        Plots the log(Z) vs. alpha curve.
        """
        if self.npa_results is None or self.npa_results.empty: return # <--- ### MODIFICATION ###
        fig, ax = plt.subplots(figsize=(6, 5))
        
        ax.plot(self.npa_results['alpha'], self.npa_results['log_Z_alpha'], 'o-', label=r"$\log_{10}(Z(\alpha))$")
        
        ax.set_xlabel(r"Conversion, $\alpha$")
        ax.set_ylabel(r"$\log_{10}(Z(\alpha))$ where $Z(\alpha) = A(\alpha)f(\alpha)$")
        ax.legend()
        ax.margins(x=0)
        fig.tight_layout()
        self._create_plot_window(fig, "NPA Function Plot")

    ### MODIFICATION END ###
    ### NEW FUNCTION END ###

class DSCGoalDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent); self.title("Choose DSC Analysis Goal"); self.result = None
        try: icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico"); self.iconbitmap(icon_path)
        except Exception: pass
        self.transient(parent); self.grab_set(); self.protocol("WM_DELETE_WINDOW", self.destroy)
        main_frame = ttk.Frame(self, padding="20"); main_frame.pack(fill="both", expand=True)
        ttk.Label(main_frame, text="What do you want to calculate?", font=("Segoe UI", 11, "bold")).pack(pady=10)
        self.goal_var = tk.StringVar(value="temps_only")
        ttk.Radiobutton(main_frame, text="Calculate Peak Temperatures Only (T_onset, T_peak)", variable=self.goal_var, value="temps_only").pack(anchor="w", padx=20, pady=5)
        ttk.Radiobutton(main_frame, text="Calculate Enthalpy (ΔH) from Pre-Normalized Data", variable=self.goal_var, value="enthalpy_normalized").pack(anchor="w", padx=20, pady=5)
        ttk.Radiobutton(main_frame, text="Calculate Enthalpy (ΔH) from Raw Data (Provide Mass)", variable=self.goal_var, value="enthalpy_raw").pack(anchor="w", padx=20, pady=5)
        btn_frame = ttk.Frame(main_frame); btn_frame.pack(fill="x", pady=(20, 0))
        ttk.Button(btn_frame, text="Next", command=self.on_ok, style="Primary.TButton").pack(side="right")
        ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton").pack(side="right", padx=(0, 10))
    def on_ok(self):
        choice = self.goal_var.get()
        if choice == 'temps_only': self.result = ('temps', False)
        elif choice == 'enthalpy_normalized': self.result = ('enthalpy', False)
        elif choice == 'enthalpy_raw': self.result = ('enthalpy', True)
        self.destroy()

class DSCChoiceDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent); self.title("Choose DSC Baseline Method"); self.result = None
        try: icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico"); self.iconbitmap(icon_path)
        except Exception: pass
        self.transient(parent); self.grab_set(); self.protocol("WM_DELETE_WINDOW", self.destroy)
        main_frame = ttk.Frame(self, padding="20"); main_frame.pack(fill="both", expand=True)
        ttk.Label(main_frame, text="Select a baseline method for enthalpy (ΔH) integration:", font=("Segoe UI", 11, "bold")).pack(pady=10)
        self.method_var = tk.StringVar(value="standard")
        ttk.Radiobutton(main_frame, text="Use Standard Baseline (Automated)", variable=self.method_var, value="standard").pack(anchor="w", padx=20, pady=5)
        ttk.Radiobutton(main_frame, text="Define Interactive Baseline (High Accuracy)", variable=self.method_var, value="interactive").pack(anchor="w", padx=20, pady=5)
        btn_frame = ttk.Frame(main_frame); btn_frame.pack(fill="x", pady=(20, 0))
        ttk.Button(btn_frame, text="OK", command=self.on_ok, style="Primary.TButton").pack(side="right")
        ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton").pack(side="right", padx=(0, 10))
    def on_ok(self): self.result = self.method_var.get(); self.destroy()

class InteractivePlotter(tk.Toplevel):
    def __init__(self, parent, df, beta, mass):
        super().__init__(parent); self.parent = parent; self.title(f"Interactive Baseline for {beta:.1f} K/min")
        self.protocol("WM_DELETE_WINDOW", self.destroy)
        try: icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico"); self.iconbitmap(icon_path)
        except Exception: pass
        self.df, self.beta, self.mass, self.result, self.baseline_points, self.clicks = df, beta, mass, None, [], []
        self.fig, self.ax = plt.subplots(); self.canvas = FigureCanvasTkAgg(self.fig, master=self)
        toolbar_frame = ttk.Frame(self); toolbar_frame.pack(side=tk.TOP, fill=tk.X)
        toolbar = NavigationToolbar2Tk(self.canvas, toolbar_frame); toolbar.update()
        self.canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        self.instructions_var = tk.StringVar(value="Click to define the start point of the baseline.")
        ttk.Label(self, textvariable=self.instructions_var, font=("Segoe UI", 10, "bold")).pack(pady=5)
        self.cid = self.canvas.mpl_connect('button_press_event', self.on_click); self.plot_initial()
    def plot_initial(self):
        self.ax.clear(); self.ax.plot(self.df['Temp_K'], self.df['DSC_norm'], label=f"{self.beta:.1f} K/min")
        self.ax.set_xlabel("Temperature (K)"); self.ax.set_ylabel("Heat flow (a.u.)"); self.ax.set_yticks([])
        self.ax.set_title("Click to Define Baseline Start & End"); self.canvas.draw()
    def on_click(self, event):
        if event.inaxes != self.ax: return
        self.clicks.append((event.xdata, event.ydata))
        if len(self.clicks) == 1:
            self.ax.axvline(x=event.xdata, color='r', linestyle='--'); self.instructions_var.set("Click to define the end point of the baseline."); self.canvas.draw()
        elif len(self.clicks) == 2:
            self.canvas.mpl_disconnect(self.cid); self.ax.axvline(x=event.xdata, color='g', linestyle='--')
            (T_start, _), (T_end, _) = self.clicks
            if T_start > T_end: T_start, T_end = T_end, T_start
            y_start, y_end = np.interp(T_start, self.df['Temp_K'], self.df['DSC_norm']), np.interp(T_end, self.df['Temp_K'], self.df['DSC_norm'])
            self.baseline_points = [(T_start, y_start), (T_end, y_end)]
            self.ax.plot([T_start, T_end], [y_start, y_end], 'k--', lw=2)
            self.instructions_var.set("Baseline defined. Calculating..."); self.canvas.draw(); self.calculate_and_close()
    def calculate_and_close(self):
        try:
            (T_start, y_start), (T_end, y_end) = self.baseline_points
            mask = (self.df['Temp_K'] >= T_start) & (self.df['Temp_K'] <= T_end)
            sub_df = self.df[mask].copy(); baseline_vals = np.linspace(y_start, y_end, len(sub_df))
            sub_df['DSC_corr'] = sub_df['DSC_norm'] - baseline_vals
            self.result = {"ΔH (J/g)": integrate.trapezoid(sub_df['DSC_corr'], sub_df['Time_min'] * 60)}
            self.after(500, self.destroy)
        except Exception as e: messagebox.showerror("Error", f"Could not calculate parameters: {e}", parent=self); self.destroy()

class AutocatalyticModelDialog(tk.Toplevel):
    def __init__(self, parent, models):
        super().__init__(parent); self.title("Autocatalytic Model Fit"); self.result = None
        try: icon_path = parent.resource_path("BIT_Kinetics_Icon_Tight.ico"); self.iconbitmap(icon_path)
        except Exception: pass
        self.transient(parent); self.grab_set(); self.protocol("WM_DELETE_WINDOW", self.destroy); self.configure(bg=parent.BG)
        main_frame = ttk.Frame(self, padding="20"); main_frame.pack(fill="both", expand=True)
        ttk.Label(main_frame, text="Select a model for multi-rate global fitting:", font=("Segoe UI", 12, "bold")).pack(pady=(0, 15), anchor="w")
        self.model_var = tk.StringVar(value="SB")
        for key, (name, eqn) in models.items():
            rb = ttk.Radiobutton(main_frame, text=name, variable=self.model_var, value=key); rb.pack(anchor="w", padx=10, pady=(8, 2))
            ttk.Label(main_frame, text=f"  Eq: {eqn}", font=("Segoe UI", 9, "italic"), foreground=parent.FG_SECONDARY).pack(anchor="w", padx=25, pady=(0, 8))
        ttk.Separator(main_frame, orient="horizontal").pack(fill="x", pady=15)
        self.cross_val_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(main_frame, text="Hold out one heating rate for cross-validation", variable=self.cross_val_var, style="Modern.TCheckbutton").pack(anchor="w", padx=10, pady=5)
        btn_frame = ttk.Frame(main_frame); btn_frame.pack(fill="x", pady=(20, 0))
        ttk.Button(btn_frame, text="Next", command=self.on_ok, style="Primary.TButton").pack(side="right")
        ttk.Button(btn_frame, text="Cancel", command=self.destroy, style="Secondary.TButton").pack(side="right", padx=(0, 10))
    def on_ok(self): self.result = {"model": self.model_var.get(), "cross_validation": self.cross_val_var.get()}; self.destroy()


if __name__ == "__main__":
    app = App()
    app.mainloop()   

